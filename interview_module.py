"""
CVOLVE PRO — AI Interview Practice Module
==========================================
Handles the full interview practice session:
 - Structured Q&A generation (Behavioral / Technical, 3 difficulty levels)
 - AI interviewer (question-by-question flow)
 - Voice answer recording + transcription (browser MediaRecorder via ST component)
 - AI evaluation (semantic, keyword, structure, completeness)
 - Downloadable feedback report (PDF + DOCX)
"""

import streamlit as st
import json
import re
import os
import time
from io import BytesIO
from datetime import datetime

import google.generativeai as genai
import openai
from streamlit import session_state as st_session
from tts_utils import tts_component_html

from credit_engine import (
    can_use_f2f,
    start_f2f_session,
    charge_f2f_block,
    refund_f2f_block,
    end_f2f_session,
)
import pricing

# ─────────────────────────────────────────────────────────────────────────────
# Gemini model (shared with cv_generator)
# ─────────────────────────────────────────────────────────────────────────────
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
_model = genai.GenerativeModel("gemini-2.5-flash")
openai.api_key = os.getenv("OPENAI_API_KEY")


# ─────────────────────────────────────────────────────────────────────────────
# Credit costs per duration
# ─────────────────────────────────────────────────────────────────────────────
DURATION_CREDITS = {
    "15 minutes": 5,
    "30 minutes": 8,
    "45 minutes": 12,
}

# Question counts: behavioral = 10 behavioral + 5 resume; technical = 10 technical + 5 resume
QUESTION_COUNTS = {
    "behavioral": {"behavioral": 10, "resume": 5},
    "technical": {"technical": 10, "resume": 5},
}


# ─────────────────────────────────────────────────────────────────────────────
# Demo AI fallback (works without any API key)
# ─────────────────────────────────────────────────────────────────────────────
def _demo_ai_response(prompt: str, json_mode: bool) -> str:
    import hashlib
    import json as _json

    seed = hashlib.md5(prompt.encode()).hexdigest()

    if "interview question bank" in prompt or "generate_structured_interview_qa" in prompt.lower():
        # Handle both label variants: "JOB DESCRIPTION (must anchor EVERY question):" and "JOB DESCRIPTION:"
        jd_label = "JOB DESCRIPTION (must anchor EVERY question):"
        if jd_label not in prompt:
            jd_label = "JOB DESCRIPTION:"
        res_label = "RESUME (for resume-specific questions):"
        if res_label not in prompt:
            res_label = "CANDIDATE RESUME:" if "CANDIDATE RESUME:" in prompt else "RESUME:"
        parts = prompt.split(jd_label)
        jd_part = parts[1].split(res_label)[0].strip() if len(parts) > 1 else ""
        resume_part = prompt.split(res_label)[-1].strip() if res_label in prompt else ""
        seed = hashlib.md5((jd_part + resume_part).encode()).hexdigest()
        rng = int(seed[:8], 16)

        jd_sentences = [s.strip() for s in jd_part.replace("\n", ". ").split(".") if len(s.strip()) > 15]
        if not jd_sentences:
            jd_sentences = [jd_part.strip()]

        is_technical_type = "technical / skills-based" in prompt.lower()

        # Detect difficulty level from prompt
        diff = "medium"
        if "difficulty level: easy" in prompt.lower():
            diff = "easy"
        elif "difficulty level: hard" in prompt.lower():
            diff = "hard"
        elif "difficulty level: medium" in prompt.lower():
            diff = "medium"

        cat1_count = 10
        res_count = 5
        needed_total = cat1_count + res_count

        demo_questions = []
        seen_questions = set()

        if is_technical_type:
            # ── TECHNICAL ROUND (Tech/Subject Expert) ─────────────────────────
            # Extract meaningful skill phrases from JD
            jd_phrases = []
            for s in jd_sentences:
                for phrase in s.split(","):
                    phrase = phrase.strip().strip("-•*")
                    if len(phrase) > 10:
                        # Clean phrase for grammatical template insertion
                        clean = phrase[0].lower() + phrase[1:] if phrase else phrase
                        clean = clean.rstrip(".")
                        jd_phrases.append(clean)

            tech_templates = {
                "easy": [
                    "Can you explain what {} means in the context of this role?",
                    "What foundational knowledge do you have about {}?",
                    "How would you define {} to someone unfamiliar with the field?",
                ],
                "medium": [
                    "Walk me through your hands-on experience with {}.",
                    "How would you approach {} in a real project?",
                    "What tools, frameworks, or methods do you use for {}?",
                ],
                "hard": [
                    "Describe a complex technical challenge you solved involving {}.",
                    "How would you optimize or troubleshoot {} under real constraints?",
                    "Explain the architectural trade-offs and advanced considerations for {}.",
                ],
            }
            pool = tech_templates.get(diff, tech_templates["medium"])
            for i, phrase in enumerate(jd_phrases[:12]):
                t = pool[(rng + i) % len(pool)].format(phrase)
                if t not in seen_questions:
                    seen_questions.add(t)
                    demo_questions.append({
                        "question": t,
                        "difficulty": diff,
                        "ideal_answer": f"I have practical experience with {phrase}. My approach involves understanding requirements, applying best practices, and iterating based on feedback. I focus on delivering reliable, maintainable results.",
                        "key_points": [f"Practical experience with {phrase}", "Best practices", "Problem-solving", "Quality focus", "Continuous improvement"],
                    })

            # Resume-based questions for technical round
            tech_resume_stems = [
                "Walk me through your most relevant project and how it connects to this role's technical requirements.",
                "What technical skill from your background is most relevant for this position?",
                "Describe a complex technical problem you solved in a previous role.",
                "How has your past technical experience prepared you for the technical challenges of this role?",
                "Tell me about a time you had to quickly learn a new technology for a project.",
                "Describe a technical process or workflow you improved significantly.",
            ]
            for stem in tech_resume_stems:
                if len(demo_questions) >= needed_total:
                    break
                if stem not in seen_questions:
                    seen_questions.add(stem)
                    demo_questions.append({
                        "question": stem,
                        "difficulty": diff,
                        "ideal_answer": "My technical background aligns well with this role. I have hands-on experience with relevant technologies and have delivered measurable results in past projects.",
                        "key_points": ["Technical depth", "Project experience", "Problem-solving", "Results orientation", "Continuous learning"],
                    })

        else:
            # ── BEHAVIORAL ROUND (HR Interview) ───────────────────────────────
            # Pure HR/personality questions — NO JD tech references
            hr_pools = {
                "easy": [
                    "Tell me about yourself and your background.",
                    "Why did you choose this career path?",
                    "What do you enjoy most about your work?",
                    "Describe your ideal work environment.",
                    "What are your biggest strengths?",
                ],
                "medium": [
                    "Tell me about yourself and why you are interested in this role.",
                    "Describe a time you worked successfully as part of a team.",
                    "Tell me about a time you handled a conflict with a colleague.",
                    "Describe a situation where you had to adapt to a significant change.",
                    "Tell me about a time you went above and beyond for a project or team.",
                    "Describe a time you received difficult feedback and how you responded.",
                    "How do you handle competing priorities and tight deadlines?",
                    "Tell me about a time you failed and what you learned from it.",
                ],
                "hard": [
                    "Tell me about yourself and why you are the right fit for this role.",
                    "Describe a time you influenced a decision without direct authority.",
                    "Tell me about a time you had to lead through ambiguity or uncertainty.",
                    "Describe a situation where you had to deliver difficult feedback to a peer or manager.",
                    "Tell me about a time you had to make an unpopular decision for the greater good.",
                    "Describe a time you had to manage a stakeholder who had conflicting expectations.",
                    "Tell me about a time you had to rebuild trust with a team member or client.",
                    "How do you approach your long-term career development?",
                ],
            }
            pool = hr_pools.get(diff, hr_pools["medium"])
            for q_text in pool:
                if len(demo_questions) >= cat1_count:
                    break
                if q_text not in seen_questions:
                    seen_questions.add(q_text)
                    demo_questions.append({
                        "question": q_text,
                        "difficulty": diff,
                        "ideal_answer": "I approach this by focusing on clear communication, empathy, and results. I believe in understanding the situation fully before acting, and I always follow up to ensure positive outcomes.",
                        "key_points": ["Self-awareness", "Communication", "Teamwork", "Problem-solving", "Growth mindset"],
                    })

            # Resume-based for behavioral (HR-style, linking to past experience)
            hr_resume_stems = [
                "Walk me through your career journey and how it led you to this role.",
                "What accomplishment from your past experience are you most proud of?",
                "Describe a challenge you faced in your previous role and how you overcame it.",
                "How has your previous experience prepared you for the responsibilities of this role?",
                "Tell me about a mentor or leader who shaped your professional approach.",
                "What areas of your professional skill set would you like to develop further?",
                "Describe a time you went beyond your job description to help your team.",
            ]
            for stem in hr_resume_stems:
                if len(demo_questions) >= needed_total:
                    break
                if stem not in seen_questions:
                    seen_questions.add(stem)
                    demo_questions.append({
                        "question": stem,
                        "difficulty": diff,
                        "ideal_answer": "My career journey has been driven by a passion for learning and growth. Each role has contributed to my skill set, and I am excited to bring this experience to the new challenges this position offers.",
                        "key_points": ["Career narrative", "Self-awareness", "Growth mindset", "Relevant experience", "Role alignment"],
                    })

        # Guarantee the advertised question count by padding with generic questions.
        generic_fillers = [
            "What motivates you to do your best work?",
            "Describe a skill you are currently developing.",
            "How do you stay updated with trends in your field?",
            "Describe your approach to continuous professional development.",
            "Describe a project you are proud of and the key decisions you made.",
            "How would you approach building something from scratch in this role?",
            "Tell me about a goal you set for yourself and how you achieved it.",
            "Describe a time you had to learn something new quickly to get a job done.",
            "How do you prioritize your work when everything is urgent?",
            "Describe a time you had to make a difficult choice with limited information.",
        ]
        generic_kp = ["Self-awareness", "Continuous learning", "Growth mindset", "Problem-solving", "Communication"]
        filler_i = 0
        while len(demo_questions) < needed_total and filler_i < 100:
            q_text = generic_fillers[filler_i % len(generic_fillers)]
            filler_i += 1
            if q_text not in seen_questions:
                seen_questions.add(q_text)
                demo_questions.append({
                    "question": q_text,
                    "difficulty": diff,
                    "ideal_answer": "I approach this by focusing on clear communication, empathy, and results. I believe in understanding the situation fully before acting, and I always follow up to ensure positive outcomes.",
                    "key_points": generic_kp,
                })

        import random as _random
        _rng = _random.Random(rng)
        shuffled = list(demo_questions)
        _rng.shuffle(shuffled)
        picked = shuffled[:needed_total]
        cat1_section = picked[:cat1_count]
        resume_section = picked[cat1_count:cat1_count + res_count]

        # Demo answers are generic placeholders; nudge users to ground them in
        # their own CV facts (real AI answers are already CV-personalized).
        if resume_part and resume_part.strip():
            _demo_pers_note = (
                "\n\n[💡 Personalise me: rewrite with a specific project, role, "
                "or metric from YOUR CV so this reads like your own answer.]"
            )
            for _dq in picked:
                if _dq.get("ideal_answer"):
                    _dq["ideal_answer"] = _dq["ideal_answer"] + _demo_pers_note

        if is_technical_type:
            demo = {"technical": cat1_section, "resume": resume_section}
        else:
            demo = {"general": cat1_section, "resume": resume_section}
        return _json.dumps(demo)

    if "evaluate" in prompt.lower() and ("answer" in prompt.lower() or "question" in prompt.lower()):
        try:
            key_points = []
            user_answer = ""
            question_text = ""
            ideal_answer = ""
            # Extract key_points JSON array
            before_answer = prompt.split("CANDIDATE'S ANSWER:")[0] if "CANDIDATE'S ANSWER:" in prompt else ""
            m_kp = re.search(r'\[.*?\]', before_answer)
            if m_kp:
                parsed = _json.loads(m_kp.group())
                if isinstance(parsed, list):
                    key_points = parsed
            # Extract user answer
            if "CANDIDATE'S ANSWER:" in prompt:
                user_answer = prompt.split("CANDIDATE'S ANSWER:", 1)[1].strip()
            # Extract question text
            mq = re.search(r'QUESTION:\s*(.+?)(?:\nSECTION:|$)', prompt)
            if mq:
                question_text = mq.group(1).strip()
            # Extract ideal answer
            mi = re.search(r'IDEAL ANSWER[^:]*:\s*(.+?)(?:\n\nKEY POINTS|\nKEY POINTS)', prompt, re.DOTALL)
            if mi:
                ideal_answer = mi.group(1).strip()

            answer_lower = user_answer.lower()
            n_words = len(user_answer.split())
            n_sentences = len(re.split(r'[.!?]+', user_answer.strip())) if user_answer.strip() else 0
            avg_words_per_sent = n_words / max(n_sentences, 1)

            if n_words < 5:
                return _json.dumps({
                    "score": 0, "meaning_match": 0, "keyword_coverage": 0,
                    "keywords_covered": [], "keywords_missed": key_points or [],
                    "structure_score": 0, "completeness_score": 0,
                    "clarity_score": 0, "relevance_score": 0, "depth_score": 0,
                    "confidence_indicators": [],
                    "strengths": [], "improvements": ["Your answer is too short. Provide a detailed response."],
                    "improved_answer": "Expand your answer with specific examples covering the key points listed above.",
                    "brief_feedback": "Answer too short. Please write at least a paragraph."
                })

            # Improved keyword matching with stemming-like approach
            import re as _re
            stop_words = {"the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
                         "have", "has", "had", "do", "does", "did", "will", "would", "could",
                         "should", "may", "might", "shall", "can", "to", "of", "in", "for",
                         "on", "with", "at", "by", "from", "as", "into", "about", "like"}
            covered = []
            for kp in key_points:
                kp_words = [w for w in kp.lower().split() if w not in stop_words]
                word_match = sum(1 for w in kp_words if w in answer_lower)
                ratio = word_match / max(len(kp_words), 1)
                if ratio >= 0.3:
                    covered.append(kp)
            missed = [kp for kp in key_points if kp not in covered]

            kp_score = int(len(covered) / max(len(key_points), 1) * 100) if key_points else 50

            # Structure score: detect STAR, logical flow, paragraph structure
            has_star = any(w in answer_lower for w in ["situation", "task", "action", "result", "star"])
            has_example = any(w in answer_lower for w in ["example", "project", "experience", "role", "position", "company"])
            has_transition = any(w in answer_lower for w in ["first", "second", "finally", "additionally", "moreover", "however", "therefore", "specifically"])
            has_para_structure = n_sentences >= 3 and avg_words_per_sent >= 8
            struct_score = 85 if has_star and has_para_structure else \
                           70 if has_star or (has_example and has_para_structure) else \
                           50 if has_example and has_transition else \
                           30 if n_sentences >= 2 else 15

            # Completeness score: more detailed answers score higher
            has_metrics = bool(_re.search(r'\d+', answer_lower))
            has_specifics = has_metrics or any(w in answer_lower for w in ["because", "resulted", "achieved", "improved", "led to", "outcome"])
            comp_score = min(100, kp_score + 25) if n_words > 50 else \
                         min(100, kp_score + 10) if n_words > 30 else \
                         min(70, kp_score)

            # Clarity score: sentence structure, readability
            clarity_score = 80 if has_para_structure and avg_words_per_sent >= 8 else \
                            60 if n_sentences >= 2 and avg_words_per_sent >= 5 else \
                            40 if n_sentences >= 1 else 20
            # Clarity penalized for very long sentences (run-on)
            if avg_words_per_sent > 35:
                clarity_score = max(30, clarity_score - 20)

            # Relevance score: answer length relative to question, avoids copying question
            question_lower = question_text.lower()
            copied_ratio = sum(1 for w in question_lower.split() if w in answer_lower and len(w) > 3) / max(len(question_lower.split()), 1)
            relevance_score = 85 if copied_ratio < 0.6 and n_words >= 20 else \
                              60 if copied_ratio < 0.8 and n_words >= 10 else \
                              40 if n_words >= 5 else 20

            # Depth score: uses specifics, examples, metrics
            depth_score = 85 if has_star and has_metrics and n_words > 60 else \
                          70 if (has_example or has_metrics) and n_words > 40 else \
                          50 if n_words > 25 else 30

            # Weighted overall
            overall = int(kp_score * 0.35 + struct_score * 0.15 + comp_score * 0.15 +
                         clarity_score * 0.10 + relevance_score * 0.10 + depth_score * 0.15)
            overall = max(10, min(100, overall))

            strengths = []
            improvements = []
            if kp_score >= 50:
                strengths.append("Addressed key points expected for this question")
            else:
                improvements.append(f"Cover these expected points: {', '.join(missed[:4])}")
            if has_star:
                strengths.append("Used STAR or concrete example structure")
            elif has_example:
                strengths.append("Referenced past experience")
            else:
                improvements.append("Use a structured format (STAR: Situation, Task, Action, Result)")
            if has_metrics:
                strengths.append("Used numbers/metrics to support your answer")
            if clarity_score >= 60:
                strengths.append("Clear and well-articulated response")
            else:
                improvements.append("Improve clarity with shorter sentences and better paragraph structure")
            if depth_score >= 50:
                strengths.append("Provided specific details and depth")
            else:
                improvements.append("Add more specific examples and details to strengthen your answer")
            if n_words < 40:
                improvements.append("Expand your answer with more detail")
            if not has_transition and n_words > 50:
                improvements.append("Use transition words (first, second, finally) for better flow")

            brief = f"Score: {overall}/100. "
            if kp_score < 60:
                brief += f"Missed key points: {', '.join(missed[:4])}. " if missed else ""
            if not has_star and n_words > 30:
                brief += "Structure your answer with concrete examples (use 'In my previous role...')."
            elif n_words < 40:
                brief += "Provide more detail in your answer."
            else:
                brief += "Good effort! Refine by connecting more directly to the key points."

            return _json.dumps({
                "score": overall, "meaning_match": kp_score, "keyword_coverage": kp_score,
                "keywords_covered": covered,
                "keywords_missed": missed,
                "structure_score": struct_score, "completeness_score": comp_score,
                "clarity_score": clarity_score, "relevance_score": relevance_score,
                "depth_score": depth_score,
                "confidence_indicators": ["Answer provided"] if n_words > 20 else [],
                "strengths": strengths or ["Attempted to answer"],
                "improvements": improvements or ["Review the ideal answer for guidance"],
                "improved_answer": "Review the key points above and rewrite your answer covering each one with a specific example from your experience.",
                "brief_feedback": brief
            })
        except Exception:
            return _json.dumps({
                "score": 50, "meaning_match": 50, "keyword_coverage": 50,
                "keywords_covered": [], "keywords_missed": [],
                "structure_score": 50, "completeness_score": 50,
                "clarity_score": 50, "relevance_score": 50, "depth_score": 50,
                "confidence_indicators": [],
                "strengths": ["Answer provided"], "improvements": ["Review the ideal answer for improvement"],
                "improved_answer": "",
                "brief_feedback": "Your answer was evaluated. Review the ideal answer and key points above to improve."
            })

    if "feedback" in prompt.lower() or "performance report" in prompt.lower() or "career coach" in prompt.lower():
        # Extract JD and resume from prompt for keyword-aware recommendations
        fb_jd = ""
        fb_resume = ""
        if "JOB DESCRIPTION:" in prompt:
            fb_jd_part = prompt.split("JOB DESCRIPTION:", 1)[1].strip()
            fb_jd = fb_jd_part  # Take everything after the label (JD is at the end of the prompt)
        if "CANDIDATE RESUME:" in prompt:
            fb_resume_part = prompt.split("CANDIDATE RESUME:", 1)[1].strip()
            fb_resume = fb_resume_part.split("JOB DESCRIPTION:")[0].strip() if "JOB DESCRIPTION:" in fb_resume_part else fb_resume_part

        fb_jd_keywords = _extract_jd_keywords(fb_jd) if fb_jd else []
        fb_covered, fb_missing = _jd_resume_keyword_gaps(fb_jd_keywords, fb_resume) if fb_jd_keywords and fb_resume else ([], fb_jd_keywords)

        fb_recs = [
            "Review the questions where you scored lowest and retry with improved answers",
            "Use the STAR method (Situation, Task, Action, Result) to structure your answers",
        ]
        if fb_missing:
            fb_recs.append(f"Your resume is missing these JD keywords: {', '.join(fb_missing[:8])}. Add relevant projects or skills to strengthen your application.")
        if fb_covered:
            fb_recs.append(f"Highlight these resume strengths during interviews: {', '.join(fb_covered[:6])}.")
        fb_recs.append("Quantify achievements with specific metrics and numbers.")

        fb_missing_str = ", ".join(fb_missing[:8]) if fb_missing else "None — great alignment!"
        fb_next = f"Work on the missing JD keywords: {fb_missing_str}." if fb_missing else "Focus on deepening your existing skills with more advanced projects."

        return _json.dumps({
            "overall_summary": f"Your interview performance shows areas for improvement. {'Your resume covers ' + str(len(fb_covered)) + '/' + str(len(fb_jd_keywords)) + ' of the JD keywords.' if fb_jd_keywords else 'Focus on connecting your experience to job requirements.'}",
            "key_strengths": ["Attempted to answer all questions"] + (["Resume aligns with JD keywords"] if fb_covered else []),
            "weak_areas": (["Missing JD keywords in resume"] if fb_missing else []) + ["Answer structure needs improvement"],
            "general_feedback": "Structure your general answers with specific examples. Use the STAR method to provide context.",
            "technical_feedback": "Demonstrate technical depth by explaining not just what you did, but how and why.",
            "resume_feedback": f"Your resume covers {len(fb_covered)}/{len(fb_jd_keywords)} JD keywords. {'Missing: ' + ', '.join(fb_missing[:6]) + '.' if fb_missing else 'Great alignment!'}",
            "recommendations": fb_recs,
            "next_steps": fb_next
        })

    return '{"result": "Demo response"}'


# ─────────────────────────────────────────────────────────────────────────────
# AI call helper (Gemini / OpenAI)
# ─────────────────────────────────────────────────────────────────────────────
_DEMO_MODE = False

def _check_demo_mode():
    global _DEMO_MODE
    if _DEMO_MODE:
        return True
    groq_key = os.getenv("GROQ_API_KEY") or ""
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or ""
    openai_key = os.getenv("OPENAI_API_KEY") or ""
    if not groq_key and not gemini_key and not openai_key:
        _DEMO_MODE = True
        return True
    return False


def _try_ollama(prompt: str, json_mode: bool):
    """Try calling local Ollama instance. Returns response text or None if unavailable."""
    try:
        import httpx as _httpx
        ollama_payload = {
            "model": "llama3.2",
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.7},
        }
        if json_mode:
            ollama_payload["format"] = "json"
        resp = _httpx.post("http://localhost:11434/api/generate", json=ollama_payload, timeout=60)
        if resp.status_code == 200:
            data = resp.json()
            return data.get("response", "")
    except Exception:
        pass
    return None


def _ai_call(prompt: str, json_mode: bool = False) -> str:
    """Route to Groq, Ollama, OpenAI, or Gemini, with demo fallback."""

    groq_key = os.getenv("GROQ_API_KEY") or ""
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or ""
    openai_key = os.getenv("OPENAI_API_KEY") or ""

    # Try Groq first if key is set (free, no credit card)
    if groq_key:
        try:
            import httpx as _httpx
            groq_model = os.getenv("GROQ_MODEL") or "llama-3.3-70b-versatile"
            groq_payload = {
                "model": groq_model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 4096,
            }
            if json_mode:
                groq_payload["response_format"] = {"type": "json_object"}
            groq_resp = _httpx.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {groq_key}", "Content-Type": "application/json"},
                json=groq_payload,
                timeout=60,
            )
            if groq_resp.status_code == 200:
                return groq_resp.json()["choices"][0]["message"]["content"]
        except Exception:
            pass

    # No cloud keys → try local Ollama
    if not groq_key and not gemini_key and not openai_key:
        ollama_resp = _try_ollama(prompt, json_mode)
        if ollama_resp is not None:
            return ollama_resp

    if _check_demo_mode():
        return _demo_ai_response(prompt, json_mode)

    try:
        if st_session.get("ai_model") == "openai":
            resp = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                response_format={"type": "json_object"} if json_mode else {"type": "text"},
            )
            return resp.choices[0].message.content
        else:
            resp = _model.generate_content(
                prompt,
                generation_config={
                    "temperature": 0.7,
                    "response_mime_type": "application/json" if json_mode else "text/plain",
                },
            )
            return resp.text
    except Exception:
        _DEMO_MODE = True
        return _demo_ai_response(prompt, json_mode)


# ─────────────────────────────────────────────────────────────────────────────
# 1. Generate structured interview Q&A
# ─────────────────────────────────────────────────────────────────────────────
def generate_structured_interview_qa(resume_text: str, job_description: str, duration: str, interview_type: str = "behavioral", difficulty: str = "medium") -> dict:
    """
    Returns a dict with categories based on interview_type:
    - "behavioral": {"behavioral": [10], "resume": [5]}  → total 15
    - "technical":  {"technical": [10], "resume": [5]}    → total 15
    """
    counts = QUESTION_COUNTS.get(interview_type, QUESTION_COUNTS["behavioral"])

    if interview_type == "behavioral":
        cat1_name = "behavioral"
        cat1_label = "Behavioral / Situational"
        cat1_count = counts["behavioral"]
        cat1_instructions = f"""
**{cat1_label} ({cat1_count} questions):**
- These must be behavioral/situational questions using the STAR format (Situation, Task, Action, Result).
- Each question must be SPECIFIC to a responsibility, challenge, or context mentioned in the job description.
- Example: If the JD mentions "cross-functional collaboration with product managers", ask: "Tell me about a time you navigated a disagreement with a product manager on technical requirements."
- Focus on soft skills: leadership, teamwork, conflict resolution, adaptability, communication, ownership.
- Questions should start with phrases like: "Tell me about a time...", "Give me an example of...", "Describe a situation where..."
- DO NOT ask any technical/coding questions in this section.
"""
    else:
        cat1_name = "technical"
        cat1_label = "Technical / Skills-Based"
        cat1_count = counts["technical"]
        cat1_instructions = f"""
**{cat1_label} ({cat1_count} questions):**
- These are for a VIRTUAL TEXT-BASED INTERVIEW — no whiteboard, no diagrams. Every question must be answerable with just a keyboard.
- EVERY question must test knowledge of a specific tool, language, framework, or concept EXPLICITLY mentioned in the job description.
- Extract ALL technical keywords from the JD first. Then write one question per keyword or combination.
- Questions must be CONVERSATIONAL and DISCUSSION-BASED — like a real interviewer asking follow-ups.
- PREFER question types like: "Explain how...", "Walk me through...", "What's your experience with...", "How would you approach...", "Can you describe a time you used...".
- AVOID "Design a..." or "Draw the architecture of..." questions — these require diagrams and don't work in a text-only format.
- AVOID coding questions that require writing code (no "write a function to...").
- Examples of GOOD questions: "Walk me through how you'd debug a slow SQL query.", "Explain the trade-offs between REST and GraphQL.", "How would you approach optimizing an API endpoint that's frequently called?"
- These must still probe depth — ask about trade-offs, real experience, and decision-making — but in a conversational way.
- NO behavioral or generic questions. Every question must be IMPOSSIBLE to ask without the specific JD.
"""

    res_count = counts["resume"]

    diff_instructions = {
        "easy": "Questions should test FOUNDATIONAL knowledge. Ask about core concepts, basic definitions, and straightforward applications. Avoid edge cases or deep trade-offs.",
        "medium": "Questions should test SOLID understanding. Ask about practical implementation, common trade-offs, and real-world application. Standard interview depth.",
        "hard": "Questions should test DEEP EXPERTISE. Ask about complex scenarios, nuanced trade-offs, advanced optimizations, troubleshooting edge cases, and architectural decisions.",
    }
    diff_guide = diff_instructions.get(difficulty, diff_instructions["medium"])

    prompt = f"""
You are an expert interviewer and career coach. Your task is to generate a highly targeted interview question bank that is DEEPLY ALIGNED with the specific job description below.

DIFFICULTY LEVEL: {difficulty.upper()}
{diff_guide}

OUTPUT FORMAT (strict JSON only, no markdown, no extra text):
{{
  "{cat1_name}": [
    {{
      "question": "...",
      "difficulty": "{difficulty}",
      "ideal_answer": "A complete well-structured answer (200-250 words)...",
      "key_points": ["key point 1", "key point 2", "key point 3", "key point 4", "key point 5"]
    }}
  ],
  "resume": [...]
}}

INTERVIEW TYPE: {cat1_label}

CRITICAL RULES — Follow these EXACTLY:

{cat1_instructions}

**Resume-based ({res_count} questions):**
- Each question must ask about a SPECIFIC project, role, technology, or achievement mentioned in the resume.
- Probe depth at the {difficulty} level: {'basic understanding and familiarity' if difficulty == 'easy' else 'practical experience and decision-making' if difficulty == 'medium' else 'deep technical depth, trade-offs, and advanced optimizations'}.
- Questions must tie the candidate's past experience to the job requirements.
- For behavioral type: use STAR format for resume questions too.
- For technical type: probe technical depth of their past projects.

ADDITIONAL REQUIREMENTS:
- ALL {cat1_count + res_count} questions must be unique and non-repetitive.
- ideal_answer: 200-250 words, comprehensive, well-structured.
- key_points: 5-7 essential concepts/keywords the answer must cover.
- Do NOT include numbering inside question text.
- If the JD or resume is thin, infer reasonable questions based on what IS provided — never fill with generic questions.

PERSONALIZATION — EVERY ideal_answer MUST be grounded in the candidate's RESUME:
- Write every ideal_answer in FIRST PERSON, in the candidate's voice, as a strong example answer they could give.
- Base it on the candidate's ACTUAL roles, companies, projects, technologies, skills, and achievements listed in the RESUME below.
- Every answer MUST reference at least one specific item from the resume (a role, company, project, tool, metric, or accomplishment).
- Tie the candidate's resume experience directly to the job description: show how their past work maps to what this job asks for.
- Structure behavioral answers using STAR (Situation, Task, Action, Result), filling in specifics from the resume.
- For technical questions, the answer must describe how the candidate has actually used the relevant tool/concept in their resume projects/roles, plus best practice.
- If the resume lacks a specific fact the answer needs (e.g. an exact metric, team size, or technology), write the example with a clearly-marked placeholder in SQUARE BRACKETS, e.g. "[e.g. reduced query latency by 40%]" or "[e.g. a 6-person team]", so the user can fill in their real number.
- NEVER write a generic model answer that ignores the resume. The answer must read like this specific candidate's own strong answer.
- The answer should be a "leading" answer the user can copy, edit, or expand — not a vague suggestion.

JOB DESCRIPTION (must anchor EVERY question):
{job_description}

RESUME (grounding source for the PERSONALIZED example answers above):
{resume_text}
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        data = json.loads(raw_clean)
    except Exception as e:
        raise ValueError(f"Failed to parse AI response as JSON: {e}\n\nRaw: {raw[:500]}")

    # Remap old demo format to match requested interview type
    if interview_type == "behavioral" and "general" in data and "behavioral" not in data:
        data["behavioral"] = data.pop("general")
        data.pop("technical", None)
    elif interview_type == "technical" and "technical" in data:
        data.pop("general", None)
        data.pop("behavioral", None)

    return data


# ─────────────────────────────────────────────────────────────────────────────
# 2. Flatten Q list for session (ordered: behavioral then technical)
# ─────────────────────────────────────────────────────────────────────────────
def flatten_questions(qa_bank: dict) -> list:
    """Returns flat list of question objects with section metadata."""
    flat = []
    for section in qa_bank.keys():
        qs = qa_bank.get(section, [])
        for q in qs:
            flat.append({
                "section": section,
                "difficulty": q.get("difficulty", st.session_state.get("interview_difficulty", "medium")),
                "question": q.get("question", ""),
                "ideal_answer": q.get("ideal_answer", ""),
                "key_points": q.get("key_points", []),
            })
    return flat


# ─────────────────────────────────────────────────────────────────────────────
# 3. AI Answer Evaluation
# ─────────────────────────────────────────────────────────────────────────────
def _demo_evaluate_answer(question: str, user_answer: str,
                          key_points: list) -> dict:
    """Rule-based evaluation for demo mode — no AI call needed."""
    answer_lower = user_answer.lower()
    n_words = len(user_answer.split())

    if not user_answer or n_words < 5:
        return {
            "score": 0, "meaning_match": 0, "keyword_coverage": 0,
            "keywords_covered": [], "keywords_missed": key_points or [],
            "structure_score": 0, "completeness_score": 0,
            "clarity_score": 0, "relevance_score": 0, "depth_score": 0,
            "confidence_indicators": [],
            "strengths": [], "improvements": ["Answer too short."],
            "improved_answer": "Expand your answer covering the key points above.",
            "brief_feedback": "Answer too short. Please write at least a paragraph."
        }

    stop_words = {"the","a","an","is","are","was","were","be","been","being",
                  "have","has","had","do","does","did","will","would","could",
                  "should","may","might","shall","can","to","of","in","for",
                  "on","with","at","by","from","as","into","about","like","its","it's"}
    answer_words = set(answer_lower.split())
    covered = []
    for kp in key_points:
        kp_words = [w for w in kp.lower().split() if w not in stop_words]
        if not kp_words:
            continue
        matches = sum(1 for w in kp_words if w in answer_words)
        if matches / len(kp_words) >= 0.3:
            covered.append(kp)
    missed = [kp for kp in key_points if kp not in covered]

    kp_score = int(len(covered) / max(len(key_points), 1) * 100)
    n_sentences = len(re.split(r'[.!?]+', user_answer.strip())) if user_answer.strip() else 0
    avg_wps = n_words / max(n_sentences, 1)

    has_star = any(w in answer_lower for w in ["situation","task","action","result","star"])
    has_example = any(w in answer_lower for w in ["example","project","experience","role"])
    has_metrics = bool(re.search(r'\d+', answer_lower))
    has_transition = any(w in answer_lower for w in ["first","second","finally","additionally","however","therefore"])
    has_para = n_sentences >= 3 and avg_wps >= 8

    struct_score = 85 if has_star and has_para else 70 if has_star or (has_example and has_para) else 50 if has_example and has_transition else 30 if n_sentences >= 2 else 15
    comp_score = min(100, kp_score + 25) if n_words > 50 else min(100, kp_score + 10) if n_words > 30 else min(70, kp_score)
    clarity_score = 80 if has_para and avg_wps >= 8 else 60 if n_sentences >= 2 and avg_wps >= 5 else 40 if n_sentences >= 1 else 20
    if avg_wps > 35:
        clarity_score = max(30, clarity_score - 20)
    relevance_score = 85 if n_words >= 20 else 60 if n_words >= 10 else 40
    depth_score = 85 if has_star and has_metrics and n_words > 60 else 70 if (has_example or has_metrics) and n_words > 40 else 50 if n_words > 25 else 30

    overall = int(kp_score * 0.35 + struct_score * 0.15 + comp_score * 0.15 +
                 clarity_score * 0.10 + relevance_score * 0.10 + depth_score * 0.15)
    overall = max(10, min(100, overall))

    strengths, improvements = [], []
    if kp_score >= 50:
        strengths.append("Addressed expected key points")
    else:
        improvements.append(f"Cover these: {', '.join(missed[:4])}")
    if has_star:
        strengths.append("Used STAR format")
    elif has_example:
        strengths.append("Referenced past experience")
    else:
        improvements.append("Use STAR format (Situation, Task, Action, Result)")
    if has_metrics:
        strengths.append("Used numbers/metrics")
    if clarity_score >= 60:
        strengths.append("Clear articulation")
    else:
        improvements.append("Improve clarity with shorter sentences")
    if depth_score >= 50:
        strengths.append("Good depth and detail")
    else:
        improvements.append("Add specific examples and details")
    if n_words < 40:
        improvements.append("Expand answer with more detail")

    brief = f"Score: {overall}/100. "
    if kp_score < 60 and missed:
        brief += f"Missed: {', '.join(missed[:3])}. "
    if not has_star and n_words > 30:
        brief += "Use concrete examples (e.g., 'In my previous role...')."
    elif n_words < 40:
        brief += "Provide more detail."
    else:
        brief += "Good effort! Refine by connecting more directly to key points."

    return {
        "score": overall, "meaning_match": kp_score, "keyword_coverage": kp_score,
        "keywords_covered": covered, "keywords_missed": missed,
        "structure_score": struct_score, "completeness_score": comp_score,
        "clarity_score": clarity_score, "relevance_score": relevance_score, "depth_score": depth_score,
        "confidence_indicators": ["Answer provided"] if n_words > 20 else [],
        "strengths": strengths or ["Attempted to answer"],
        "improvements": improvements or ["Review the ideal answer for guidance"],
        "improved_answer": "Review the key points above and rewrite your answer covering each one with a specific example.",
        "brief_feedback": brief,
    }


def evaluate_answer(question: str, ideal_answer: str, key_points: list,
                    user_answer: str, section: str, difficulty: str) -> dict:
    if not user_answer or len(user_answer.strip()) < 10:
        return {
            "score": 0, "meaning_match": 0, "keyword_coverage": 0,
            "keywords_covered": [], "keywords_missed": key_points,
            "structure_score": 0, "completeness_score": 0,
            "clarity_score": 0, "relevance_score": 0, "depth_score": 0,
            "confidence_indicators": [],
            "strengths": [], "improvements": ["No answer was provided."],
            "improved_answer": ideal_answer,
            "brief_feedback": "No answer provided. Please attempt to answer all questions."
        }

    if _check_demo_mode():
        return _demo_evaluate_answer(question, user_answer, key_points)

    prompt = f"""
You are a STRICT interview coach evaluating a candidate's answer. You do NOT inflate scores. You compare what the candidate said against the ideal answer and key points.

QUESTION: {question}
SECTION: {section}
DIFFICULTY: {difficulty}

IDEAL ANSWER (reference for comparison):
{ideal_answer}

KEY POINTS the answer should cover:
{json.dumps(key_points)}

CANDIDATE'S ANSWER:
{user_answer}

First, carefully read the question and the ideal answer. Then compare the candidate's answer to the ideal answer. Score based on how closely the candidate's answer matches the meaning, covers the key points, and demonstrates clarity, relevance, and depth.

Output ONLY valid JSON (no markdown):
{{
  "score": <overall 0-100>,
  "meaning_match": <0-100, how closely the meaning matches the ideal>,
  "keyword_coverage": <0-100, % of key points explicitly addressed>,
  "keywords_covered": ["<key point explicitly covered>", ...],
  "keywords_missed": ["<key point NOT addressed>", ...],
  "structure_score": <0-100, logical flow, STAR or clear structure>,
  "completeness_score": <0-100, how complete and detailed>,
  "clarity_score": <0-100, how clear and well-articulated the answer is>,
  "relevance_score": <0-100, how directly on-topic and non-repetitive>,
  "depth_score": <0-100, depth of explanation and use of specific examples>,
  "confidence_indicators": ["<positive indicator>", ...],
  "strengths": ["<strength 1>", "<strength 2>", ...],
  "improvements": ["<specific improvement>", ...],
  "improved_answer": "<a better version of the candidate's actual answer, 150-200 words>",
  "brief_feedback": "<2-3 sentence coach feedback>"
}}

STRICT SCORING RUBRIC:
- 0-20: Answer is blank, gibberish, completely unrelated, or random text with no meaningful connection to the question
- 21-40: Answer touches the general topic but misses most key points; vague or generic
- 41-60: Answer addresses some key points but is incomplete; partial understanding shown
- 61-80: Answer covers most key points with reasonable depth; minor gaps remain
- 81-90: Answer covers nearly all key points with good detail and structure
- 91-100: Answer matches the ideal answer closely; all key points covered with excellent depth and examples

CRITICAL RULES:
- If the candidate wrote random text, nonsense, or copied the question back, score 0-15.
- Do NOT give partial credit for unrelated content. The answer must actually address the key points.
- A vague 1-sentence answer that mentions one key point should score no more than 25.
- "Improved_answer" must start from the candidate's actual answer (if any substance), not from the ideal answer.
- Be STRICT. A score of 50 means the answer was mediocre.
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        return json.loads(raw_clean)
    except Exception:
        return {
            "score": 50, "meaning_match": 50, "keyword_coverage": 50,
            "keywords_covered": [], "keywords_missed": [],
            "structure_score": 50, "completeness_score": 50,
            "clarity_score": 50, "relevance_score": 50, "depth_score": 50,
            "confidence_indicators": [],
            "strengths": ["Answer provided"], "improvements": ["Could not parse evaluation"],
            "improved_answer": ideal_answer,
            "brief_feedback": "Evaluation could not be parsed. Please try again."
        }


# ─────────────────────────────────────────────────────────────────────────────
# 4a. JD vs Resume Keyword Analyzer
# ─────────────────────────────────────────────────────────────────────────────
def _extract_jd_keywords(job_description: str) -> list:
    """Extract important skill keywords from a job description."""
    tech_terms = [
        "python", "java", "javascript", "typescript", "react", "node", "angular", "vue",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "redis", "docker", "kubernetes",
        "aws", "gcp", "azure", "ci/cd", "git", "linux", "api", "rest", "graphql",
        "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch",
        "flask", "django", "fastapi", "spring", "dotnet", "c++", "c#", "go", "rust",
        "html", "css", "sass", "webpack", "babel", "jest", "pytest", "selenium",
        "agile", "scrum", "jira", "confluence", "microservices", "serverless",
        "arduino", "raspberry pi", "sensor", "actuator", "microcontroller", "robotics",
        "iot", "embedded", "firmware", "circuit", "electronics", "breadboard",
        "scratch", "blockly", "coding", "programming", "logic", "algorithm",
        "data analysis", "pandas", "numpy", "matplotlib", "tableau", "power bi",
        "communication", "leadership", "teamwork", "problem-solving", "analytical",
        "project management", "stakeholder", "mentoring", "training", "teaching",
    ]
    jd_lower = job_description.lower()
    found = []
    for term in tech_terms:
        if term in jd_lower and term not in found:
            found.append(term)
    # Also extract capitalized phrases (potential tools/companies/technologies)
    caps = re.findall(r'\b[A-Z][a-z]*(?:\+|[0-9]+)?(?:\s+[A-Z][a-z]*)*\b', job_description)
    for c in caps:
        c_lower = c.lower()
        if len(c) > 3 and c_lower not in found and c_lower not in ("this", "that", "with", "from", "what", "your", "will", "have", "been", "their", "they", "them", "each", "about", "which", "would", "could", "should", "there", "these", "those", "being", "doing", "having", "making"):
            found.append(c_lower)
    return found


def _jd_resume_keyword_gaps(jd_keywords: list, resume_text: str) -> tuple:
    """Returns (covered_keywords, missing_keywords) comparing JD keywords against resume."""
    resume_lower = resume_text.lower()
    covered = [kw for kw in jd_keywords if kw in resume_lower]
    missing = [kw for kw in jd_keywords if kw not in resume_lower]
    return covered, missing


# ─────────────────────────────────────────────────────────────────────────────
# 4b. Generate Full Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_feedback_report(session_results: list, duration: str, resume_text: str = "", job_description: str = "") -> dict:
    """
    session_results: list of {question_obj, evaluation}
    Returns a structured feedback report dict.
    """
    total_score = 0
    general_scores = []
    technical_scores = []
    resume_scores = []
    all_keywords_covered = set()
    all_keywords_missed = set()
    well_answered = []
    incomplete_answers = []

    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        score = ev.get("score", 0)
        total_score += score

        section = q_obj.get("section", "")
        if section in ("behavioral", "general"):
            general_scores.append(score)
        elif section in ("technical",):
            technical_scores.append(score)
        else:
            resume_scores.append(score)

        all_keywords_covered.update(ev.get("keywords_covered", []))
        all_keywords_missed.update(ev.get("keywords_missed", []))

        q_text = q_obj.get("question", "")
        if score >= 70:
            well_answered.append({"question": q_text, "score": score})
        else:
            incomplete_answers.append({
                "question": q_text,
                "score": score,
                "improvements": ev.get("improvements", []),
                "improved_answer": ev.get("improved_answer", ""),
            })

    n = len(session_results)
    overall = round(total_score / n) if n > 0 else 0
    general_avg = round(sum(general_scores) / len(general_scores)) if general_scores else 0
    technical_avg = round(sum(technical_scores) / len(technical_scores)) if technical_scores else 0
    resume_avg = round(sum(resume_scores) / len(resume_scores)) if resume_scores else 0

    # Remove keywords from missed if they were covered in other questions
    all_keywords_missed -= all_keywords_covered

    # JD vs Resume keyword gap analysis
    jd_keywords = _extract_jd_keywords(job_description)
    resume_covered_kw, resume_missing_kw = _jd_resume_keyword_gaps(jd_keywords, resume_text)

    # Performance band
    if overall >= 85:
        band = "Excellent"
        band_color = "🟢"
    elif overall >= 70:
        band = "Good"
        band_color = "🔵"
    elif overall >= 55:
        band = "Average"
        band_color = "🟡"
    else:
        band = "Needs Improvement"
        band_color = "🔴"

    # Generate AI-powered insights
    results_summary = []
    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        results_summary.append({
            "question": q_obj.get("question", "")[:100],
            "section": q_obj.get("section", ""),
            "difficulty": q_obj.get("difficulty", ""),
            "score": ev.get("score", 0),
            "strengths": ev.get("strengths", []),
            "improvements": ev.get("improvements", []),
        })

    itype = st.session_state.get("interview_type", "behavioral")
    gen_section_label = "BEHAVIORAL / SITUATIONAL" if itype == "behavioral" else "GENERAL"
    gen_feedback_label = "behavioral/situational" if itype == "behavioral" else "general"

    resume_section = f"\nCANDIDATE RESUME:\n{resume_text}" if resume_text else ""
    jd_section = f"\nJOB DESCRIPTION:\n{job_description}" if job_description else ""

    ai_prompt = f"""
You are a senior career coach reviewing a {duration} mock interview session.

OVERALL SCORE: {overall}/100 ({band})
{gen_section_label} AVG: {general_avg}/100
TECHNICAL AVG: {technical_avg}/100
RESUME-BASED AVG: {resume_avg}/100

SESSION RESULTS SUMMARY:
{json.dumps(results_summary, indent=2)}
{resume_section}
{jd_section}

Generate a concise, honest, and constructive performance report. Output ONLY valid JSON:
{{
  "overall_summary": "<2-3 sentence overall assessment>",
  "key_strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "weak_areas": ["<weak area 1>", "<weak area 2>"],
  "general_feedback": "<specific feedback on {gen_feedback_label} interview answers>",
  "technical_feedback": "<specific feedback on technical answers>",
  "resume_feedback": "<specific feedback on resume-based answers>",
  "recommendations": ["<actionable recommendation 1>", "<recommendation 2>", "<recommendation 3>", "<recommendation 4>"],
  "next_steps": "<what to focus on in the next practice session>"
}}

IMPORTANT: Make the 'recommendations' and 'next_steps' SPECIFIC to the candidate's resume and the job description. Reference actual projects, technologies, and skills from the resume. Avoid generic advice.
"""

    try:
        raw = _ai_call(ai_prompt, json_mode=True)
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        ai_insights = json.loads(raw_clean)
    except Exception:
        resume_lines = [l.strip() for l in resume_text.split("\n") if l.strip() and len(l.strip()) > 20] if resume_text else []
        jd_lines = [l.strip() for l in job_description.split("\n") if l.strip() and len(l.strip()) > 20] if job_description else []
        recommendations = [
            f"Review the specific questions where you scored below 70 and practice improving those answers."
        ]
        if resume_lines:
            top_projects = resume_lines[:3]
            for proj in top_projects:
                recommendations.append(f"Prepare a detailed STAR story around: '{proj[:80]}...'")
            recommendations.append(f"Quantify achievements on your resume with specific metrics and numbers.")
        else:
            recommendations.append("Add specific metrics and numbers to each achievement on your resume.")
        recommendations.append("Practice structuring answers using the STAR method (Situation, Task, Action, Result).")

        recs = list(recommendations)
        if resume_missing_kw:
            recs.append(f"Your resume is missing these JD keywords: {', '.join(resume_missing_kw[:8])}. Add relevant projects or skills.")
        if resume_covered_kw:
            recs.append(f"Highlight these strengths from your resume during interviews: {', '.join(resume_covered_kw[:6])}.")
        recs.append("Practice the STAR method to structure your answers with specific situations from your experience.")

        ai_insights = {
            "overall_summary": f"You scored {overall}/100 overall. {'Your resume covers ' + str(len(resume_covered_kw)) + '/' + str(len(jd_keywords)) + ' of the JD keywords.' if jd_keywords else 'Focus on connecting your experience to the job requirements.'}",
            "key_strengths": ["Completed the practice session"] + (["Strong resume alignment with JD"] if resume_covered_kw else ["Good effort"]) + (["Strong resume content"] if resume_lines else []),
            "weak_areas": (["Resume missing key JD keywords: " + ", ".join(resume_missing_kw[:4])] if resume_missing_kw else []) + ["Answers need more specific metrics and structure"],
            "general_feedback": "Work on structuring general answers using the STAR method with examples from your actual experience.",
            "technical_feedback": "Strengthen technical depth with concrete examples from your past projects.",
            "resume_feedback": f"Your resume covers {len(resume_covered_kw)}/{len(jd_keywords)} keywords from the job description. {'Focus on adding: ' + ', '.join(resume_missing_kw[:6]) + '.' if resume_missing_kw else 'Great alignment with the role!'}",
            "recommendations": recs,
            "next_steps": f"Work on the {len(resume_missing_kw)} missing JD keywords: {', '.join(resume_missing_kw[:8])}. Practice answering questions about these areas."
        }

    # Merge all keywords: per-question eval keywords + JD-resume gap keywords
    merged_covered = sorted(set(all_keywords_covered) | set(resume_covered_kw))
    merged_missed = sorted((set(all_keywords_missed) | set(resume_missing_kw)) - set(merged_covered))

    return {
        "overall_score": overall,
        "performance_band": band,
        "band_color": band_color,
        "general_score": general_avg,
        "technical_score": technical_avg,
        "resume_score": resume_avg,
        "total_questions": n,
        "duration": duration,
        "well_answered": well_answered,
        "incomplete_answers": incomplete_answers,
        "keywords_covered": merged_covered,
        "keywords_missed": merged_missed,
        "jd_keywords": jd_keywords,
        "resume_covered_kw": resume_covered_kw,
        "resume_missing_kw": resume_missing_kw,
        "session_results": session_results,
        **ai_insights,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 5. Export report to PDF & DOCX
# ─────────────────────────────────────────────────────────────────────────────
def export_feedback_report(report: dict):
    """Returns (pdf_buffer, docx_buffer)."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from docx import Document
    from docx.shared import Pt, RGBColor

    # ─── PDF ────────────────────────────────────────────────────────────────
    pdf_buf = BytesIO()
    doc = SimpleDocTemplate(pdf_buf, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("TitleS", fontSize=22, leading=28, fontName="Helvetica-Bold",
                              textColor=colors.HexColor("#1a1a2e"), spaceAfter=6)
    h2_s = ParagraphStyle("H2S", fontSize=14, leading=18, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#16213e"), spaceAfter=4, spaceBefore=12)
    h3_s = ParagraphStyle("H3S", fontSize=11, leading=14, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#0f3460"), spaceAfter=3, spaceBefore=8)
    body_s = ParagraphStyle("BodyS", fontSize=10, leading=14, spaceAfter=4)
    bullet_s = ParagraphStyle("BulletS", fontSize=10, leading=14, leftIndent=16, spaceAfter=3)
    score_s = ParagraphStyle("ScoreS", fontSize=28, fontName="Helvetica-Bold",
                             textColor=colors.HexColor("#e94560"), spaceAfter=4)
    sub_s = ParagraphStyle("SubS", fontSize=10, leading=13, textColor=colors.grey, spaceAfter=8)

    story = []

    # Header
    story.append(Paragraph("CVOLVE PRO — Interview Feedback Report", title_s))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}", sub_s))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#e94560")))

    # Score summary table
    story.append(Spacer(1, 10))
    score_data = [
        ["Overall", "General", "Technical", "Resume", "Band"],
        [f"{report['overall_score']}/100", f"{report['general_score']}/100",
         f"{report['technical_score']}/100", f"{report['resume_score']}/100",
         report['performance_band']],
    ]
    score_table = Table(score_data, colWidths=[1.2*inch, 1.1*inch, 1.1*inch, 1.1*inch, 1.1*inch])
    score_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16213e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 11),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f4ff")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccddff")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(score_table)
    story.append(Spacer(1, 14))

    # Overall summary
    story.append(Paragraph("📋 Overall Assessment", h2_s))
    story.append(Paragraph(report.get("overall_summary", ""), body_s))

    # Strengths
    story.append(Paragraph("💪 Key Strengths", h2_s))
    for s in report.get("key_strengths", []):
        story.append(Paragraph(f"• {s}", bullet_s))

    # Weak areas
    story.append(Paragraph("⚠️ Areas to Improve", h2_s))
    for w in report.get("weak_areas", []):
        story.append(Paragraph(f"• {w}", bullet_s))

    # Section feedback
    story.append(Paragraph("📋 General Questions Feedback", h2_s))
    story.append(Paragraph(report.get("general_feedback", ""), body_s))
    story.append(Paragraph("⚙️ Technical Questions Feedback", h2_s))
    story.append(Paragraph(report.get("technical_feedback", ""), body_s))
    story.append(Paragraph("📄 Resume-based Questions Feedback", h2_s))
    story.append(Paragraph(report.get("resume_feedback", ""), body_s))

    # Questions answered well
    if report.get("well_answered"):
        story.append(Paragraph("✅ Questions Answered Well", h2_s))
        for item in report["well_answered"]:
            story.append(Paragraph(f"• {item['question']} — Score: {item['score']}/100", bullet_s))

    # Incomplete answers with improved versions
    if report.get("incomplete_answers"):
        story.append(Paragraph("📌 Questions Needing Improvement", h2_s))
        for item in report["incomplete_answers"]:
            story.append(Paragraph(item["question"], h3_s))
            story.append(Paragraph(f"Score: {item['score']}/100", sub_s))
            for imp in item.get("improvements", []):
                story.append(Paragraph(f"• {imp}", bullet_s))
            if item.get("improved_answer"):
                story.append(Paragraph("Suggested Improved Answer:", h3_s))
                story.append(Paragraph(item["improved_answer"], body_s))

    # Keywords
    story.append(Paragraph("🔑 Keywords Covered", h2_s))
    covered_text = ", ".join(report.get("keywords_covered", [])) or "None recorded"
    story.append(Paragraph(covered_text, body_s))

    story.append(Paragraph("❌ Keywords Missed", h2_s))
    missed_text = ", ".join(report.get("keywords_missed", [])) or "None — great coverage!"
    story.append(Paragraph(missed_text, body_s))

    # Recommendations
    story.append(Paragraph("🎯 Recommendations for Further Practice", h2_s))
    for r in report.get("recommendations", []):
        story.append(Paragraph(f"• {r}", bullet_s))

    story.append(Paragraph("🚀 Next Steps", h2_s))
    story.append(Paragraph(report.get("next_steps", ""), body_s))

    doc.build(story)
    pdf_buf.seek(0)

    # ─── DOCX ───────────────────────────────────────────────────────────────
    docx_buf = BytesIO()
    wd = Document()
    wd.core_properties.title = "CVOLVE PRO Interview Feedback Report"

    def add_h(doc, text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)

    def add_p(doc, text, bold=False, size=11):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    def add_b(doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    add_h(wd, "CVOLVE PRO — Interview Feedback Report", 1, (26, 26, 46))
    add_p(wd, f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}")
    add_h(wd, f"Overall Score: {report['overall_score']}/100 — {report['performance_band']}", 2, (233, 69, 96))
    add_p(wd, f"General: {report['general_score']}/100 | Technical: {report['technical_score']}/100 | Resume: {report['resume_score']}/100 | Questions: {report['total_questions']}")

    add_h(wd, "Overall Assessment", 2)
    add_p(wd, report.get("overall_summary", ""))

    add_h(wd, "Key Strengths", 2)
    for s in report.get("key_strengths", []):
        add_b(wd, s)

    add_h(wd, "Areas to Improve", 2)
    for w in report.get("weak_areas", []):
        add_b(wd, w)

    add_h(wd, "General Questions Feedback", 2)
    add_p(wd, report.get("general_feedback", ""))

    add_h(wd, "Technical Questions Feedback", 2)
    add_p(wd, report.get("technical_feedback", ""))

    add_h(wd, "Resume-based Questions Feedback", 2)
    add_p(wd, report.get("resume_feedback", ""))

    if report.get("well_answered"):
        add_h(wd, "Questions Answered Well", 2)
        for item in report["well_answered"]:
            add_b(wd, f"{item['question']} — {item['score']}/100")

    if report.get("incomplete_answers"):
        add_h(wd, "Questions Needing Improvement", 2)
        for item in report["incomplete_answers"]:
            add_h(wd, item["question"], 3)
            add_p(wd, f"Score: {item['score']}/100")
            for imp in item.get("improvements", []):
                add_b(wd, imp)
            if item.get("improved_answer"):
                add_p(wd, "Suggested Improved Answer:", bold=True)
                add_p(wd, item["improved_answer"])

    add_h(wd, "Keywords Covered", 2)
    add_p(wd, ", ".join(report.get("keywords_covered", [])) or "None")

    add_h(wd, "Keywords Missed", 2)
    add_p(wd, ", ".join(report.get("keywords_missed", [])) or "None — great coverage!")

    add_h(wd, "Recommendations", 2)
    for r in report.get("recommendations", []):
        add_b(wd, r)

    add_h(wd, "Next Steps", 2)
    add_p(wd, report.get("next_steps", ""))

    wd.save(docx_buf)
    docx_buf.seek(0)

    return pdf_buf, docx_buf


# ─────────────────────────────────────────────────────────────────────────────
# 6. Speech-to-text via inline JavaScript (direct textarea fill)
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# 7. Main Streamlit UI — show_interview_practice_page()
# ─────────────────────────────────────────────────────────────────────────────
def show_interview_practice_page(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn):
    """
    Main entry point called from app.py.
    Passes down helper functions to avoid circular imports.
    """
    _init_session()

    # ── Phase router ──────────────────────────────────────────────────────────
    phase = st.session_state.get("interview_phase", "setup")

    if phase == "setup":
        _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn)
    elif phase == "session":
        _phase_session()
    elif phase == "f2f":
        _phase_f2f(export_qa_fn)
    elif phase == "report":
        _phase_report()


def _init_session():
    defaults = {
        "interview_phase": "setup",
        "interview_qa_bank": None,
        "interview_questions_flat": None,
        "interview_current_idx": 0,
        "interview_session_results": [],
        "interview_report": None,
        "interview_duration": "30 minutes",
        "interview_type": "behavioral",
        "interview_difficulty": "medium",
        "interview_jd": "",
        "interview_resume_text": "",
        "voice_transcript_buffer": "",
        "interview_show_feedback": False,
        "interview_last_evaluation": None,
        "interview_last_answer": "",
        # F2F (live voice) interview
        "f2f_session_id": None,
        "f2f_is_free": False,
        "f2f_max_minutes": 0,
        "f2f_questions": [],
        "f2f_idx": 0,
        "f2f_start_ts": None,
        "f2f_blocks_charged": 0,
        "f2f_results": [],
        "f2f_status": "",
        "f2f_notice": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────────────────────────────────────
# Phase 1: Setup
# ─────────────────────────────────────────────────────────────────────────────
def _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn):

    st.markdown("""
    <div class="interview-header">
        <h1 style="display:inline-block; vertical-align:middle; margin:0;">🤖 AI Interview Practice</h1>
        <p>A real-time interview simulation powered by AI — generate questions, practice answers, and receive a detailed performance report.</p>
    </div>
    """, unsafe_allow_html=True)

    # ─── Credit Info Banner ───────────────────────────────────────────────────
    col_c1, col_c2, col_c3 = st.columns(3)
    with col_c1:
        st.markdown("""<div class="interview-credit-card">
            <div class="interview-credit-value">5 Credits</div>
            <div class="interview-credit-sub">15-minute session</div></div>""", unsafe_allow_html=True)
    with col_c2:
        st.markdown("""<div class="interview-credit-card">
            <div class="interview-credit-value">8 Credits</div>
            <div class="interview-credit-sub">30-minute session</div></div>""", unsafe_allow_html=True)
    with col_c3:
        st.markdown("""<div class="interview-credit-card">
            <div class="interview-credit-value">12 Credits</div>
            <div class="interview-credit-sub">45-minute session</div></div>""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ⚙️ Session Configuration")

    col1, col2 = st.columns([1, 1])

    with col1:
        duration = st.selectbox(
            "⏱️ Interview Duration",
            options=["15 minutes", "30 minutes", "45 minutes"],
            index=1,
            help="Longer sessions include more questions across all difficulty levels.",
            key="interview_duration_select"
        )
        st.session_state.interview_duration = duration

        interview_type = st.radio(
            "🎯 Interview Type",
            options=["behavioral", "technical"],
            index=0,
            format_func=lambda x: "🧠 Behavioral / Situational" if x == "behavioral" else "💻 Technical / Skills-Based",
            help="Behavioral: STAR-format situational questions. Technical: Deep-dive into JD-specific skills.",
            key="interview_type_select"
        )
        st.session_state.interview_type = interview_type

        difficulty = st.select_slider(
            "📊 Difficulty Level",
            options=["easy", "medium", "hard"],
            value="medium",
            key="interview_difficulty_select",
            help="Easy: Foundational concepts. Medium: Standard depth. Hard: Advanced, requires deep expertise."
        )
        st.session_state.interview_difficulty = difficulty

        jd = st.text_area(
            "📋 Job Description",
            height=160,
            placeholder="Paste the full job description here...",
            key="interview_jd_input",
            value=st.session_state.interview_jd,
        )
        st.session_state.interview_jd = jd

    with col2:
        uploaded = st.file_uploader(
            "📄 Upload Your Resume (PDF / DOCX)",
            type=["pdf", "docx"],
            key="interview_resume_upload"
        )

        resume_text_available = bool(st.session_state.get("interview_resume_text"))
        if resume_text_available:
            st.success("✅ Resume pre-loaded from your optimized CV (upload optional).")

        st.markdown("#### 📊 What You'll Get")
        type_label = "Behavioral + Resume-based" if interview_type == "behavioral" else "Technical + Resume-based"
        st.markdown(f"""
        <div class="interview-card" style="padding: 1.1rem 1.25rem;">
        <ul style="list-style:none;padding:0;margin:0;line-height:1.75;color:var(--text,#1e293b);">
          <li>✅ {type_label} questions tailored to JD</li>
          <li>✅ <b>15 questions total</b> per session</li>
          <li>✅ AI Interviewer — question by question</li>
          <li>✅ Type or speak your answers</li>
          <li>✅ Per-question timer</li>
          <li>✅ AI evaluation on meaning, keywords, structure</li>
          <li>✅ Full feedback report (PDF + DOCX)</li>
          <li>✅ Suggested improved answers</li>
        </ul>
        </div>
        """, unsafe_allow_html=True)

    # ─── Live F2F Mock Interview ──────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🎙️ Live F2F Mock Interview")
    st.caption(
        f"Pay-as-you-go live voice interview — {pricing.F2F_BLOCK_CREDITS} credits per "
        f"{pricing.F2F_BLOCK_MINUTES}-minute block. Interview Pro: up to "
        f"{pricing.F2F_MAX_MINUTES_INTERVIEW_PRO} min. Free plan: one 3-minute voice interview."
    )

    user_email = st.session_state.user_data["email"]
    f2f_account_type = "business" if st.session_state.get("account_type") == "business" else "individual"
    try:
        gate = can_use_f2f(f2f_account_type, user_email)
    except Exception:
        gate = {"allowed": False, "free_once": False, "reason": "requires_interview_pro"}

    if not gate["allowed"]:
        if gate.get("reason") == "free_used":
            st.warning("You've already used your one-time free voice interview. Upgrade to Interview Pro or add a pack to run more.")
        else:
            st.warning("F2F interviews require the **Interview Pro** plan or an active **credit pack**.")
    elif gate.get("free_once"):
        st.info("🎉 You have one free 3-minute voice interview available.")
    else:
        st.info(f"F2F interview available. You'll be billed {pricing.F2F_BLOCK_CREDITS} credits per {pricing.F2F_BLOCK_MINUTES}-minute block.")

    if st.button("🎙️ Start Live F2F Interview", type="primary", key="start_f2f_btn",
                 use_container_width=True, disabled=not gate["allowed"]):
        if not gate["allowed"]:
            st.error("F2F interviews require Interview Pro or an active credit pack.")
        else:
            with st.spinner("Starting your live F2F interview..."):
                start = start_f2f_session(f2f_account_type, user_email)
                if not start["ok"]:
                    st.error(f"Could not start F2F session: {start.get('reason')}")
                else:
                    resume_text = st.session_state.interview_resume_text
                    jd = st.session_state.interview_jd
                    if not resume_text:
                        resume_text = st.session_state.get("f2f_resume_text", "")
                    try:
                        qa_bank = generate_structured_interview_qa(
                            resume_text or "No resume provided",
                            jd or "General professional role",
                            "15 minutes", "behavioral", "medium",
                        )
                        flat = flatten_questions(qa_bank) or []
                    except Exception:
                        flat = _demo_questions()
                    if not flat:
                        end_f2f_session(start["session_id"])
                        st.error("Could not prepare interview questions. No credits were charged.")
                    else:
                        st.session_state.f2f_session_id = start["session_id"]
                        st.session_state.f2f_is_free = start["is_free"]
                        st.session_state.f2f_max_minutes = start["max_minutes"]
                        st.session_state.f2f_questions = flat
                        st.session_state.f2f_idx = 0
                        st.session_state.f2f_results = []
                        st.session_state.f2f_start_ts = time.time()
                        st.session_state.f2f_blocks_charged = 0
                        st.session_state.f2f_status = "active"
                        st.session_state.f2f_notice = ""
                        st.session_state.interview_phase = "f2f"
                        if not start["is_free"]:
                            block = charge_f2f_block(start["session_id"])
                            if block.get("ok"):
                                st.session_state.f2f_blocks_charged = block.get("blocks_charged", 0)
                            else:
                                end_f2f_session(start["session_id"])
                                st.error(f"Could not charge F2F block: {block.get('reason')}")
                                return
                        st.success("Live interview started. Speak your answers — you'll be billed per 15-minute block.")
                        st.rerun()

    # ─── Generate & Download Q&A Bank ────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 Step 1: Generate & Download Q&A Bank")
    st.markdown("*You can download the full Q&A bank (PDF/DOCX) before or after practicing.*")

    gen_col1, gen_col2 = st.columns([2, 1])

    if jd.strip() and (uploaded or resume_text_available):
        with gen_col1:
            if st.button("📚 Generate Q&A Bank + Start Practice Session",
                         type="primary", key="start_practice_btn",
                         use_container_width=True):
                credits_needed = DURATION_CREDITS[duration]
                if not check_access_fn(required_credits=credits_needed):
                    st.error(f"⚠️ You need {credits_needed} credits for a {duration} session. Please top up.")
                    return

                with st.spinner("🤖 AI is generating your personalized interview questions..."):
                    try:
                        if uploaded:
                            resume_text = extract_resume_fn(uploaded)
                        else:
                            resume_text = st.session_state.get("interview_resume_text", "")
                        st.session_state.interview_resume_text = resume_text

                        qa_bank = generate_structured_interview_qa(resume_text, jd, duration, st.session_state.interview_type, st.session_state.get("interview_difficulty", "medium"))
                        st.session_state.interview_qa_bank = qa_bank

                        flat = flatten_questions(qa_bank)
                        if not flat:
                            st.session_state.interview_qa_bank = None
                            st.error("⚠️ No interview questions could be generated. No credits were charged. Please try again.")
                            return

                        st.session_state.interview_questions_flat = flat
                        st.session_state.interview_current_idx = 0
                        st.session_state.interview_session_results = []
                        st.session_state.interview_report = None

                        # Deduct credits
                        user_email = st.session_state.user_data["email"]
                        if not deduct_credits_fn(user_email, credits_needed, feature="Interview Practice"):
                            st.error("⚠️ Credit deduction failed. Your session was not started and no credits were used. Please try again.")
                            return

                        st.session_state.interview_phase = "session"
                        st.success(f"✅ {len(flat)} questions generated! {credits_needed} credits used. Starting session...")
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ Failed to generate questions: {str(e)}")

    else:
        st.info("👆 Please upload your resume and enter a job description to begin.")

    # Download-only option (if bank already exists)
    if st.session_state.interview_qa_bank:
        st.markdown("---")
        st.markdown("### 📁 Download Previously Generated Q&A Bank")
        bank = st.session_state.interview_qa_bank
        resume_text = st.session_state.interview_resume_text
        jd_text = st.session_state.interview_jd

        # Convert structured bank to flat text for export
        flat_text = _qa_bank_to_text(bank)
        try:
            pdf_buf, docx_buf = export_qa_fn(flat_text)
            dc1, dc2 = st.columns(2)
            with dc1:
                st.download_button("📥 Download Q&A PDF", data=pdf_buf,
                                   file_name="interview_qa_bank.pdf", mime="application/pdf",
                                   key="dl_qa_pdf_setup")
            with dc2:
                st.download_button("📥 Download Q&A DOCX", data=docx_buf,
                                   file_name="interview_qa_bank.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key="dl_qa_docx_setup")
        except Exception as e:
            st.warning(f"Export failed: {e}")


def _transcribe_audio_deepgram(audio_bytes: bytes) -> str:
    """Send audio to Deepgram and return transcript text."""
    import httpx
    key = os.getenv("DEEPGRAM_API_KEY")
    if not key:
        st.error("Deepgram API key not found. Set DEEPGRAM_API_KEY in .env")
        return ""
    try:
        for ct in ["audio/webm", "audio/ogg", "audio/mp4", "audio/wav"]:
            try:
                with httpx.Client(timeout=30) as client:
                    resp = client.post(
                        "https://api.deepgram.com/v1/listen",
                        headers={"Authorization": f"Token {key}", "Content-Type": ct},
                        content=audio_bytes,
                        params={"model": "nova-2", "language": "en", "smart_format": "true"},
                    )
                data = resp.json()
                transcript = data.get("results", {}).get("channels", [{}])[0].get("alternatives", [{}])[0].get("transcript", "")
                if transcript:
                    return transcript.strip()
            except Exception:
                continue
        st.error("Deepgram transcription failed after trying all formats")
        return ""
    except Exception as e:
        st.error(f"Deepgram error: {e}")
        return ""


def _render_voice_recorder(session_key: str):
    """Record voice using native st.audio_input, transcribe via Deepgram, auto-fill text area."""
    audio = st.audio_input("Record your answer (click mic to start/stop)", key=f"audio_{session_key}")
    if audio is not None:
        audio_bytes = audio.getvalue()
        if not audio_bytes or len(audio_bytes) < 100:
            st.warning("Recording too short, please try again")
            return
        audio_hash = hash(audio_bytes)
        cache_key = f"_dg_{session_key}"
        if st.session_state.get(cache_key) != audio_hash:
            st.session_state[cache_key] = audio_hash
            with st.spinner("Transcribing via Deepgram..."):
                transcript = _transcribe_audio_deepgram(audio_bytes)
                if transcript:
                    st.session_state[session_key] = transcript
                    st.rerun()
                else:
                    st.warning("No speech detected. Try speaking more clearly or check your mic.")


def _timer_component_html(seconds: int, question_idx: int = 0) -> str:
    """HTML/JS countdown timer that shows remaining time per question. Resets per question via unique idx."""
    seed = question_idx % 10000
    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            display: flex; align-items: center; justify-content: flex-end;
            min-height: 70px; background: transparent;
        }}
        .timer {{
            padding: 10px 20px; border-radius: 12px;
            background: #ffffff;
            border: 1px solid #e2e8f0;
            color: #ff8c00; font-size: 28px; font-weight: 800;
            text-align: center; letter-spacing: 1px;
            box-shadow: 0 2px 8px rgba(15, 23, 42, 0.08);
            min-width: 140px;
        }}
        .timer.warning {{ color: #f59e0b; border-color: #fde68a; background: #fffbeb; }}
        .timer.danger {{ color: #ef4444; border-color: #fecaca; background: #fef2f2; animation: blink 0.8s infinite; }}
        .label {{ font-size: 11px; color: #64748b; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; margin-top: 4px; text-align: center; }}
        @keyframes blink {{ 0%,100%{{opacity:1}} 50%{{opacity:0.4}} }}
    </style>
</head>
<body>
    <div>
        <div class="timer" id="timer_{seed}">{seconds // 60}:{seconds % 60:02d}</div>
        <div class="label">Q{question_idx+1} — Remaining</div>
    </div>
    <script>
    (function() {{
        var total = {seconds};
        var el = document.getElementById('timer_{seed}');
        function tick() {{
            if (total <= 0) return;
            total--;
            var m = Math.floor(total / 60);
            var s = total % 60;
            el.textContent = m + ':' + (s < 10 ? '0' : '') + s;
            el.className = 'timer' + (total <= 10 ? ' danger' : total <= 30 ? ' warning' : '');
            setTimeout(tick, 1000);
        }}
        tick();
    }})();
    </script>
</body>
</html>
    """


def _qa_bank_to_text(bank: dict) -> str:
    """Convert structured Q&A bank dict to displayable text for existing export_interview_qa."""
    lines = []
    idx = 1
    section_titles = {
        "general": "General Questions",
        "technical": "Technical Questions (JD-based)",
        "behavioral": "Behavioral / Situational Questions",
        "resume": "Resume-based Questions",
    }
    for section in bank.keys():
        section_title = section_titles.get(section, section.title())
        lines.append(f"\n=== {section_title} ===\n")
        qs = bank.get(section, [])
        for q in qs:
            lines.append(f"{idx}. {q.get('question', '')}")
            lines.append(f"Answer: {q.get('ideal_answer', '')}")
            lines.append("")
            idx += 1
    return "\n".join(lines)


def _render_inline_feedback(q_obj: dict):
    eval_data = st.session_state.interview_last_evaluation or {}
    answer = st.session_state.interview_last_answer or ""
    score = eval_data.get("score", 0)
    band = "Excellent" if score >= 85 else "Good" if score >= 70 else "Average" if score >= 55 else "Needs Improvement"
    bc = "#10b981" if score >= 85 else "#3b82f6" if score >= 70 else "#f59e0b" if score >= 55 else "#ef4444"

    # Score bar
    st.markdown(f"""
    <div style="border-left:5px solid {bc};background:var(--surface, #ffffff);border-radius:0 12px 12px 0;border:1px solid var(--border, #e2e8f0);border-left-width:5px;padding:16px 20px;margin-bottom:16px;box-shadow:var(--shadow);">
      <div style="display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:8px">
        <div>
          <span style="font-size:13px;font-weight:600;color:var(--muted, #64748b)">📊 Score</span>
          <span style="background:{bc};color:#fff;padding:3px 12px;border-radius:12px;font-size:12px;font-weight:700;margin-left:8px">{band}</span>
          <div style="color:var(--text, #1e293b);font-size:14px;margin-top:6px;line-height:1.5;font-weight:500">{eval_data.get("brief_feedback", "")}</div>
        </div>
        <div style="font-size:36px;font-weight:800;color:{bc};line-height:1">{score}<span style="font-size:16px;color:var(--muted, #64748b)">/100</span></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Keywords covered / missed as cards
    covered = eval_data.get("keywords_covered", [])
    missed = eval_data.get("keywords_missed", [])
    if covered or missed:
        kw1, kw2 = st.columns(2)
        with kw1:
            tags = "".join(f"<span class='interview-badge interview-badge-success' style='margin:3px 6px 3px 0'>{k}</span>" for k in covered)
            st.markdown(f"""
            <div style="border:1px solid #a7f3d0;border-radius:10px;padding:12px 14px;background:#ecfdf5;margin-bottom:10px">
              <div style="font-size:13px;font-weight:700;color:#047857;margin-bottom:6px">✅ Covered</div>
              <div>{tags if tags else '<span style="color:#64748b;font-size:12px">None</span>'}</div>
            </div>
            """, unsafe_allow_html=True)
        with kw2:
            tags = "".join(f"<span class='interview-badge interview-badge-danger' style='margin:3px 6px 3px 0'>{k}</span>" for k in missed)
            st.markdown(f"""
            <div style="border:1px solid #fecaca;border-radius:10px;padding:12px 14px;background:#fef2f2;margin-bottom:10px">
              <div style="font-size:13px;font-weight:700;color:#b91c1c;margin-bottom:6px">❌ Missed</div>
              <div>{tags if tags else '<span style="color:#64748b;font-size:12px">None — great coverage!</span>'}</div>
            </div>
            """, unsafe_allow_html=True)

    # Strengths / Improvements as cards
    strengths = eval_data.get("strengths", [])
    improvements = eval_data.get("improvements", [])
    if strengths or improvements:
        sc, ic = st.columns(2)
        with sc:
            items = "".join(f"<div style='padding:3px 0;font-size:13px;color:#047857;font-weight:500'>✅ {s}</div>" for s in strengths)
            st.markdown(f"""
            <div style="border:1px solid #a7f3d0;border-radius:10px;padding:12px 14px;background:#ecfdf5;margin-bottom:10px">
              <div style="font-size:13px;font-weight:700;color:#047857;margin-bottom:4px">✅ Strengths</div>
              {items if items else '<div style="color:#64748b;font-size:12px">None highlighted</div>'}
            </div>
            """, unsafe_allow_html=True)
        with ic:
            items = "".join(f"<div style='padding:3px 0;font-size:13px;color:#b45309;font-weight:500'>📈 {imp}</div>" for imp in improvements)
            st.markdown(f"""
            <div style="border:1px solid #fde68a;border-radius:10px;padding:12px 14px;background:#fffbeb;margin-bottom:10px">
              <div style="font-size:13px;font-weight:700;color:#b45309;margin-bottom:4px">📈 Areas to Improve</div>
              {items if items else '<div style="color:#64748b;font-size:12px">None highlighted</div>'}
            </div>
            """, unsafe_allow_html=True)

    with st.expander("📝 Your Answer", expanded=False):
        st.markdown(answer or "*Skipped — no answer provided*")

    improved = eval_data.get("improved_answer", "")
    if improved and len(improved) > 20:
        with st.expander("💡 Suggested Improved Answer", expanded=False):
            st.info(improved)

    indicators = eval_data.get("confidence_indicators", [])
    if indicators:
        with st.expander("🔋 Confidence Signals", expanded=False):
            for ind in indicators:
                st.markdown(f"• {ind}")

    total_q = len(st.session_state.interview_questions_flat or [])
    is_last = st.session_state.interview_current_idx >= total_q - 1
    if st.button("➡️ Continue to Next Question" if not is_last else "📊 View Full Report",
                 type="primary", use_container_width=True):
        st.session_state.interview_current_idx += 1
        st.session_state.interview_show_feedback = False
        st.session_state.voice_transcript_buffer = ""
        if st.session_state.interview_current_idx >= total_q:
            _wrap_up_session()
        else:
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 2: Interview Session
# ─────────────────────────────────────────────────────────────────────────────
def _demo_questions():
    """Lightweight fallback questions for F2F when the AI generator fails."""
    return [
        {"section": "behavioral", "difficulty": "Simple",
         "question": "Tell me about yourself and your professional background.",
         "ideal_answer": "A concise STAR-structured summary covering experience, key achievements, and why you fit the role.",
         "key_points": ["Relevant experience", "Measurable achievements", "Alignment with the role"]},
        {"section": "behavioral", "difficulty": "Hard",
         "question": "Describe a time you faced a major challenge at work and how you handled it.",
         "ideal_answer": "Use the STAR format: Situation, Task, Action, Result, with a concrete measurable outcome.",
         "key_points": ["Clear situation", "Your specific action", "Quantified result"]},
        {"section": "behavioral", "difficulty": "Simple",
         "question": "Why do you want this role, and what makes you a good fit?",
         "ideal_answer": "Connect your skills and experience to the role's key requirements with specific examples.",
         "key_points": ["Company/role research", "Skill-to-job mapping", "Enthusiasm"]},
        {"section": "behavioral", "difficulty": "Very Hard",
         "question": "Tell me about a time you had to lead or influence a team without formal authority.",
         "ideal_answer": "Describe how you built trust, communicated a vision, and delivered results through others.",
         "key_points": ["Influence strategy", "Stakeholder management", "Outcome"]},
        {"section": "behavioral", "difficulty": "Simple",
         "question": "Where do you see yourself in five years?",
         "ideal_answer": "A growth-oriented answer that aligns your ambitions with the company's trajectory.",
         "key_points": ["Career direction", "Growth plan", "Fit with company"]},
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Phase: Live F2F (voice) interview — pay-as-you-go per 15-min block
# ─────────────────────────────────────────────────────────────────────────────
def _phase_f2f(export_qa_fn=None):
    session_id = st.session_state.get("f2f_session_id")
    if not session_id:
        st.error("No active F2F session. Returning to setup.")
        st.session_state.interview_phase = "setup"
        st.rerun()
        return

    questions = st.session_state.get("f2f_questions", [])
    idx = st.session_state.get("f2f_idx", 0)
    status = st.session_state.get("f2f_status", "active")

    # ---- block billing check (wall-clock based) ----
    elapsed_min = 0
    if st.session_state.get("f2f_start_ts"):
        elapsed_min = (time.time() - st.session_state["f2f_start_ts"]) / 60.0
    blocks_charged = st.session_state.get("f2f_blocks_charged", 0)
    blocks_due = int(elapsed_min // pricing.F2F_BLOCK_MINUTES)

    if status == "active" and not st.session_state.get("f2f_is_free") and blocks_due >= blocks_charged:
        res = charge_f2f_block(session_id)
        if res.get("ok"):
            if res.get("reason") == "max_minutes":
                st.session_state.f2f_status = "ended"
                st.session_state.f2f_notice = "⏰ Session reached its time limit. Thanks for practicing!"
                status = "ended"
            else:
                st.session_state.f2f_blocks_charged = res.get("blocks_charged", blocks_charged)
                blocks_charged = st.session_state["f2f_blocks_charged"]
        else:
            st.session_state.f2f_status = "ended"
            st.session_state.f2f_notice = "❌ Credits ran out — session ended. Add credits or a pack and start again."
            status = "ended"
            end_f2f_session(session_id)

    # ---- header ----
    is_free = st.session_state.get("f2f_is_free", False)
    total_cost = (blocks_charged if blocks_charged else 0) * pricing.F2F_BLOCK_CREDITS
    st.markdown("""
    <div class="interview-header" style="margin-bottom:16px;">
        <h2 style="margin:0; font-size:1.8rem;">🎙️ Live F2F Interview</h2>
        <p style="margin:6px 0 0; font-size:0.95rem;">Speak naturally — real-time voice interview simulation.</p>
    </div>
    """, unsafe_allow_html=True)

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.metric("Elapsed", f"{int(elapsed_min)} min")
    with m2:
        st.metric("Blocks Billed", f"{blocks_charged} × {pricing.F2F_BLOCK_MINUTES} min")
    with m3:
        st.metric("Cost So Far", f"{total_cost} credits" if not is_free else "Free")
    with m4:
        cap = st.session_state.get("f2f_max_minutes") or "—"
        st.metric("Session Cap", f"{cap} min")

    if status == "ended":
        st.info(st.session_state.get("f2f_notice", "Session ended."))
        results = st.session_state.get("f2f_results", [])
        if results:
            st.markdown("### 📋 Your F2F Session Summary")
            for i, r in enumerate(results):
                ev = r.get("evaluation", {})
                score = ev.get("score", 0)
                st.markdown(f"**Q{i+1}.** {r.get('question', '')[:90]} — *{score}/100*")

        # Download the question bank used in this session
        f2f_qs = st.session_state.get("f2f_questions", [])
        if f2f_qs and export_qa_fn:
            bank = {}
            for q in f2f_qs:
                bank.setdefault(q.get("section", "general"), []).append(q)
            flat_text = _qa_bank_to_text(bank)
            try:
                pdf_buf, docx_buf = export_qa_fn(flat_text)
                st.markdown("### 📥 Download Question Bank")
                dc1, dc2 = st.columns(2)
                with dc1:
                    st.download_button("📥 Download Q&A PDF", data=pdf_buf,
                                       file_name="f2f_question_bank.pdf", mime="application/pdf",
                                       key="dl_f2f_pdf")
                with dc2:
                    st.download_button("📥 Download Q&A DOCX", data=docx_buf,
                                       file_name="f2f_question_bank.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       key="dl_f2f_docx")
            except Exception as e:
                st.warning(f"Export failed: {e}")

        if st.button("↩️ Return to Setup", key="f2f_back_setup", use_container_width=True):
            _reset_interview_session()
            st.rerun()
        return

    # ---- remaining time in current block ----
    block_remaining = pricing.F2F_BLOCK_MINUTES - (int(elapsed_min) % pricing.F2F_BLOCK_MINUTES)
    st.progress(min(1.0, (int(elapsed_min) % pricing.F2F_BLOCK_MINUTES) / pricing.F2F_BLOCK_MINUTES))
    st.caption(f"⏳ {block_remaining} min remaining in current block. Answer freely; time is billed in 15-min blocks.")

    if idx >= len(questions):
        st.session_state.f2f_status = "ended"
        st.session_state.f2f_notice = "All questions covered — session complete!"
        if not is_free:
            end_f2f_session(session_id)
        st.rerun()
        return

    q_obj = questions[idx]
    section = q_obj["section"].title()
    difficulty = q_obj["difficulty"]
    question_text = q_obj["question"]

    st.markdown(f"""
    <div class="interview-card interview-card-accent" style="padding:18px 20px;margin-bottom:12px;">
        <div style="display:flex;gap:8px;margin-bottom:8px;flex-wrap:wrap;align-items:center;">
            <span class="interview-badge interview-badge-neutral">{section}</span>
            <span class="interview-badge interview-badge-primary">{difficulty}</span>
        </div>
        <p style="font-size:18px;font-weight:600;color:var(--text, #1e293b);margin:6px 0 0;line-height:1.5;">{question_text}</p>
    </div>
    """, unsafe_allow_html=True)

    c_left, c_right = st.columns([1, 1])
    with c_left:
        st.components.v1.html(tts_component_html(question_text), height=64)
    with c_right:
        st.components.v1.html(_timer_component_html(120, idx), height=64)

    voice_key = f"f2f_voice_answer_{idx}"
    typed_key = f"f2f_typed_answer_{idx}"

    tab_type, tab_voice = st.tabs(["⌨️ Type Answer", "🎙️ Speak Answer"])
    with tab_type:
        typed_answer = st.text_area("Type your answer here:", height=200,
                                    key=typed_key, label_visibility="collapsed",
                                    placeholder="Speak or type your answer naturally.")
    with tab_voice:
        _render_voice_recorder(voice_key)
        voice_answer = st.text_area("Your answer (editable):", height=140,
                                    key=voice_key, label_visibility="collapsed",
                                    placeholder="Transcript auto-fills here after speaking.")

    final_answer = typed_answer.strip() if typed_answer.strip() else voice_answer.strip()

    btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])
    with btn_col1:
        if st.button("✅ Submit Answer", type="primary", key=f"f2f_submit_{idx}",
                     disabled=not final_answer, use_container_width=True):
            _f2f_submit_answer(q_obj, final_answer, session_id)
    with btn_col2:
        if st.button("⏭️ Next Question", key=f"f2f_next_{idx}", use_container_width=True):
            st.session_state.f2f_idx = idx + 1
            st.rerun()
    with btn_col3:
        if st.button("🔚 End Session", key=f"f2f_end_{idx}", use_container_width=True):
            st.session_state.f2f_status = "ended"
            st.session_state.f2f_notice = "Session ended by you. Great practice!"
            if not is_free:
                end_f2f_session(session_id)
            st.rerun()

    # ── CV-personalized example answer (copy / edit / expand) ────────────────
    _render_example_answer(q_obj, typed_key, show_draft_button=True, idx=idx)

    if st.session_state.get("f2f_results"):
        with st.expander(f"📋 Answered so far — {len(st.session_state['f2f_results'])}", expanded=False):
            for i, r in enumerate(st.session_state["f2f_results"]):
                ev = r.get("evaluation", {})
                score = ev.get("score", 0)
                st.markdown(f"**Q{i+1}.** {r.get('question', '')[:90]} — *{score}/100*")


def _f2f_submit_answer(q_obj: dict, answer: str, session_id):
    """Evaluate an F2F answer, store it, and advance to the next question."""
    try:
        with st.spinner("🤖 AI is evaluating your answer..."):
            evaluation = evaluate_answer(
                question=q_obj["question"],
                ideal_answer=q_obj["ideal_answer"],
                key_points=q_obj["key_points"],
                user_answer=answer,
                section=q_obj["section"],
                difficulty=q_obj["difficulty"],
            )
    except Exception:
        evaluation = {"score": 0, "feedback": "Evaluation unavailable.", "suggestions": []}

    results = list(st.session_state.get("f2f_results", []))
    results.append({
        "question": q_obj["question"],
        "answer": answer,
        "evaluation": evaluation,
    })
    st.session_state.f2f_results = results
    st.session_state.f2f_idx = st.session_state.get("f2f_idx", 0) + 1
    st.rerun()


def _use_example_as_draft(typed_key, example_answer):
    """Pre-fill the typed-answer box with the CV-personalized example answer.

    Must run as an on_click callback so the state is set before the widget is
    instantiated on the next rerun (Streamlit constraint).
    """
    try:
        st.session_state[typed_key] = example_answer
    except Exception:
        pass


def _render_example_answer(q_obj, typed_key, show_draft_button=True, idx=None):
    """Render the 'Example Answer (from your CV)' panel for a question."""
    example_answer = q_obj.get("ideal_answer", "") or ""
    if not example_answer:
        return
    with st.expander("💡 Example Answer (from your CV)", expanded=False):
        st.markdown(
            "<small><em>A strong, personalised answer based on your resume. "
            "Copy it, edit it, or expand on it.</em></small>",
            unsafe_allow_html=True,
        )
        st.markdown(example_answer)
        if show_draft_button and typed_key:
            label = "📋 Use as my starting draft"
            key = f"use_draft_{idx}" if idx is not None else f"use_draft_{typed_key}"
            if st.button(label, key=key):
                _use_example_as_draft(typed_key, example_answer)


def _phase_session():
    questions = st.session_state.interview_questions_flat
    idx = st.session_state.interview_current_idx
    total = len(questions)

    if not questions:
        st.error("No questions found. Please go back and regenerate.")
        if st.button("↩️ Back to Setup"):
            st.session_state.interview_phase = "setup"
            st.rerun()
        return

    # Progress header
    progress_pct = idx / total
    st.markdown(f"""
    <div class="interview-header" style="padding: 1.25rem 1.5rem; margin-bottom: 1.25rem;">
        <div style="display:flex; justify-content:space-between; align-items:center; flex-wrap:wrap; gap:10px;">
            <div style="text-align:left;">
                <h2 style="color:#1e293b !important; margin:0; font-size:1.6rem; font-weight:700;">🤖 AI Interviewer</h2>
                <p style="color:#1e293b !important; margin:2px 0 0; font-size:0.9rem; opacity:0.9;">Answer each question as if in a real interview</p>
            </div>
            <div style="text-align:right;">
                <div style="color:#1e293b; font-size:1.25rem; font-weight:800;">Question {idx+1}/{total}</div>
                <div style="color:#1e293b; font-size:0.8rem; font-weight:600; opacity:0.85;">Session in progress</div>
            </div>
        </div>
        <div class="interview-progress-wrap">
            <div class="interview-progress-fill" style="width:{progress_pct*100:.1f}%;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if idx >= total:
        # All questions answered
        _wrap_up_session()
        return

    q_obj = questions[idx]
    section = q_obj["section"].title()
    difficulty = q_obj["difficulty"]
    question_text = q_obj["question"]
    typed_key = f"typed_answer_{idx}"

    # Difficulty badge color
    diff_colors = {"Simple": "#10b981", "Easy": "#10b981", "Medium": "#f59e0b", "Hard": "#f59e0b", "Very Hard": "#ef4444"}
    diff_color = diff_colors.get(difficulty, "#ff8c00")

    # Determine per-question time limit (total duration in sec / number of questions)
    total_duration_map = {"15 minutes": 900, "30 minutes": 1800, "45 minutes": 2700}
    total_sec = total_duration_map.get(st.session_state.interview_duration, 1800)
    sec_per_q = max(60, total_sec // max(len(questions), 1))

    # Question card
    st.markdown(f"""
    <div style="border-left:5px solid {diff_color}; background: var(--surface, #ffffff); border: 1px solid var(--border, #e2e8f0); border-left-width: 5px; border-radius: 0 12px 12px 0; padding: 18px 20px; margin-bottom: 14px; box-shadow: var(--shadow);">
        <div style="display:flex; gap:8px; margin-bottom:8px; flex-wrap:wrap; align-items:center;">
            <span class="interview-badge interview-badge-neutral">{section}</span>
            <span class="interview-badge interview-badge-primary">{difficulty}</span>
        </div>
        <p style="font-size:18px; font-weight:600; color:var(--text, #1e293b); margin:0; line-height:1.5;">{question_text}</p>
    </div>
    """, unsafe_allow_html=True)

    # Play button (left edge) + Timer (right edge)
    col_l, col_r = st.columns([1, 1])
    with col_l:
        st.components.v1.html(tts_component_html(question_text), height=64)
    with col_r:
        st.components.v1.html(_timer_component_html(sec_per_q, idx), height=64)

    # ── Answer input OR inline feedback (after submission) ────────────────────
    if st.session_state.interview_show_feedback:
        _render_inline_feedback(q_obj)
    else:
        st.markdown("#### ✍️ Your Answer")

        voice_key = f"voice_answer_{idx}"

        tab_type, tab_voice = st.tabs(["⌨️ Type Answer", "🎙️ Speak Answer"])

        with tab_type:
            typed_answer = st.text_area(
                "Type your answer here:",
                height=200,
                placeholder="Give a structured answer. Use STAR format. Be specific and detailed.",
                key=typed_key,
                label_visibility="collapsed"
            )

        with tab_voice:
            _render_voice_recorder(voice_key)

            voice_answer = st.text_area(
                "Your answer (editable):",
                height=140,
                placeholder="Transcript auto-fills here after speaking. Edit if needed...",
                key=voice_key,
                label_visibility="collapsed",
            )

        # Combine: typed preferred, fallback to voice
        final_answer = typed_answer.strip() if typed_answer.strip() else voice_answer.strip()

        # ── Skip / Submit ─────────────────────────────────────────────────────────
        btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])

        with btn_col1:
            submit_disabled = not final_answer
            if st.button("✅ Submit Answer & Next Question",
                         type="primary", key=f"submit_{idx}",
                         disabled=submit_disabled,
                         use_container_width=True):
                _submit_answer(q_obj, final_answer)

        with btn_col2:
            if st.button("⏭️ Skip Question", key=f"skip_{idx}", use_container_width=True):
                _submit_answer(q_obj, "")  # empty = skipped

        with btn_col3:
            if st.button("🔚 End Session Early", key=f"end_{idx}", use_container_width=True):
                _wrap_up_session(early=True)

    # ── Hints panel (collapsed once answered) ─────────────────────────────────
    with st.expander("💡 Hints — Key Points to Cover", expanded=not st.session_state.interview_show_feedback):
        st.markdown("*Cover these concepts in your answer to score higher:*")
        for kp in q_obj.get("key_points", []):
            st.markdown(f"• {kp}")

    # ── CV-personalized example answer (copy / edit / expand) ────────────────
    _render_example_answer(
        q_obj,
        typed_key,
        show_draft_button=not st.session_state.interview_show_feedback,
        idx=idx,
    )

    # ── Previous results quick view ───────────────────────────────────────────
    if st.session_state.interview_session_results:
        with st.expander(f"📋 Previous Answers — {len(st.session_state.interview_session_results)} answered", expanded=False):
            for i, res in enumerate(st.session_state.interview_session_results):
                ev = res.get("evaluation", {})
                score = ev.get("score", 0)
                color = "#10b981" if score >= 70 else "#f59e0b" if score >= 50 else "#ef4444"
                q_text = res["question_obj"]["question"][:80]
                sub_scores = f"M:{ev.get('meaning_match',0)} S:{ev.get('structure_score',0)} C:{ev.get('clarity_score',0)} D:{ev.get('depth_score',0)}"
                with st.container():
                    cols = st.columns([3, 1, 1])
                    cols[0].markdown(f"<span style='font-size:13px;color:var(--text, #1e293b);font-weight:500'>Q{i+1}. {q_text}...</span>", unsafe_allow_html=True)
                    cols[1].markdown(f"<span style='font-size:11px;color:var(--muted, #64748b)'>{sub_scores}</span>", unsafe_allow_html=True)
                    cols[2].markdown(f"<span style='font-size:13px;font-weight:700;color:{color}'>{score}/100</span>", unsafe_allow_html=True)


def _submit_answer(q_obj: dict, answer: str):
    """Evaluate answer, store result, and show detailed feedback inline."""
    with st.spinner("🤖 AI is evaluating your answer..."):
        evaluation = evaluate_answer(
            question=q_obj["question"],
            ideal_answer=q_obj["ideal_answer"],
            key_points=q_obj["key_points"],
            user_answer=answer,
            section=q_obj["section"],
            difficulty=q_obj["difficulty"],
        )

    st.session_state.interview_session_results.append({
        "question_obj": q_obj,
        "user_answer": answer,
        "evaluation": evaluation,
    })

    st.session_state.voice_transcript_buffer = ""
    st.session_state.interview_show_feedback = True
    st.session_state.interview_last_evaluation = evaluation
    st.session_state.interview_last_answer = answer
    st.rerun()


def _wrap_up_session(early: bool = False):
    """Generate report and move to report phase."""
    results = st.session_state.interview_session_results
    if not results:
        st.warning("No answers recorded. Please answer at least one question.")
        return

    if early:
        st.info(f"Session ended early. {len(results)} questions answered.")

    with st.spinner("📊 Generating your feedback report..."):
        report = generate_feedback_report(
            results,
            st.session_state.interview_duration,
            st.session_state.get("interview_resume_text", ""),
            st.session_state.get("interview_jd", ""),
        )
        st.session_state.interview_report = report
        st.session_state.interview_phase = "report"

    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 3: Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def _phase_report():
    report = st.session_state.interview_report
    if not report:
        st.error("No report available.")
        return

    score = report["overall_score"]
    band = report["performance_band"]
    band_color_map = {"Excellent": "#047857", "Good": "#1d4ed8", "Average": "#b45309", "Needs Improvement": "#b91c1c"}
    band_color = band_color_map.get(band, "#ff8c00")

    # ── Score card ────────────────────────────────────────────────────────────
    st.markdown(f"""
    <div class="interview-score-banner">
        <h1 class="interview-score-num">{score}<span style="font-size:1.6rem; opacity:0.8;">/100</span></h1>
        <div class="interview-score-band" style="color:{band_color};">{band}</div>
        <p style="color:#1e293b; margin:4px 0 0; font-size:1.05rem; font-weight:600; opacity:0.95;">Your Interview Performance Score</p>
    </div>
    """, unsafe_allow_html=True)

    # Section scores (dynamic labels based on interview type)
    itype = st.session_state.get("interview_type", "behavioral")
    gen_label = "🧠 Behavioral" if itype == "behavioral" else "📋 General"
    gen_key = "general_score"
    tech_label = "⚙️ Technical"
    tech_key = "technical_score"

    mc1, mc2, mc3, mc4, mc5 = st.columns(5)
    with mc1:
        st.metric(gen_label, f"{report[gen_key]}/100")
    with mc2:
        st.metric(tech_label, f"{report[tech_key]}/100")
    with mc3:
        st.metric("📄 Resume", f"{report['resume_score']}/100")
    with mc4:
        st.metric("✅ Questions Done", report["total_questions"])
    with mc5:
        st.metric("✅ Answered Well", len(report["well_answered"]))

    st.markdown("---")

    # ── Overall Summary ───────────────────────────────────────────────────────
    st.markdown("### 📋 Overall Assessment")
    st.info(report.get("overall_summary", ""))

    col_s, col_w = st.columns(2)
    with col_s:
        st.markdown("#### 💪 Strengths")
        for s in report.get("key_strengths", []):
            st.markdown(f"✅ {s}")

    with col_w:
        st.markdown("#### ⚠️ Areas to Improve")
        for w in report.get("weak_areas", []):
            st.markdown(f"🔸 {w}")

    # ── Section feedback ──────────────────────────────────────────────────────
    col_g, col_t, col_r = st.columns(3)
    with col_g:
        st.markdown(f"#### {gen_label}")
        st.markdown(report.get("general_feedback", ""))
    with col_t:
        st.markdown(f"#### {tech_label}")
        st.markdown(report.get("technical_feedback", ""))
    with col_r:
        st.markdown("#### 📄 Resume")
        st.markdown(report.get("resume_feedback", ""))

    st.markdown("---")

    # ── Per-question results ──────────────────────────────────────────────────
    st.markdown("### 📊 Question-by-Question Results")

    for i, res in enumerate(report.get("session_results", [])):
        q_obj = res.get("question_obj", {})
        ev = res.get("evaluation", {})
        user_ans = res.get("user_answer", "")
        s = ev.get("score", 0)
        border_color = "#10b981" if s >= 70 else "#f59e0b" if s >= 50 else "#ef4444"

        with st.expander(f"Q{i+1}. {q_obj.get('question','')[:80]}... — {s}/100", expanded=False):
            st.markdown(f"""
            <div style="border-left:5px solid {border_color}; background:var(--surface, #ffffff); border:1px solid var(--border, #e2e8f0); border-left-width:5px; border-radius:0 10px 10px 0; padding:14px 18px; margin-bottom:12px; box-shadow:var(--shadow);">
                <div style="margin-bottom:6px;">
                    <span class="interview-badge interview-badge-neutral">{q_obj.get('section','').title()}</span>
                    <span class="interview-badge interview-badge-primary" style="margin-left:6px;">{q_obj.get('difficulty','')}</span>
                </div>
                <div style="font-size:16px; font-weight:600; color:var(--text, #1e293b);">{q_obj.get('question','')}</div>
            </div>
            """, unsafe_allow_html=True)

            r1c1, r1c2, r1c3, r1c4 = st.columns(4)
            r1c1.metric("Overall", f"{ev.get('score',0)}/100")
            r1c2.metric("Meaning", f"{ev.get('meaning_match',0)}/100")
            r1c3.metric("Keywords", f"{ev.get('keyword_coverage',0)}/100")
            r1c4.metric("Structure", f"{ev.get('structure_score',0)}/100")
            r2c1, r2c2, r2c3 = st.columns(3)
            r2c1.metric("Clarity", f"{ev.get('clarity_score',0)}/100")
            r2c2.metric("Relevance", f"{ev.get('relevance_score',0)}/100")
            r2c3.metric("Depth", f"{ev.get('depth_score',0)}/100")

            if user_ans:
                st.markdown("**Your Answer:**")
                st.info(user_ans)
            else:
                st.warning("*Question skipped*")

            kc1, kc2 = st.columns(2)
            with kc1:
                st.markdown("**✅ Keywords Covered:**")
                covered = ev.get("keywords_covered", [])
                st.markdown(", ".join(covered) if covered else "_None_")
            with kc2:
                st.markdown("**❌ Keywords Missed:**")
                missed = ev.get("keywords_missed", [])
                st.markdown(", ".join(missed) if missed else "_None — great!_")

            st.markdown("**💡 Coach Feedback:**")
            st.markdown(ev.get("brief_feedback", ""))

            if ev.get("strengths"):
                st.markdown("**Strengths:**")
                for s_item in ev["strengths"]:
                    st.markdown(f"✔ {s_item}")

            if ev.get("improvements"):
                st.markdown("**Improvements:**")
                for imp in ev["improvements"]:
                    st.markdown(f"• {imp}")

            if ev.get("improved_answer"):
                st.markdown("**📝 Suggested Improved Answer:**")
                st.success(ev["improved_answer"])

    st.markdown("---")

    # ── JD vs Resume Keyword Gap ──────────────────────────────────────────────
    jd_kw = report.get("jd_keywords", [])
    resume_covered = report.get("resume_covered_kw", [])
    resume_missing = report.get("resume_missing_kw", [])
    all_keywords_covered = report.get("keywords_covered", [])
    all_keywords_missed = report.get("keywords_missed", [])

    st.markdown("### 🔑 JD Keyword Analysis")
    st.markdown("*Keywords extracted from the job description — shows what your resume covers vs misses:*")

    if jd_kw:
        jd_covered_badges = " ".join([f'<span class="interview-badge interview-badge-success" style="margin:3px;display:inline-block">✅ {k}</span>' for k in resume_covered])
        jd_missed_badges = " ".join([f'<span class="interview-badge interview-badge-danger" style="margin:3px;display:inline-block">❌ {k}</span>' for k in resume_missing])

        mc1, mc2 = st.columns(2)
        with mc1:
            st.markdown(f"**✅ Found in Resume ({len(resume_covered)}/{len(jd_kw)})**")
            st.markdown(jd_covered_badges if resume_covered else "_None_", unsafe_allow_html=True)
        with mc2:
            st.markdown(f"**❌ Missing from Resume ({len(resume_missing)}/{len(jd_kw)})**")
            st.markdown(jd_missed_badges if resume_missing else "_None — great alignment!_ 🎉", unsafe_allow_html=True)

        # Per-question keywords (shown as additional info)
        if all_keywords_covered or all_keywords_missed:
            st.markdown("**Per-Question Answer Keywords:**")
            pc1, pc2 = st.columns(2)
            with pc1:
                st.markdown("Covered: " + ", ".join(all_keywords_covered) if all_keywords_covered else "")
            with pc2:
                st.markdown("Missed: " + ", ".join(all_keywords_missed) if all_keywords_missed else "")
    else:
        # Fallback: show per-question keywords
        st.markdown("**Keywords Covered:**")
        st.markdown(", ".join(all_keywords_covered) if all_keywords_covered else "_None_")
        st.markdown("**Keywords Missed:**")
        st.markdown(", ".join(all_keywords_missed) if all_keywords_missed else "_None_")

    st.markdown("---")

    # ── Topics To Prepare ──────────────────────────────────────────────────────
    st.markdown("### 📚 Topics to Prepare for This Interview")
    if resume_missing:
        st.markdown("*Focus on these areas to strengthen your candidacy:*")
        for i, kw in enumerate(resume_missing[:12], 1):
            st.markdown(f"**{i}. {kw.title()}** — *Research common interview questions and prepare a detailed answer with specific examples.*")
    else:
        st.markdown("*Your resume covers all detected JD keywords. Focus on deepening your answers with specific metrics and outcomes.*")

    st.markdown("---")

    # ── Recommendations ───────────────────────────────────────────────────────
    st.markdown("### 🎯 Recommendations For Further Practice")
    recs = report.get("recommendations", [])
    if not recs and resume_missing:
        recs = [f"Add {kw} projects or experience to your resume" for kw in resume_missing[:6]]
        recs.append("Practice STAR-format answers for each missing keyword area")
    for rec in recs:
        st.markdown(f"▶ {rec}")

    st.markdown("### 🚀 Next Steps")
    nxt = report.get("next_steps", "")
    if not nxt and resume_missing:
        nxt = f"Focus on these {len(resume_missing)} missing keywords: {', '.join(resume_missing[:10])}. Build projects or gain experience in these areas, then practice answering related interview questions."
    st.info(nxt or "Review your answers above and retry the questions where you scored lowest.")

    st.markdown("---")

    # ── Downloads ─────────────────────────────────────────────────────────────
    st.markdown("### 📥 Download Your Feedback Report")
    try:
        with st.spinner("Preparing report files..."):
            pdf_buf, docx_buf = export_feedback_report(report)

        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("📄 Download PDF Report", data=pdf_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                               mime="application/pdf", key="dl_report_pdf", use_container_width=True)
        with dl2:
            st.download_button("📝 Download DOCX Report", data=docx_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="dl_report_docx", use_container_width=True)
    except Exception as e:
        st.error(f"Report export failed: {e}")

    st.markdown("---")

    # ── Restart ───────────────────────────────────────────────────────────────
    rc1, rc2 = st.columns(2)
    with rc1:
        if st.button("🔁 Practice Again (New Session)", type="primary", use_container_width=True):
            _reset_interview_session()
            st.rerun()
    with rc2:
        if st.button("⚙️ Change Settings", use_container_width=True):
            st.session_state.interview_phase = "setup"
            st.rerun()


def _reset_interview_session():
    keys = [
        "interview_phase", "interview_qa_bank", "interview_questions_flat",
        "interview_current_idx", "interview_session_results", "interview_report",
        "voice_transcript_buffer", "interview_show_feedback", "interview_last_evaluation",
        "interview_last_answer", "f2f_session_id", "f2f_questions", "f2f_idx", "f2f_results",
    ]
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]

    to_delete = [
        k for k in st.session_state.keys()
        if k.startswith(("voice_answer_", "typed_answer_", "f2f_voice_answer_", "f2f_typed_answer_", "audio_", "_dg_"))
    ]
    for k in to_delete:
        del st.session_state[k]
