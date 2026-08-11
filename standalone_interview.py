"""
Standalone launcher — runs ONLY the Interview Practice module locally.
Bypasses the PostgreSQL database, login, and credit engine (test-account
bypass gives unlimited credits). AI falls back to demo mode automatically
when no GROQ/GEMINI/OPENAI keys are set.

Run:  streamlit run standalone_interview.py
"""
import io

import streamlit as st

# Dummy logged-in user — "tester@cvolvepro.com" trips the credit-engine
# test-account bypass, so credits are always available and nothing is charged.
st.session_state.user_data = {"email": "tester@cvolvepro.com", "name": "Local Tester"}
st.session_state.account_type = "individual"

from interview_module import show_interview_practice_page  # noqa: E402


def check_access_fn(required_credits, feature="Interview"):
    """Test account => always allowed."""
    return True


def deduct_credits_fn(email, credits, feature="Interview"):
    """Test account => always succeeds."""
    return True


def extract_resume_fn(uploaded_file) -> str:
    """Extract text from an uploaded PDF/DOCX resume."""
    if uploaded_file is None:
        return ""
    name = (getattr(uploaded_file, "name", "") or "").lower()
    raw = uploaded_file.getvalue()
    if not raw:
        return ""
    try:
        if name.endswith(".pdf"):
            from PyPDF2 import PdfReader
            return "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(raw)).pages)
        if name.endswith(".docx"):
            from docx import Document
            return "\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs)
    except Exception:
        pass
    return ""


def export_qa_fn(text: str):
    """Return (pdf_buffer, docx_buffer) for the Q&A bank text.

    Uses the production exporter (cv_generator.export_interview_qa) which renders
    section headings, bold questions, and bold STAR labels in both PDF and DOCX.
    """
    try:
        from cv_generator import export_interview_qa
        return export_interview_qa(text)
    except Exception:
        pass
    # Fallback: flat text (only reached if cv_generator can't be imported)
    import io as _io
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate
    pdf_buf = _io.BytesIO()
    doc = SimpleDocTemplate(pdf_buf)
    doc.build([Paragraph(t.replace("\n", "<br/>"), getSampleStyleSheet()["BodyText"]) for t in text.split("\n\n")])
    pdf_buf.seek(0)
    docx_buf = _io.BytesIO()
    from docx import Document
    wd = Document()
    for line in text.splitlines():
        wd.add_paragraph(line)
    wd.save(docx_buf)
    docx_buf.seek(0)
    return pdf_buf, docx_buf


st.set_page_config(
    page_title="Interview Practice — Local",
    page_icon="🎤",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Seed a sample resume so a session can start without an upload.
if not st.session_state.get("interview_resume_text"):
    st.session_state.interview_resume_text = (
        "Software Engineer with 5 years of experience building web applications. "
        "Skilled in Python, FastAPI, React, PostgreSQL, Docker, and AWS. "
        "Led a team project that reduced API latency by 40%. "
        "Built REST and GraphQL APIs, CI/CD pipelines with GitHub Actions. "
        "B.Tech in Computer Science."
    )

st.markdown(
    '<h1 style="text-align:center;">🎤 AI Interview Practice — Local Demo</h1>'
    '<p style="text-align:center;color:#64748b;">'
    'Standalone mode (no database / no API keys). A sample resume is pre-loaded.</p>',
    unsafe_allow_html=True,
)

show_interview_practice_page(
    check_access_fn=check_access_fn,
    deduct_credits_fn=deduct_credits_fn,
    extract_resume_fn=extract_resume_fn,
    export_qa_fn=export_qa_fn,
)
