# api_server.py  (clean, straightforward, reusable FastAPI wrapper)
# Drop-in replacement / cleaned version based on your existing files.
#
# Path note: the local file path for your environment is /opt/cvolvepro/CVOLVE-PRO/api_server.py

import os
import io
import time
import base64
import traceback
import inspect
import logging
import re
from typing import Optional, Any, Dict
from io import BytesIO
import secrets
from werkzeug.security import generate_password_hash, check_password_hash
from database import jobsqa_set_email_otp, jobsqa_verify_email_otp
import resend
from payment import create_jobsqa_checkout_session
from payment import handle_jobsqa_payment
from database import payment_exists, save_payment
import stripe
from fastapi import Request



resend.api_key = os.getenv("RESEND_API_KEY")
FROM_EMAIL = os.getenv("FROM_EMAIL", "JobsQA <no-reply@cvolvepro.com>")

# ------------------ Load Streamlit secrets (if present) ------------------
STREAMLIT_SECRETS_PATH = "/opt/cvolvepro/CVOLVE-PRO/.streamlit/secrets.toml"

def _load_toml(path: str):
    try:
        import toml
        return toml.load(path)
    except Exception:
        try:
            import tomli  # type: ignore[reportMissingImports]
            with open(path, "rb") as f:
                return tomli.load(f)
        except Exception:
            return {}

def load_streamlit_secrets(path: str = STREAMLIT_SECRETS_PATH):
    try:
        if not os.path.exists(path):
            logging.info("Streamlit secrets not found at %s", path)
            return

        secrets = _load_toml(path)
        if not secrets:
            logging.info("No secrets parsed from %s", path)
            return

        mapping = {
            "OPENAI_API_KEY": "OPENAI_API_KEY",
            "GEMINI_API_KEY": "GEMINI_API_KEY",
            "GOOGLE_API_KEY": "GEMINI_API_KEY",
            "JWT_SECRET": "JWT_SECRET",
            "FLASK_SECRET": "JWT_SECRET",
            "STRIPE_SECRET_KEY": "STRIPE_SECRET_KEY",
            "STRIPE_PUBLIC_KEY": "STRIPE_PUBLIC_KEY",
            "SMTP_PASSWORD": "SMTP_PASSWORD",
            "SMTP_USERNAME": "SMTP_USERNAME",
            "SMTP_SERVER": "SMTP_SERVER",
            "SMTP_PORT": "SMTP_PORT",
            "SMTP_FROM": "SMTP_FROM",
            "CVOLVE_API_KEY": "CVOLVE_API_KEY",
            "STRIPE_JOBSQA_WEBHOOK_SECRET": "STRIPE_JOBSQA_WEBHOOK_SECRET",
        }

        for key, envname in mapping.items():
            if key in secrets and not os.environ.get(envname):
                os.environ[envname] = str(secrets[key])

        for k, v in secrets.items():
            env_k = str(k).upper()
            if not os.environ.get(env_k):
                if isinstance(v, bool):
                    os.environ[env_k] = "1" if v else "0"
                else:
                    os.environ[env_k] = str(v)

        logging.info("Loaded Streamlit secrets from %s", path)

    except Exception:
        logging.exception("Failed to load Streamlit secrets from %s", path)

load_streamlit_secrets()

STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_JOBSQA_WEBHOOK_SECRET = os.getenv("STRIPE_JOBSQA_WEBHOOK_SECRET")

if not STRIPE_SECRET_KEY:
    raise RuntimeError("STRIPE_SECRET_KEY missing")

stripe.api_key = STRIPE_SECRET_KEY
# ------------------ End secrets loader ------------------

from fastapi import FastAPI, HTTPException, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import jwt

# docx imports for Word generation (copied formatting logic from app.py)
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document: Any = None
    OxmlElement: Any = None
    qn: Any = None
    Inches: Any = None
    Pt: Any = None
    WD_ALIGN_PARAGRAPH: Any = None

# reuse your modules (must be on PYTHONPATH / same folder)
from auth import authenticate_email
from database import (
    get_user_data,
    get_user_credits,
    update_user_credits,
    save_cv_generation,
    jobsqa_get_user_by_email,
    jobsqa_create_user,
    jobsqa_authenticate,
    jobsqa_get_credits,
    jobsqa_update_credits,
    jobsqa_save_interview,
    get_alignment_answers,
    save_alignment_answers
)
import cv_generator
from templates import apply_template
from utils import enforce_page_limit
from credit_engine import has_enough, spend_credits, wallet_balance

# Config (env-first)
JWT_SECRET = os.getenv("JWT_SECRET", os.getenv("FLASK_SECRET", "please_change_me_in_prod"))
CV_CREDIT_COST = int(os.getenv("CV_CREDIT_COST", "3"))
DEFAULT_TEMPLATE = os.getenv("DEFAULT_TEMPLATE", "professional")
DEFAULT_MAX_PAGES = int(os.getenv("CV_MAX_PAGES", "2"))

app = FastAPI(title="CVolve Pro API (clean wrapper)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://cvolvepro.com", "chrome-extension://fbcioogbhdchilmhlkohffeanpdgakgi", "https://jobsqa.com", "https://www.jobsqa.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Request models ----------
class LoginRequest(BaseModel):
    email: str
    password: str

class GenerateCVRequest(BaseModel):
    job_description: Optional[str] = None
    target_match: Optional[int] = 90
    resume_base64: Optional[str] = None
    resume_filename: Optional[str] = None
    template: Optional[str] = None
    sections: Optional[Any] = None
    quantitative_focus: Optional[Any] = None
    action_verb_intensity: Optional[Any] = None
    keyword_matching: Optional[Any] = None
    language: Optional[str] = "English"
    model: Optional[str] = None
    extras: Optional[dict] = None
    output_format: Optional[str] = None  # accept docx/pdf from extension

class GenerateCLRequest(BaseModel):
    job_description: str
    resume_base64: Optional[str] = None
    resume_filename: Optional[str] = None
    language: Optional[str] = "English"
    extras: Optional[dict] = None

class InterviewQARequest(BaseModel):
    resume_base64: Optional[str] = None
    resume_filename: Optional[str] = None
    job_description: str

class JobsQASignupRequest(BaseModel):
    email: str
    password: str

class JobsQALoginRequest(BaseModel):
    email: str
    password: str

class CVJDGapRequest(BaseModel):
    job_description: str
    resume_base64: Optional[str] = None
    resume_filename: Optional[str] = None
    language: Optional[str] = "English"
    max_gaps: Optional[int] = 5

class SaveAlignmentRequest(BaseModel):
    jd_hash: str
    gaps: Optional[list] = []
    answers: Optional[dict] = {}


# ---------- Helpers ----------
def create_jwt_for_user(email: str) -> str:
    payload = {"sub": email, "iat": int(time.time())}
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    if isinstance(token, bytes):
        token = token.decode("utf-8")
    return token

def create_jwt_for_jobsqa_user(email: str) -> str:
    payload = {
        "sub": email,
        "app": "jobsqa",
        "iat": int(time.time())
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    if isinstance(token, bytes):
        token = token.decode("utf-8")
    return token


def verify_bearer_token(authorization: str = Header(None)) -> str:
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    if not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Invalid Authorization header")
    token = authorization.split(" ", 1)[1].strip()
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        email = payload.get("sub")
        if not email:
            raise HTTPException(status_code=401, detail="Invalid token payload")
        user = get_user_data(email)
        if not user:
            raise HTTPException(status_code=401, detail="Unknown token user")
        return email
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
    

def verify_bearer_token_jobsqa(authorization: str = Header(None)) -> str:
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")

    if not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Invalid Authorization header")

    token = authorization.split(" ", 1)[1].strip()

    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])

        # ✅ Ensure token is issued for JobsQA only
        if payload.get("app") != "jobsqa":
            raise HTTPException(status_code=401, detail="Invalid JobsQA token")

        email = payload.get("sub")
        if not email:
            raise HTTPException(status_code=401, detail="Invalid token")

        user = jobsqa_get_user_by_email(email)
        if not user:
            raise HTTPException(status_code=401, detail="JobsQA user not found")

        return email

    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")



def _to_sections_dict(value: Any) -> Optional[Dict[str, bool]]:
    if value is None:
        return None
    if isinstance(value, dict):
        return {k: bool(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return {str(v): True for v in value}
    if isinstance(value, str):
        return {value: True}
    return None

def _inspect_generate_cv_defaults() -> Dict[str, Any]:
    try:
        sig = inspect.signature(cv_generator.generate_cv)
        defaults = {}
        for name, param in sig.parameters.items():
            if param.default is not inspect._empty:
                defaults[name] = param.default
        return defaults
    except Exception:
        return {}

def _build_generate_args(req: GenerateCVRequest, resume_text: str = "") -> Dict[str, Any]:
    defaults = _inspect_generate_cv_defaults()
    req_dict = req.model_dump(exclude_none=True) if hasattr(req, "model_dump") else req.dict(exclude_none=True)
    args: Dict[str, Any] = {}

    sections_input = req_dict.get("sections")
    if sections_input is None and "sections" in defaults:
        sections_input = defaults.get("sections")
    sections_converted = _to_sections_dict(sections_input)

    model_input = req_dict.get("model") or (req_dict.get("extras") or {}).get("model") or defaults.get("model")

    name_map = {
        "resume_text": resume_text,
        "job_description": req_dict.get("job_description"),
        "target_match": req_dict.get("target_match"),
        "template": req_dict.get("template") or defaults.get("template") or DEFAULT_TEMPLATE,
        "sections": sections_converted,
        "quantitative_focus": req_dict.get("quantitative_focus") if "quantitative_focus" in req_dict else defaults.get("quantitative_focus"),
        "action_verb_intensity": req_dict.get("action_verb_intensity") if "action_verb_intensity" in req_dict else defaults.get("action_verb_intensity"),
        "keyword_matching": req_dict.get("keyword_matching") if "keyword_matching" in req_dict else defaults.get("keyword_matching"),
        "language": req_dict.get("language") or defaults.get("language") or "English",
        "model": model_input,
        "model_choice": model_input,
        "extra_context": (req_dict.get("extras") or {}).get("extra_context"),
        "output_format": req_dict.get("output_format") if "output_format" in req_dict else ( (req_dict.get("extras") or {}).get("output_format") or defaults.get("output_format") )
    }

    extras = req_dict.get("extras") or {}
    sig = inspect.signature(cv_generator.generate_cv)
    for name in sig.parameters.keys():
        if name in name_map and name_map[name] is not None:
            args[name] = name_map[name]
        elif name in req_dict and req_dict.get(name) is not None:
            args[name] = req_dict.get(name)
        elif name in extras:
            args[name] = extras[name]
        elif name in defaults and defaults[name] is not None:
            if name == "sections":
                args[name] = _to_sections_dict(defaults[name])
            else:
                args[name] = defaults[name]

    if "sections" in sig.parameters and "sections" in args and args["sections"] is None and req_dict.get("sections") is not None:
        raise ValueError("Unable to convert provided 'sections' into a dict mapping (expected dict/list/string)")

    args = {k: v for k, v in args.items() if v is not None}
    return args

# ------------ Normalization helpers for DOCX-friendly text --------------
# This normalizer enforces the conventions the DOCX builder expects:
# - canonical section headings (UPPERCASE + ':')
# - company header lines as single lines with ' | ' separators
# - bullets normalized to '• '
# - inline <b> and <strong> converted to **bold**
# - collapse excessive blank lines
HEADING_KEYS = [
    "professional summary", "skills", "key skills", "work experience",
    "experience", "education", "projects", "certifications",
    "contact", "contact info", "summary", "achievements"
]
HEADING_RE = re.compile(r'^\s*(' + r'|'.join(re.escape(k) for k in HEADING_KEYS) + r')\s*[:\-]*\s*$', flags=re.I)

def normalize_text_for_docx(text: str) -> str:
    """
    Robust normalizer:
    - convert <b>/<strong> to **bold**
    - normalize bullets to '• '
    - canonicalize headings to UPPERCASE + ':'
    - collapse excessive blank lines
    - remove duplicate header/contact block (including concatenated duplicates on same line)
    - remove consecutive duplicate lines
    - collapse repeated adjacent words (with and without space)
    - collapse bold+plain duplicates: '**X** X' or 'X **X**' -> '**X**'
    """
    if text is None:
        return ""
    text = str(text)

    # Convert HTML bold tags to **...**
    text = re.sub(r'<\s*(b|strong)\s*>(.*?)<\s*/\s*\1\s*>', r'**\2**', text, flags=re.I | re.S)

    # Normalize line endings and split into lines (rstrip each)
    lines = [ln.rstrip() for ln in re.split(r'\r\n|\r|\n', text)]

    # Trim leading/trailing blank lines
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    # --- Robust top-block dedupe (handles same-line repeats & spaced repeats) ---
    non_empty = [i for i, ln in enumerate(lines) if ln.strip()]
    if len(non_empty) >= 1:
        i1 = non_empty[0]
        line1 = lines[i1].strip()

        # If the first line is long, build a token-sequence to detect double repetition even if spacing differs.
        # tokens e.g. ['AAKASH','YADAV']
        tokens = [t for t in re.split(r'\s+', line1) if t]
        if tokens and len(tokens) <= 6:
            # pattern to match token sequence repeated (with optional whitespace between repeats)
            token_seq_ws = r'\s+'.join([re.escape(t) for t in tokens])
            # two cases: repeated with whitespace, or concatenated without whitespace
            patt_repeat_ws = re.compile(r'(^|\s)(' + token_seq_ws + r')(\s{2,}|' + r'\s+' + r'|\b)(' + token_seq_ws + r')(\b|\s)', flags=re.I)
            if patt_repeat_ws.search(line1):
                # replace by single occurrence of the token sequence (preserve original token joiner ' ')
                single = " ".join(tokens)
                lines[i1] = single
                line1 = lines[i1].strip()
            else:
                # concatenated without whitespace e.g. NameName
                concatenated = ''.join(tokens)
                if concatenated.lower().count(tokens[0].lower()) >= 2 and concatenated.lower().startswith(tokens[0].lower()):
                    # attempt safe collapse by halving if even length of chars (fallback)
                    # prefer safe approach: look for immediate duplication in the raw string
                    raw_no_space = re.sub(r'\s+', '', lines[i1])
                    L = len(tokens)
                    # if the raw_no_space contains two equal halves of token sequence, collapse
                    try:
                        # attempt the half-split approach: find mid where halves equal (case-insensitive)
                        for cut in range(3, len(raw_no_space)//2 + 1):
                            if raw_no_space[:cut].lower() == raw_no_space[cut:cut*2].lower():
                                # map back to approximate first half original substring
                                half_candidate = lines[i1][:len(lines[i1])//2].strip()
                                lines[i1] = half_candidate or " ".join(tokens)
                                line1 = lines[i1].strip()
                                break
                    except Exception:
                        pass

        # If there's a second non-empty line, dedupe identical or repeated contact segments
        if len(non_empty) >= 2:
            i2 = non_empty[1]
            line2 = lines[i2].strip()
            if line1 == line2:
                # exact duplicate: remove second
                lines.pop(i2)
            else:
                # collapse repeated pipe-separated segments in contact line e.g. A | B | C | A | B | C
                parts = [p.strip() for p in re.split(r'\s*\|\s*', line2) if p.strip()]
                if parts and len(parts) > 1:
                    # try smallest repeating unit
                    collapsed_once = False
                    for L in range(1, len(parts)//2 + 1):
                        if len(parts) % L != 0:
                            continue
                        ok = True
                        for i in range(0, len(parts), L):
                            if parts[i:i+L] != parts[0:L]:
                                ok = False
                                break
                        if ok:
                            lines[i2] = " | ".join(parts[0:L])
                            collapsed_once = True
                            break
                    if not collapsed_once:
                        # also attempt collapse if exact halves equal even without divisibility
                        half = len(parts)//2
                        if half > 0 and parts[:half] == parts[half:half*2]:
                            lines[i2] = " | ".join(parts[:half])

    # ---- main normalization loop ----
    out_lines = []
    prev_line = None
    prev_blank = False

    for raw in lines:
        if raw is None:
            continue
        line = raw

        # blank handling
        if not line.strip():
            if not prev_blank:
                out_lines.append("")
            prev_blank = True
            prev_line = ""
            continue
        prev_blank = False

        # bullets normalization
        stripped = line.lstrip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            normalized_line = "• " + stripped[2:].strip()
        else:
            normalized_line = line

        # remove exact consecutive duplicates
        if prev_line is not None and normalized_line.strip() == prev_line.strip():
            continue

        # collapse adjacent duplicate words with spaces 'foo foo' -> 'foo'
        normalized_line = re.sub(r'\b(\w+(?:[-]\w+)*)\b(?:\s+\1\b)+', r'\1', normalized_line, flags=re.I)

        # collapse concatenated duplicates without spaces e.g. OperationsOperations -> Operations
        def _collapse_concat_dup(s):
            out = s
            for _ in range(5):
                t = re.sub(r'([A-Za-z][A-Za-z\-\']{2,})(\1\b)', r'\1', out)
                if t == out:
                    break
                out = t
            # additional pattern: repeated token sequences (two-word repeats) like "A B A B" already handled earlier
            return out
        normalized_line = _collapse_concat_dup(normalized_line)

        # collapse bold+plain duplicates:
        # '**Word** Word' -> '**Word**'
        # 'Word **Word**' -> '**Word**'
        try:
            normalized_line = re.sub(r'\*\*(.+?)\*\*\s+(\1)\b', r'**\1**', normalized_line, flags=re.I)
            normalized_line = re.sub(r'\b(\w[\w\-\']{0,})\b\s+\*\*(\1)\*\*', r'**\1**', normalized_line, flags=re.I)
            normalized_line = re.sub(r'\*\*(.+?)\*\*\s+\*\*(\1)\*\*', r'**\1**', normalized_line, flags=re.I)
            normalized_line = re.sub(r'\*\*(.+?)\*\*\s*[-–—:/]?\s*(\1)\b', r'**\1**', normalized_line, flags=re.I)
        except Exception:
            pass

        # canonical headings to UPPERCASE + ':'
        m = HEADING_RE.match(normalized_line)
        if m:
            normalized_line = m.group(1).strip().upper() + ':'

        out_lines.append(normalized_line)
        prev_line = normalized_line

    normalized = "\n".join(out_lines)

    # collapse runs of 3+ blanks to exactly 2 newlines
    normalized = re.sub(r'\n{3,}', '\n\n', normalized)

    # final pass: collapse remaining concatenated duplicates
    normalized = re.sub(r'([A-Za-z][A-Za-z\-\']{2,})\1\b', r'\1', normalized)

    return normalized.strip()

# ----- Word-generation helpers (copied & improved from app.py so formatting matches) -----
def add_bottom_border(paragraph):
    """
    Create a single bottom border on a paragraph (same as app.py).
    """
    if OxmlElement is None or qn is None:
        return
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')     # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    borders.append(bottom)
    pPr.append(borders)

def create_word_document(content: str) -> BytesIO:
    """
    Create a DOCX BytesIO following the same rules as the streamlit app,
    but with improved header deduplication, left-aligned contact row,
    no inline **bold** conversion (per your request), reduced spacing,
    and preserving bold only for section headers and company lines.
    """
    if Document is None:
        raise Exception("python-docx not installed. Install via `pip install python-docx`")

    def _remove_simple_repeated_prefix(s: str) -> str:
        """
        Remove trivial repeated halves or repeated word-sequences.
        Examples fixed:
          "AAKASH YADAVAAKASH YADAV" -> "AAKASH YADAV"
          "JOHN DOE JOHN DOE" -> "JOHN DOE"
        This is heuristic but effective for most generator-caused duplicates.
        """
        if not s or len(s) < 6:
            return s
        # Try exact half-repeat (concatenated)
        n = len(s)
        if n % 2 == 0:
            half = n // 2
            if s[:half].strip().upper() == s[half:].strip().upper():
                return s[:half].strip()
        # Try repeated word sequence (space separated)
        words = s.split()
        if len(words) >= 2:
            for k in range(1, len(words)//2 + 1):
                first = words[:k]
                second = words[k:2*k]
                if first and second and [w.upper() for w in first] == [w.upper() for w in second]:
                    return " ".join(first)
        return s

    def _dedupe_preserve_order(lines):
        seen = set()
        out = []
        for l in lines:
            norm = re.sub(r'\s+', ' ', l.strip()).lower()
            if not norm:
                continue
            if norm in seen:
                continue
            seen.add(norm)
            out.append(l)
        return out

    # Build document
    current_section = ""
    doc = Document()

    # Set narrow margins (same as streamlit)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # Set base font and spacing
    style: Any = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Parse lines and collect header block + sections (simple pass)
    lines = [ln.rstrip() for ln in content.splitlines() if ln and ln.strip()]
    header_block = []
    body_lines = []
    section_started = False
    current_section_name = None
    section_map = {}  # section_name (with colon) -> list of lines

    for ln in lines:
        stripped = ln.strip()
        # Identify section header if it is fully uppercase and ends with ':'
        if stripped.endswith(':') and stripped == stripped.upper():
            section_started = True
            current_section_name = stripped if stripped.endswith(':') else stripped + ':'
            section_map.setdefault(current_section_name, [])
        else:
            if not section_started:
                header_block.append(stripped)
            else:
                if current_section_name:
                    section_map[current_section_name].append(stripped)
                else:
                    # fallback to header
                    header_block.append(stripped)

    # Post-process header: dedupe + try to normalize duplicated names / lines
    header_block = [h for h in header_block if h and h.strip()]
    header_block = _dedupe_preserve_order(header_block)
    # Normalize simple repeated halves inside each header line
    header_block = [_remove_simple_repeated_prefix(h) for h in header_block]

    # If contact details appear multiple times (same email/phone), merge them into one line.
    # Heuristic: keep first line as name (if uppercase or contains two words), and merge the rest as contact.
    header_name = None
    contact_parts = []
    if header_block:
        # take first line as primary name if it contains letters (prefer uppercase)
        header_name_candidate = header_block[0].strip()
        header_name = header_name_candidate
        # rest lines -> contact_parts
        for extra in header_block[1:]:
            contact_parts.append(extra.strip())

    # If the first header line itself contains contact (e.g., "NAME | phone | email"), try split once
    if header_name and ('|' in header_name):
        parts = [p.strip() for p in header_name.split('|') if p.strip()]
        if parts:
            header_name = parts[0]
            contact_parts = parts[1:] + contact_parts

    # Deduplicate contact_parts (case-insensitive)
    contact_parts = _dedupe_preserve_order(contact_parts)

    # Write header to document
    if header_name:
        p_name = doc.add_paragraph()
        r = p_name.add_run(header_name.strip())
        r.bold = True
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # small spacing after name
        p_name.paragraph_format.space_after = Pt(2)

    if contact_parts:
        contact_line = " | ".join(contact_parts)
        p_contact = doc.add_paragraph()
        p_contact.add_run(contact_line)
        p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_contact.paragraph_format.space_after = Pt(4)
    else:
        # If no separate contact parts but there are additional header lines, add them
        # (skip the first because it's the name already written)
        if len(header_block) > 1:
            for extra in header_block[1:]:
                p = doc.add_paragraph()
                p.add_run(extra)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)

    # Now write sections — keep spacing tighter than before
    for section_name, section_lines in section_map.items():
        if not section_lines:
            continue

        # Section header paragraph
        header_para = doc.add_paragraph()
        run = header_para.add_run(section_name)
        run.bold = True
        header_para.paragraph_format.space_before = Pt(6)
        header_para.paragraph_format.space_after = Pt(2)
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_bottom_border(header_para)

        # Content lines
        for ln in section_lines:
            if not ln.strip():
                continue

            # Small top spacing before company lines in Work Experience
            add_top_space = False
            is_company_line = False
            if section_name.lower().startswith("work experience") and "|" in ln and not ln.strip().startswith("•"):
                is_company_line = True
                add_top_space = True

            if add_top_space:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(1)

            para = doc.add_paragraph()
            # Per your request: remove inline **bold** markers and do NOT create inline bold runs
            clean_text = ln.replace("**", "")

            run = para.add_run(clean_text)

            # Keep company/role lines bold for readability (change if you want none bold)
            if is_company_line:
                run.bold = True

            # Projects non-bulleted first lines — keep bold header-ish
            if section_name.lower().startswith("projects") and not clean_text.startswith("•"):
                # optional: keep bold for project title lines
                run.bold = True

            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Reduced spacing between lines and sections
            para.paragraph_format.space_after = Pt(1)   # tightened from Pt(2)
            para.paragraph_format.line_spacing = 1.0

    # Finalize buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------- End Word helpers ----------

def create_cover_letter_docx(text: str) -> BytesIO:
    """
    Final & Stable Cover Letter DOCX builder.
    Fully fixes:
    - Greeting on separate line
    - Closing block (Sincerely, Name, Phone, Email) grouped correctly
    - Blank-line issues
    - Last-line font mismatch (forces same font)
    """

    if Document is None:
        raise Exception("python-docx not installed.")

    doc = Document()

    # consistent margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Normal style
    base: Any = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    clean = text.strip()

    # ----------------------------------------
    # 1) Split into paragraphs (blank line = new para)
    # ----------------------------------------
    paras = [p.strip() for p in re.split(r"\n{2,}", clean) if p.strip()]

    # ----------------------------------------
    # 2) Detect greeting
    # ----------------------------------------
    greeting_re = re.compile(r"^(hello|dear|hi)\b", flags=re.I)
    closing_re = re.compile(r"^(sincerely|regards|thanks|thank you|yours)\b", flags=re.I)
    email_re = re.compile(r"[\w\.-]+@[\w\.-]+\.\w+", flags=re.I)
    phone_re = re.compile(r"(\+?\d[\d\-\s]{6,}\d)")

    greeting = None
    body = []
    closing_bucket = []

    # GREETING EXTRACTION
    for i, p in enumerate(paras):
        if greeting_re.match(p.lower()):
            # If greeting has trailing text, split it
            if "," in p:
                g, rest = p.split(",", 1)
                greeting = g.strip() + ","
                if rest.strip():
                    body.append(rest.strip())
            else:
                greeting = p.strip()
            # remaining paras go to body
            body += paras[i+1:]
            break
    else:
        body = paras[:]  # no greeting found

    # ----------------------------------------
    # 3) Detect closing block (Sincerely → end)
    # ----------------------------------------
    new_body = []
    closing_started = False

    for p in body:
        low = p.lower()

        # If paragraph contains closing token OR email OR phone → closing begins
        if closing_re.search(low) or email_re.search(p) or phone_re.search(p):
            closing_started = True

        if closing_started:
            closing_bucket.append(p)
        else:
            new_body.append(p)

    body = new_body

    # ----------------------------------------
    # 4) Merge closing bucket into single line for parsing
    # ----------------------------------------
    closing_text = " ".join(closing_bucket)
    closing_text = re.sub(r"\s+", " ", closing_text).strip()

    # ----------------------------------------
    # 5) Extract salutation, name, email, phone from closing text
    # ----------------------------------------
    salutation = None
    name = None
    email = None
    phone = None

    # detect salutation
    m_sal = closing_re.match(closing_text.lower())
    if m_sal:
        salutation = m_sal.group(1).title() + ","
        closing_text = closing_text[len(m_sal.group(1)):].strip(" ,")

    # detect email
    m_email = email_re.search(closing_text)
    if m_email:
        email = m_email.group(0)
        closing_text = closing_text.replace(email, "").strip(" ,")

    # detect phone
    m_phone = phone_re.search(closing_text)
    if m_phone:
        phone = m_phone.group(0)
        closing_text = closing_text.replace(phone, "").strip(" ,")

    # remaining = probable name
    name = closing_text.strip(" ,") if closing_text else None

    final_closing_lines = []
    if salutation:
        final_closing_lines.append(salutation)
    if name:
        final_closing_lines.append(name)
    if email:
        final_closing_lines.append(email)
    if phone:
        final_closing_lines.append(phone)

    # ----------------------------------------
    # 6) Write DOCX
    # ----------------------------------------

    # Greet line
    if greeting:
        p = doc.add_paragraph()
        run = p.add_run(greeting)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(12)

    # Body paras
    for block in body:
        p = doc.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            run = p.add_run(line.strip())
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            if i < len(block.split("\n")) - 1:
                run.add_break()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)

    # Closing lines
    if final_closing_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

        for cl in final_closing_lines:
            ln = doc.add_paragraph()
            run = ln.add_run(cl)
            run.font.name = "Calibri"   # <— fixes font mismatch forever
            run.font.size = Pt(11)
            ln.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ln.paragraph_format.space_after = Pt(1)
            ln.paragraph_format.line_spacing = 1.1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------- Endpoints ----------
@app.get("/health")
async def health():
    return {"ok": True}

@app.get("/api/jobsqa/health")
def jobsqa_health():
    return {"ok": True, "service": "jobsqa"}


@app.post("/api/login")
async def api_login(req: LoginRequest, request: Request):
    try:
        email = (req.email or "").strip().lower()
        password = req.password or ""
        if not email or not password:
            raise HTTPException(status_code=400, detail="Missing email or password")

        try:
            user = authenticate_email(email, password)
        except Exception:
            user = None

        if not user:
            raise HTTPException(status_code=401, detail="Invalid credentials or email not verified")

        token = create_jwt_for_user(user["email"])
        return {"success": True, "token": token, "email": user["email"], "credits": user.get("credits", 0)}
    except HTTPException:
        raise
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Internal server error")

    
@app.get("/api/jobsqa/me")
def jobsqa_me(Authorization: str = Header(...)):
    email = verify_bearer_token_jobsqa(Authorization)

    user = jobsqa_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    credits = jobsqa_get_credits(user["id"])

    return {
        "email": email,
        "credits": credits
    }


@app.post("/api/generate_cv")
async def api_generate_cv(req: GenerateCVRequest, Authorization: Optional[str] = Header(None)):
    # 1) auth
    if not Authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    try:
        user_email = verify_bearer_token(Authorization)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

    # 2) credits check
    try:
        if not has_enough("individual", user_email, amount=CV_CREDIT_COST, feature="CV"):
            raise HTTPException(status_code=402, detail="Insufficient credits")
    except HTTPException:
        raise
    except Exception:
        credits_now = None
        try:
            credits_now = get_user_credits(user_email)
        except Exception:
            credits_now = None
        if credits_now is not None and credits_now < CV_CREDIT_COST:
            raise HTTPException(status_code=402, detail="Insufficient credits")

    # 3) validate input
    if not req.job_description or not req.job_description.strip():
        raise HTTPException(status_code=400, detail="Missing job_description")
    if not req.resume_base64:
        raise HTTPException(status_code=400, detail="Missing resume_base64: CV generator requires an uploaded resume (base64)")

    # 4) optional resume extraction
    resume_text = ""
    if req.resume_base64:
        try:
            decoded = base64.b64decode(req.resume_base64)
            bio = io.BytesIO(decoded)
            bio.name = req.resume_filename or "resume.pdf"
            if hasattr(cv_generator, "extract_resume_text"):
                resume_text = cv_generator.extract_resume_text(bio)
        except Exception:
            resume_text = ""

    # 5) prepare args and call generate_cv
    try:
        if req.sections is None:
            req.sections = {
                "Professional Summary": True,
                "Skills": True,
                "Experience": True,
                "Education": True
            }

        if req.quantitative_focus is None:
            req.quantitative_focus = "Medium"

        if req.action_verb_intensity is None:
            req.action_verb_intensity = "Moderate"

        if req.keyword_matching is None:
            req.keyword_matching = "Balanced"

        call_kwargs = _build_generate_args(req, resume_text=resume_text)
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to prepare generation arguments")

    try:
        defaults_server = {
            "template": call_kwargs.get("template") or DEFAULT_TEMPLATE,
            "sections": call_kwargs.get("sections") or {
                "Professional Summary": True,
                "Key Skills": True,
                "Work Experience": True,
                "Education": True
            },
            "quantitative_focus": call_kwargs.get("quantitative_focus") or "Medium",
            "action_verb_intensity": call_kwargs.get("action_verb_intensity") or "Moderate",
            "keyword_matching": call_kwargs.get("keyword_matching") or "Balanced"
        }
        for k, v in defaults_server.items():
            if k not in call_kwargs or call_kwargs.get(k) is None:
                call_kwargs[k] = v

        if 'model' not in call_kwargs and 'model_choice' not in call_kwargs:
            model_from_req = None
            try:
                req_dict = req.dict(exclude_none=True)
                model_from_req = req_dict.get("model") or (req_dict.get("extras") or {}).get("model")
            except Exception:
                model_from_req = None
            call_kwargs['model'] = (model_from_req or "premium").lower()
            call_kwargs['model_choice'] = call_kwargs['model']
        elif 'model' in call_kwargs and 'model_choice' not in call_kwargs:
            call_kwargs['model_choice'] = call_kwargs['model']

        if 'language' not in call_kwargs:
            try:
                call_kwargs['language'] = (req.language or "English")
            except Exception:
                call_kwargs['language'] = "English"

        if 'sections' in call_kwargs:
            if call_kwargs['sections'] is None:
                call_kwargs['sections'] = {
                    "Professional Summary": True,
                    "Key Skills": True,
                    "Work Experience": True,
                    "Education": True
                }
            elif isinstance(call_kwargs['sections'], (list, tuple)):
                call_kwargs['sections'] = {str(s): True for s in call_kwargs['sections']}
            elif isinstance(call_kwargs['sections'], str):
                call_kwargs['sections'] = {call_kwargs['sections']: True}
            elif isinstance(call_kwargs['sections'], dict):
                call_kwargs['sections'] = {str(k): bool(v) for k, v in call_kwargs['sections'].items()}
            else:
                call_kwargs['sections'] = {
                    "Professional Summary": True,
                    "Key Skills": True,
                    "Work Experience": True,
                    "Education": True
                }

        if not call_kwargs.get("extra_context") and req.job_description:
            try:
                jd_h = cv_generator.hash_jd(req.job_description)
                call_kwargs["extra_context"] = get_alignment_answers(user_email, jd_h).get("answers", {})
            except Exception:
                pass

        call_kwargs["optimization_depth"] = "max_ats"
        call_kwargs["return_metadata"] = True

        cv_result = cv_generator.generate_cv(**call_kwargs)
    except TypeError as te:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"CV generation failed: {str(te)}")
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="CV generation failed")

    # ---------------- New normalization block: support docx when requested (uses same rules as Streamlit) ----------------
    try:
        # determine desired output from request/extras/call_kwargs
        desired_output = None
        try:
            req_dict = req.dict(exclude_none=True)
            desired_output = req_dict.get("output_format") or (req_dict.get("extras") or {}).get("output_format")
        except Exception:
            desired_output = None
        desired_output = (call_kwargs.get("output_format") or desired_output or "pdf").lower()

        file_bytes = None
        file_mime = None
        file_ext = None
        text_to_use = None

        # If generator returned a dict with explicit docx/pdf keys, honour them first
        cv_metadata = cv_result if isinstance(cv_result, dict) else {}

        if isinstance(cv_result, dict):
            if "docx_bytes" in cv_result:
                raw = cv_result["docx_bytes"]
                file_bytes = raw.getvalue() if hasattr(raw, "getvalue") else raw
                file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_ext = "docx"
            elif "docx" in cv_result:
                raw = cv_result["docx"]
                file_bytes = raw.getvalue() if hasattr(raw, "getvalue") else raw
                file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_ext = "docx"
            elif "pdf_bytes" in cv_result:
                raw = cv_result["pdf_bytes"]
                file_bytes = raw.getvalue() if hasattr(raw, "getvalue") else raw
                file_mime = "application/pdf"
                file_ext = "pdf"
            elif "pdf" in cv_result:
                file_bytes = cv_result["pdf"]
                file_mime = "application/pdf"
                file_ext = "pdf"
            elif "text" in cv_result:
                text_to_use = cv_result["text"]
            elif "cv_text" in cv_result:
                text_to_use = cv_result["cv_text"]
            elif "optimized_content" in cv_result:
                text_to_use = cv_result["optimized_content"]
            else:
                # fallback to string representation (may contain text)
                text_to_use = str(cv_result)

        elif isinstance(cv_result, str):
            text_to_use = cv_result
        elif isinstance(cv_result, bytes):
            # raw bytes — assume PDF (generator returned PDF)
            file_bytes = cv_result
            file_mime = "application/pdf"
            file_ext = "pdf"
        elif hasattr(cv_result, "getvalue"):
            # buffer-like (likely pdf bytes)
            try:
                file_bytes = cv_result.getvalue()
                file_mime = "application/pdf"
                file_ext = "pdf"
            except Exception:
                text_to_use = str(cv_result)
        else:
            text_to_use = str(cv_result)

        # ---------------- FORCED DOCX PATH WHEN REQUESTED ----------------
        # If the client explicitly requested a DOCX, ensure we produce a proper DOCX:
        if desired_output == 'docx':
            # If we already have docx bytes, use them
            if file_ext == "docx" and file_bytes:
                pass
            else:
                # Try to find usable text to build the DOCX
                if not text_to_use or not str(text_to_use).strip():
                    # look for common text keys in dict result
                    if isinstance(cv_result, dict):
                        text_to_use = cv_result.get("optimized_content") or cv_result.get("text") or cv_result.get("cv_text") or cv_result.get("plain_text") or cv_result.get("result") or None

                    # If still nothing and we have pdf bytes, try to extract text (if helper exists)
                    if (not text_to_use or not str(text_to_use).strip()) and file_bytes and file_ext == "pdf":
                        try:
                            if hasattr(cv_generator, "extract_resume_text"):
                                bio = io.BytesIO(file_bytes)
                                bio.name = "generated.pdf"
                                extracted = cv_generator.extract_resume_text(bio)
                                if extracted and isinstance(extracted, str) and extracted.strip():
                                    text_to_use = extracted
                        except Exception:
                            text_to_use = None

                    # If still nothing, attempt generic string conversion
                    if not text_to_use or not str(text_to_use).strip():
                        try:
                            text_to_use = str(cv_result)
                        except Exception:
                            text_to_use = ""

                # Normalize the text to canonical structure for DOCX
                try:
                    normalized_text = normalize_text_for_docx(text_to_use or "")
                    word_buf = create_word_document(normalized_text)
                    file_bytes = word_buf.getvalue()
                    file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    file_ext = "docx"
                    # ensure we return the normalized_text for logging/debug if needed
                except Exception:
                    traceback.print_exc()
                    # if create_word_document fails, fallback to creating a PDF (best-effort)
                    try:
                        txt = enforce_page_limit((text_to_use or ""), DEFAULT_MAX_PAGES)
                        pdf_buffer = apply_template(txt, call_kwargs.get("template", DEFAULT_TEMPLATE))
                        file_bytes = pdf_buffer.getvalue()
                        file_mime = "application/pdf"
                        file_ext = "pdf"
                    except Exception:
                        traceback.print_exc()
                        raise

        # ---------------- END FORCED DOCX PATH ----------------

        # Safety: If we have text and the client requested DOCX and file isn't docx yet, try again
        if text_to_use is not None and desired_output == 'docx' and (file_ext != "docx"):
            try:
                normalized_text = normalize_text_for_docx(text_to_use or "")
                word_buf = create_word_document(normalized_text)
                file_bytes = word_buf.getvalue()
                file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_ext = "docx"
            except Exception:
                traceback.print_exc()

        # If still no file_bytes, and text exists and desired_output==pdf, create PDF as before
        if not file_bytes and text_to_use is not None:
            txt = enforce_page_limit(text_to_use, DEFAULT_MAX_PAGES)
            pdf_buffer = apply_template(txt, call_kwargs.get("template", DEFAULT_TEMPLATE))
            file_bytes = pdf_buffer.getvalue()
            file_mime = "application/pdf"
            file_ext = "pdf"

        if not file_bytes:
            raise Exception("No usable output from generator")

    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to prepare generated file")

    # 7) encode and save and deduct credits (best-effort)
    try:
        file_b64 = base64.b64encode(file_bytes).decode("utf-8")
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to encode generated file")

    try:
        spend_credits("individual", user_email, "CV", amount=CV_CREDIT_COST)
    except Exception:
        traceback.print_exc()

    try:
        save_cv_generation(
            user_email=user_email,
            job_description=(req.job_description or "")[:20000],
            original_resume=(req.resume_base64 or "")[:200000],
            generated_cv=(file_b64 or "")[:500000],
            template_used=call_kwargs.get("template", DEFAULT_TEMPLATE),
            ats_score=int(cv_metadata.get("ats_score") or 0),
            target_match=(req.target_match or 0),
            processing_time=0.0
        )
    except Exception:
        traceback.print_exc()

    updated_credits = None
    try:
        updated_credits = wallet_balance("individual", user_email)["total"]
    except Exception:
        try:
            updated_credits = get_user_credits(user_email)
        except Exception:
            updated_credits = None

    meta_payload = {
        "measured_ats_score": cv_metadata.get("ats_score"),
        "target_ats_score": cv_metadata.get("target_ats_score", req.target_match or 100),
        "keyword_match": cv_metadata.get("keyword_match"),
        "missing_keywords": cv_metadata.get("missing_keywords", []),
        "optimization_summary": {
            "repair_passes_used": cv_metadata.get("repair_passes_used", 0),
            "fixes_applied": cv_metadata.get("fixes_applied", []),
            "unsupported_gaps": cv_metadata.get("unsupported_gaps", []),
        },
    }
    if file_ext == "docx":
        return {"success": True, "docx_base64": file_b64, "file_mime": file_mime, "file_ext": file_ext, "credits": updated_credits, **meta_payload}
    else:
        return {"success": True, "pdf_base64": file_b64, "file_mime": file_mime, "file_ext": file_ext, "credits": updated_credits, **meta_payload}

# Optional debug endpoint to preview normalized text (useful during testing)
@app.post("/api/debug/preview_cv")
async def debug_preview(req: GenerateCVRequest, Authorization: Optional[str] = Header(None)):
    if not Authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    try:
        user_email = verify_bearer_token(Authorization)
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
    # Build args but do not deduct credits or save
    try:
        call_kwargs = _build_generate_args(req, resume_text="")
        cv_result = cv_generator.generate_cv(**call_kwargs)
        if isinstance(cv_result, dict):
            text = cv_result.get("text") or cv_result.get("cv_text") or str(cv_result)
        elif isinstance(cv_result, str):
            text = cv_result
        elif hasattr(cv_result, "getvalue"):
            text = str(cv_result.getvalue())
        else:
            text = str(cv_result)
        normalized = normalize_text_for_docx(text)
        return {"ok": True, "normalized_preview": normalized}
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Debug preview failed")

# ---------- CV↔JD gap analysis + alignment endpoints ----------

@app.post("/api/cv_jd_gaps")
async def api_cv_jd_gaps(req: CVJDGapRequest, Authorization: Optional[str] = Header(None)):
    """Gap analysis: compare resume against JD, return structured gaps.
    Costs 1 credit (1 LLM call). Returns JSON with gaps and follow-up questions.
    """
    # 1) Auth
    if not Authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    try:
        user_email = verify_bearer_token(Authorization)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

    # 2) Credits check — gap analysis costs 1 credit
    try:
        if not has_enough("individual", user_email, amount=1, feature="CV"):
            raise HTTPException(status_code=402, detail="Insufficient credits (need ≥1 for gap analysis)")
    except HTTPException:
        raise
    except Exception:
        credits_now = None
        try:
            credits_now = get_user_credits(user_email)
        except Exception:
            credits_now = None
        if credits_now is not None and credits_now < 1:
            raise HTTPException(status_code=402, detail="Insufficient credits (need ≥1 for gap analysis)")

    # 3) Validate JD
    if not req.job_description or not req.job_description.strip():
        raise HTTPException(status_code=400, detail="Missing job_description")

    # 4) Extract resume text from base64
    resume_text = ""
    if req.resume_base64:
        try:
            decoded = base64.b64decode(req.resume_base64)
            bio = io.BytesIO(decoded)
            bio.name = req.resume_filename or "resume.pdf"
            if hasattr(cv_generator, "extract_resume_text"):
                resume_text = cv_generator.extract_resume_text(bio)
        except Exception:
            resume_text = ""

    # 5) Run gap analysis
    try:
        from cv_generator import analyze_cv_jd_gaps, hash_jd
        gap_result = analyze_cv_jd_gaps(
            resume_text,
            req.job_description,
            language=req.language or "English",
            max_gaps=req.max_gaps or 5,
        )
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Gap analysis failed")

    # 6) Deduct 1 credit
    try:
        spend_credits("individual", user_email, "CV", amount=1)
    except Exception:
        traceback.print_exc()

    # 7) Return result
    updated_credits = None
    try:
        updated_credits = wallet_balance("individual", user_email)["total"]
    except Exception:
        try:
            updated_credits = get_user_credits(user_email)
        except Exception:
            pass

    jd_hash = hash_jd(req.job_description)

    return {
        "success": True,
        "jd_hash": jd_hash,
        "sufficient": gap_result.get("sufficient", True),
        "overall_match": gap_result.get("overall_match"),
        "gaps": gap_result.get("gaps", []),
        "credits": updated_credits,
    }


@app.post("/api/save_alignment_answers")
async def api_save_alignment(req: SaveAlignmentRequest, Authorization: Optional[str] = Header(None)):
    """Persist the user's answers to gap-analysis follow-up questions.
    Keyed by JD hash so the same answers enrich CV, cover letter, and interview prep.
    """
    if not Authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    try:
        user_email = verify_bearer_token(Authorization)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

    if not req.jd_hash or not req.jd_hash.strip():
        raise HTTPException(status_code=400, detail="Missing jd_hash")

    try:
        from database import save_alignment_answers
        save_alignment_answers(user_email, req.jd_hash, req.gaps or [], req.answers or {})
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Failed to save alignment answers")

    return {"success": True}

@app.post("/api/jobsqa/signup")
async def jobsqa_signup(req: JobsQASignupRequest):
    try:
        email = req.email.strip().lower()
        password = req.password.strip()

        if jobsqa_get_user_by_email(email):
            raise HTTPException(status_code=409, detail="User already exists")

        password_hash = generate_password_hash(password)
        jobsqa_create_user(email, password_hash)

        otp = secrets.token_hex(3).upper()
        jobsqa_set_email_otp(email, otp)

        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": email,
            "subject": "JobsQA Email Verification OTP",
            "html": f"""
        <div style="font-family:Arial,sans-serif; max-width:600px">
            <h2>Welcome to JobsQA 👋</h2>
            <p>Use the OTP below to verify your JobsQA account:</p>

            <div style="
                font-size:28px;
                font-weight:bold;
                letter-spacing:6px;
                margin:20px 0;
            ">
                {otp}
            </div>

            <p>This OTP is valid for <b>10 minutes</b>.</p>

            <hr style="margin:24px 0"/>

            <p style="font-size:12px;color:#666">
                This email was sent by CVolvePro on behalf of JobsQA.<br>
                If you didn’t request this, you can safely ignore it.
            </p>
        </div>
    """
        })

        return {"success": True, "message": "OTP sent"}

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/jobsqa/verify-otp")
async def jobsqa_verify_otp(payload: dict):
    try:
        email = payload.get("email", "").strip().lower()
        otp = payload.get("otp", "").strip().upper()

        if not email or not otp:
            raise HTTPException(status_code=400, detail="Email and OTP required")

        # 1️⃣ Verify OTP (with expiry handled in DB)
        if not jobsqa_verify_email_otp(email, otp):
            raise HTTPException(status_code=400, detail="Invalid or expired OTP")

        return {"success": True}

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))




@app.post("/api/jobsqa/login")
def jobsqa_login(req: JobsQALoginRequest):
    user = None
    try:
        user = jobsqa_authenticate(req.email, req.password)
    except ValueError as e:
        if str(e) == "EMAIL_NOT_VERIFIED":
            raise HTTPException(
                status_code=403,
                detail="Email not verified. Please verify your email first."
            )

    if not user:
        raise HTTPException(
            status_code=401,
            detail="Invalid email or password"
        )

    token = create_jwt_for_jobsqa_user(user["email"])

    return {
        "success": True,
        "token": token
    }



@app.post("/api/jobsqa/generate_interview_qa")
async def jobsqa_generate_interview_qa(
    req: InterviewQARequest,
    Authorization: str = Header(...)
):
    # 1️⃣ Auth
    email = verify_bearer_token_jobsqa(Authorization)

    user = jobsqa_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    user_id = user["id"]

    # 2️⃣ Credit check (READ ONLY – NO DEDUCTION)
    credits = jobsqa_get_credits(user_id)
    if credits is None or credits < 3:
        raise HTTPException(status_code=402, detail="Insufficient credits")

    # 3️⃣ Resume extraction (safe)
    resume_text = ""
    if req.resume_base64:
        try:
            decoded = base64.b64decode(req.resume_base64)
            bio = io.BytesIO(decoded)
            bio.name = req.resume_filename or "resume.pdf"

            if hasattr(cv_generator, "extract_resume_text"):
                resume_text = cv_generator.extract_resume_text(bio)
        except Exception:
            resume_text = ""

    # 4️⃣ GENERATE INTERVIEW Q&A (STRICT)
    qa_extra = {}
    try:
        jd_h = cv_generator.hash_jd(req.job_description)
        qa_extra = get_alignment_answers(email, jd_h).get("answers", {})
    except Exception:
        pass
    try:
        qa_text = cv_generator.generate_interview_qa(
            resume_text,
            req.job_description,
            extra_context=qa_extra
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Interview Q&A generation failed: {str(e)}"
        )

    # 5️⃣ HARD VALIDATION (THIS IS CRITICAL)
    if (
        qa_text is None
        or not isinstance(qa_text, str)
        or not qa_text.strip()
    ):
        raise HTTPException(
            status_code=500,
            detail="Interview Q&A generation returned empty or invalid output"
        )

    # 6️⃣ SAVE FIRST (NO CREDIT LOSS IF SAVE FAILS)
    try:
        jobsqa_save_interview(
            user_id,
            req.resume_filename or "",
            req.job_description,
            qa_text
        )
    except Exception:
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Interview Q&A generated but failed to save. Credits not deducted."
        )

    # 7️⃣ DEDUCT CREDITS — ONLY AFTER EVERYTHING SUCCEEDS
    try:
        jobsqa_update_credits(user_id, -3, "interview_qa")
    except Exception:
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="Interview saved but credit deduction failed"
        )

    # 8️⃣ Return success
    return {
        "success": True,
        "interview_qa": qa_text,
        "credits": jobsqa_get_credits(user_id)
    }



# ----------------- NEW: Cover Letter helpers & endpoint ------------------

def _split_into_sentences(text: str):
    """
    Split text into sentences using a conservative regex.
    Keeps abbreviations naive-handled; this is best-effort.
    """
    if not text:
        return []
    text = text.strip()
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text)
    # Split on sentence boundaries: ., ?, ! followed by space and uppercase or end-of-line
    # Keep the punctuation
    parts = re.split(r'(?<=[.!?])\s+', text)
    parts = [p.strip() for p in parts if p.strip()]
    return parts

def enforce_cl_structure(text: str, min_paragraphs=3, max_paragraphs=4, max_words=450) -> str:
    """
    Defensive cover-letter normalizer:
      - aggressively convert pipe-like separators into paragraph breaks
      - treat various HTML breaks as paragraphs
      - if still a single long paragraph, split by sentence and distribute into 3-4 paragraphs
      - ensure final output has <= max_paragraphs and <= max_words (approx)
      - final safety: remove any leftover '|' characters that remain inside text
    """
    if text is None:
        return ""

    t = str(text)

    # Normalize weird unicode pipes and similar separators into ASCII pipe
    t = re.sub(r'[\u2016\u2758\u007C]+', '|', t)

    # Convert patterns like: ". |", ".|", " |—", "word|Word" into paragraph breaks
    t = re.sub(r'([.!?])\s*\|\s*', r'\1\n\n', t)                      # end-sentence + pipe -> paragraph break
    t = re.sub(r'([a-z0-9\)\]\%])\s*\|\s*([A-Z])', r'\1\n\n\2', t)    # mid-sentence pipe before capital -> paragraph
    # general fallback: any remaining ' | ' or '||' -> paragraph
    t = re.sub(r'\s*\|\s*', '\n\n', t)

    # Convert common HTML separators to paragraph breaks, strip other tags
    t = re.sub(r'<\s*(br|p|div)\s*\/?>', '\n\n', t, flags=re.I)
    t = re.sub(r'<[^>]+>', '', t)

    # Normalize CR/LF and compress >2 newlines down to exactly two
    t = t.replace('\r', '\n')
    t = re.sub(r'\n{3,}', '\n\n', t)
    t = re.sub(r'[ \t]{2,}', ' ', t).strip()

    # Build paragraphs by splitting on two-or-more newlines
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', t) if p.strip()]

    # If already in the desired range, enforce max words and return
    if min_paragraphs <= len(paragraphs) <= max_paragraphs:
        joined = "\n\n".join(paragraphs)
        words = joined.split()
        if len(words) > max_words:
            joined = " ".join(words[:max_words])
        try:
            joined = normalize_text_for_docx(joined)
        except Exception:
            pass
        joined = re.sub(r'\s*\|\s*', ' ', joined)
        return joined.strip()

    # If too few paragraphs, attempt sentence splitting
    sentences = []
    for p in paragraphs:
        sents = _split_into_sentences(p)
        if not sents:
            sents = [x.strip() for x in re.split(r',\s*', p) if x.strip()]
        sentences.extend(sents)

    # If no paragraphs existed at all, split raw text into sentences
    if not sentences:
        sentences = _split_into_sentences(t)

    # If still nothing, fallback to trimmed text (truncate if needed)
    if not sentences:
        joined = re.sub(r'\s*\|\s*', ' ', t).strip()
        words = joined.split()
        if len(words) > max_words:
            joined = " ".join(words[:max_words])
        try:
            joined = normalize_text_for_docx(joined)
        except Exception:
            pass
        return joined

    # Decide paragraph count: prefer 3 unless many sentences -> 4
    num_sentences = len(sentences)
    if num_sentences >= 9:
        target_paragraphs = min(4, max_paragraphs)
    else:
        target_paragraphs = 3 if num_sentences >= 3 else 1
    target_paragraphs = max(1, min(target_paragraphs, max_paragraphs))

    # Distribute sentences roughly evenly across paragraphs
    per_para = max(1, num_sentences // target_paragraphs)
    paras = []
    idx = 0
    for i in range(target_paragraphs):
        if i == target_paragraphs - 1:
            chunk = sentences[idx:]
        else:
            chunk = sentences[idx: idx + per_para]
        idx += len(chunk)
        para_text = " ".join(chunk).strip()
        if para_text:
            paras.append(para_text)

    # If fewer paragraphs than target, split the longest paragraph
    while len(paras) < target_paragraphs:
        longest_idx = max(range(len(paras)), key=lambda i: len(_split_into_sentences(paras[i])))
        long_sents = _split_into_sentences(paras[longest_idx])
        if len(long_sents) < 2:
            break
        half = len(long_sents) // 2
        first = " ".join(long_sents[:half]).strip()
        second = " ".join(long_sents[half:]).strip()
        paras[longest_idx:longest_idx+1] = [first, second]

    # Final join and enforce max words while preserving paragraph boundaries
    joined = "\n\n".join([p.strip() for p in paras if p.strip()])
    words = joined.split()
    if len(words) > max_words:
        joined = " ".join(words[:max_words])
        sents = _split_into_sentences(joined)
        if sents:
            target = min(3, max_paragraphs)
            per = max(1, len(sents) // target)
            new_paras = []
            i = 0
            for j in range(target):
                if j == target - 1:
                    chunk = sents[i:]
                else:
                    chunk = sents[i:i+per]
                i += len(chunk)
                new_paras.append(" ".join(chunk).strip())
            joined = "\n\n".join([p for p in new_paras if p])

    # Final normalizer + remove stray pipes
    try:
        joined = normalize_text_for_docx(joined)
    except Exception:
        pass
    joined = re.sub(r'\s*\|\s*', ' ', joined)
    joined = re.sub(r'\n{3,}', '\n\n', joined).strip()
    return joined

@app.post("/api/generate_cl")
async def api_generate_cl(req: GenerateCLRequest, Authorization: Optional[str] = Header(None)):
    if not Authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    try:
        user_email = verify_bearer_token(Authorization)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

    if not req.job_description or not req.job_description.strip():
        raise HTTPException(status_code=400, detail="Missing job_description")

    resume_text = ""
    if req.resume_base64:
        try:
            decoded = base64.b64decode(req.resume_base64)
            bio = io.BytesIO(decoded)
            bio.name = req.resume_filename or "resume.pdf"
            if hasattr(cv_generator, "extract_resume_text"):
                resume_text = cv_generator.extract_resume_text(bio)
        except Exception:
            resume_text = ""

    try:
        cl_extra = (req.extras or {}).get("extra_context", {})
        if not cl_extra and req.job_description:
            try:
                jd_h = cv_generator.hash_jd(req.job_description)
                cl_extra = get_alignment_answers(user_email, jd_h).get("answers", {})
            except Exception:
                pass
        # generate via available function
        if hasattr(cv_generator, "generate_cover_letter"):
            cover_letter = cv_generator.generate_cover_letter(resume_text, req.job_description, language=req.language or "English", extra_context=cl_extra)
        else:
            cover_letter = getattr(cv_generator, "generate_cl")(resume_text, req.job_description, language=req.language or "English", extra_context=cl_extra)
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Cover letter generation failed")

    try:
        if isinstance(cover_letter, str):
            cl_text_raw = cover_letter.replace("**", "")
        else:
            cl_text_raw = str(cover_letter)
    except Exception:
        cl_text_raw = str(cover_letter)

    # Server-side: replace pipes with paragraph breaks and enforce CL paragraph structure + length
    try:
        cl_text_structured = enforce_cl_structure(cl_text_raw, min_paragraphs=3, max_paragraphs=4, max_words=450)
    except Exception:
        # fallback: replace pipes with double newlines so UI sees paragraphs
        try:
            cl_text_structured = re.sub(r'\s*\|\s*', '\n\n', cl_text_raw)
        except Exception:
            cl_text_structured = cl_text_raw

    try:
        spend_credits("individual", user_email, "Cover Letter", amount=2)
    except Exception:
        traceback.print_exc()

    current_credits = None
    try:
        current_credits = wallet_balance("individual", user_email)["total"]
    except Exception:
        try:
            current_credits = get_user_credits(user_email)
        except Exception:
            current_credits = None

    # --- Try to produce a real .docx for the cover letter (best-effort) ---
    docx_b64 = None
    file_ext = None
    file_mime = None
    try:
        # Normalize for docx builder
        try:
            normalized = normalize_text_for_docx(cl_text_structured or "")
        except Exception:
            normalized = cl_text_structured or ""

        # create_word_document is expected to return an io.BytesIO with a valid .docx
        word_buf = None
        if 'create_word_document' in globals():
            word_buf = create_cover_letter_docx(normalized)
        elif hasattr(cv_generator, "create_word_document"):
            word_buf = getattr(cv_generator, "create_word_document")(normalized)
        else:
            word_buf = None

        if word_buf:
            file_bytes = word_buf.getvalue()
            file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_ext = "docx"
            docx_b64 = base64.b64encode(file_bytes).decode("utf-8")
    except Exception:
        # don't fail the whole request if docx creation fails
        traceback.print_exc()
        docx_b64 = None

    resp = {"success": True, "cover_letter": cl_text_structured, "credits": current_credits}
    if docx_b64:
        resp.update({"docx_base64": docx_b64, "file_mime": file_mime, "file_ext": file_ext})
    

    return resp


from fastapi import Request, Header, HTTPException

from pydantic import BaseModel

class CheckoutRequest(BaseModel):
    country_code: Optional[str] = None

@app.post("/api/jobsqa/create_checkout")
async def jobsqa_create_checkout(
    checkout_data: CheckoutRequest,
    request: Request,
    Authorization: str = Header(...)
):
    email = verify_bearer_token_jobsqa(Authorization)
    user = jobsqa_get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    # Get country from frontend (since not using Cloudflare)
    country = checkout_data.country_code or ""
    is_india = (country == "IN")


    if is_india:
        amount = 89900     # 899 in paise
        currency = "inr"
    else:
        amount = 1499     # $14.99 in cents
        currency = "usd"

    checkout_url = create_jobsqa_checkout_session(
        user_email=email,
        amount=amount,
        currency=currency,
        success_url="https://jobsqa.com/payment-success?session_id={CHECKOUT_SESSION_ID}",
        cancel_url="https://jobsqa.com"
    )

    return {
        "checkout_url": checkout_url
    }








from fastapi import Request, Response

@app.post("/api/jobsqa/webhook")
async def jobsqa_stripe_webhook(request: Request):
    """
    Stripe webhook handler for JobsQA payments.
    Automatically adds credits when payment succeeds.
    
    Events handled:
    - checkout.session.completed: When customer completes payment
    """
    try:
        # Get raw body for signature verification
        payload = await request.body()
        sig_header = request.headers.get("stripe-signature")
        
        if not STRIPE_JOBSQA_WEBHOOK_SECRET:
            logging.error("❌ STRIPE_JOBSQA_WEBHOOK_SECRET not configured")
            raise HTTPException(status_code=500, detail="Webhook secret not configured")
        
        # Verify webhook signature
        try:
            event = stripe.Webhook.construct_event(
                payload, sig_header, STRIPE_JOBSQA_WEBHOOK_SECRET
            )
        except ValueError as e:
            logging.error(f"❌ Invalid webhook payload: {str(e)}")
            raise HTTPException(status_code=400, detail="Invalid payload")
        except stripe.error.SignatureVerificationError as e:  # type: ignore[reportAttributeAccessIssue]
            logging.error(f"❌ Invalid webhook signature: {str(e)}")
            raise HTTPException(status_code=400, detail="Invalid signature")
        
        # Log the event type
        event_type = event.get("type")
        logging.info(f"📥 Received Stripe webhook: {event_type}")
        
        # Handle checkout.session.completed event
        if event_type == "checkout.session.completed":
            session = event["data"]["object"]
            
            # Log session details
            customer_email = session.get("customer_email")
            metadata = session.get("metadata", {})
            session_id = session.get("id")
            
            logging.info(f"💳 Payment completed: session={session_id}, email={customer_email}, metadata={metadata}")
            
            # Check if this is a JobsQA payment
            if metadata.get("service") == "jobsqa":
                try:
                    # Call the payment handler to add credits
                    success = handle_jobsqa_payment(session)
                    
                    if success:
                        logging.info(f"✅ Credits successfully added for {customer_email}")
                    else:
                        logging.error(f"❌ Failed to add credits for {customer_email}")
                    
                except Exception as e:
                    logging.error(f"❌ Error handling JobsQA payment: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    # Don't return 500 to Stripe - we got the payment notification
            else:
                logging.info(f"⏭️  Skipping non-JobsQA payment: service={metadata.get('service')}")
        
        else:
            logging.info(f"⏭️  Ignoring event type: {event_type}")
        
        # Always return 200 to Stripe
        return {"success": True, "event_type": event_type}
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ Webhook error: {str(e)}")
        import traceback
        traceback.print_exc()
        # Return 200 anyway to prevent Stripe from retrying
        return {"success": False, "error": str(e)}





if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)))
