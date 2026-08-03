import streamlit as st
import time
import json
import re
import os
from io import BytesIO
from docx import Document
import plotly.graph_objects as go
from reportlab.lib.pagesizes import letter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import secrets
from database import register_user
from payment import create_checkout_session
from auth import get_current_user
import phonenumbers
import pycountry
from phonenumbers import parse as pn_parse, is_valid_number, format_number, PhoneNumberFormat
import resend
import urllib.parse
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash
load_dotenv()


# Import custom modules
from database import init_db, get_user_special_discount, get_user_data, save_user_session, get_business_plan_info, get_user_credits, get_business_credits, get_db_connection, create_business_user, get_business_user, authenticate_business_user, save_business_payment, save_payment, payment_exists, register_user, verify_user_email, set_email_otp, verify_email_otp, save_cv_generation, record_credit_usage, reset_credits_if_expired, save_alignment_answers, get_alignment_answers
from auth import authenticate_user, logout_user, get_current_user, hash_password
from payment import process_payment, check_subscription, apply_discount_code, create_checkout_session
from cv_generator import generate_cv,recommend_jobs_from_resume_ai, generate_cover_letter, extract_resume_text, analyze_cv_ats_score, export_interview_qa, analyze_cv_jd_gaps, hash_jd, export_cover_letter
import job_aggregator as ja
from templates import apply_template
from utils import optimize_keywords, enforce_page_limit, get_gemini_response, get_all_country_dial_codes
from interview_module import show_interview_practice_page, _init_session as _init_interview_session
import pricing
import extra_streamlit_components as stx
from session_auth import (
    issue_session,
    validate_session,
    revoke_session,
    purge_expired_sessions,
    SESSION_COOKIE_NAME,
)
import voucher_engine
from voucher_engine import (
    is_admin as is_voucher_admin,
    generate_voucher,
    redeem_voucher,
    list_vouchers as list_all_vouchers,
    list_voucher_redemptions,
    revoke_voucher as admin_revoke_voucher,
    DEFAULT_PLAN as VOUCHER_DEFAULT_PLAN,
)
from credit_engine import (
    wallet_balance,
    spend_credits,
    has_enough,
    ats_charge_or_free,
    purchase_plan,
    purchase_pack,
    get_credit_packs,
    recent_transactions,
    backfill_wallets,
    can_use_f2f,
)



BUSINESS_PLANS = {
    "Corporate Starter": {
        "credits": 500,
        "price": 149.99,
        "duration": "3 months"
    },

    "Corporate Growth": {
        "credits": 1000,
        "price": 299.00,
        "duration": "3 months"
    },

    "Corporate Pro": {
        "credits": 2500,
        "price": 449.00,
        "duration": "6 months"
    },

    "Corporate Plus": {
        "credits": 5000,
        "price": 699.00,
        "duration": "6 months"
    },

    "Corporate Advanced": {
        "credits": 7500,
        "price": 899.00,
        "duration": "1 year"
    },

    "Corporate Enterprise": {
        "credits": 10000,
        "price": 999.00,
        "duration": "1 year"
    }
}

def add_home_button():
    home_href = "https://cvolvepro.com"
    st.sidebar.markdown(
        f'<a href="{home_href}" target="_self"><button style="width:100%">🏠 Home</button></a>',
        unsafe_allow_html=True
    )

resend.api_key = os.getenv("RESEND_API_KEY")
FROM_EMAIL = os.getenv("APP_FROM_EMAIL", "onboarding@resend.dev")

def _handle_tracking_hop():
    """If URL has ?trk=..., briefly show an interstitial and bounce to ?next=..."""
    # Works with both old/new Streamlit query APIs
    params = st.query_params if hasattr(st, "query_params") else getattr(st, "experimental_get_query_params")()

    def _get(key):
        v = params.get(key)
        return v[0] if isinstance(v, list) else v

    trk  = _get("trk")
    nxt  = _get("next")

    if trk and nxt:
        st.markdown("#### Redirecting…")
        st.caption(f"Tracking event: {trk}")
        # simple 0-sec meta refresh
        st.markdown(f'<meta http-equiv="refresh" content="0; url={nxt}">', unsafe_allow_html=True)
        st.stop()  # do not render rest of the app on this hop

# call it immediately so it catches both click and success hops
_handle_tracking_hop()

def _qp_get(key: str, default=""):
    v = st.query_params.get(key, default)
    return v[0] if isinstance(v, list) else v

def handle_stripe_return_globally():
    if _qp_get("service") == "jobsqa":
        return
    import stripe
    from database import save_payment, payment_exists as db_payment_exists

    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")

    # Legacy plan names → current plans (for old Stripe metadata / special offer)
    LEGACY_PLAN_MAP = {
        "Premium": "Career Pro",
        "Premium + Premium Classic": "Interview Pro",
    }
    PACK_CREDITS_TO_NAME = {p["credits"]: p["name"] for p in pricing.PACKS}

    success    = (_qp_get("success", "").lower() == "true")
    typ        = _qp_get("type", "")
    session_id = _qp_get("session_id", "")
    credits_qp = int(_qp_get("credits", "0") or 0)
    plan_qp    = _qp_get("plan", "")

    if not (success and session_id and typ in ("subscription", "credits", "business")):
        return

    # prevent double-processing on reruns
    if "processed_sessions" not in st.session_state:
        st.session_state.processed_sessions = set()
    if session_id in st.session_state.processed_sessions:
        return

    try:
        sess = stripe.checkout.Session.retrieve(session_id)
        if sess.get("payment_status") != "paid":
            return

        md = sess.get("metadata") or {}
        user_email = (md.get("user_email") or "").strip().lower()
        if not user_email:
            st.warning("Missing user email in Stripe metadata; cannot credit.")
            return

        # DB-level idempotency: never credit the same Stripe session twice.
        try:
            if db_payment_exists(session_id):
                st.session_state.processed_sessions.add(session_id)
                return
        except Exception:
            pass

        amount_paid = (sess.get("amount_total") or 0) / 100.0

        if typ == "subscription":
            plan = (md.get("plan") or plan_qp or "Career Pro").strip()
            plan = LEGACY_PLAN_MAP.get(plan, plan)
            if plan not in pricing.PLANS:
                st.warning(f"Unknown plan '{plan}'; payment recorded but not activated.")
                save_payment(user_email, amount_paid, "subscription", session_id, credits_purchased=0)
                return
            res = purchase_plan("individual", user_email, plan, stripe_session_id=session_id)
            if res.get("ok"):
                save_payment(user_email, amount_paid, "subscription", session_id,
                             credits_purchased=res["credits"])
                st.success(f"🎉 {plan} active for {user_email}. {res['credits']} credits added.")

        elif typ == "credits":
            pack_name = (md.get("pack") or "").strip()
            if not pack_name:
                pack_name = PACK_CREDITS_TO_NAME.get(int(md.get("credits") or credits_qp or 0))
            if not pack_name:
                save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=0)
                st.warning("Payment succeeded but no pack name was provided.")
                return
            res = purchase_pack("individual", user_email, pack_name, stripe_session_id=session_id)
            if res.get("ok"):
                save_payment(user_email, amount_paid, "credits", session_id,
                             credits_purchased=res["credits"])
                st.success(f"🎉 {res['pack']} ({res['credits']} credits) added for {user_email}.")
            else:
                save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=0)
                st.error(f"Could not add pack: {res.get('reason')}")

        # ======================================================
        # BUSINESS PLAN HANDLER
        # ======================================================

        elif typ == "business":
            from database import save_business_payment

            plan_name = md.get("plan_name", "Starter")
            if plan_name not in pricing.CORPORATE_PLANS:
                st.warning(f"Unknown business plan '{plan_name}'; payment recorded but not activated.")
                save_business_payment(user_email, amount_paid, "business_plan", session_id, credits_purchased=0)
                return

            res = purchase_plan("business", user_email, plan_name, stripe_session_id=session_id)
            if res.get("ok"):
                save_business_payment(user_email, amount_paid, "business_plan", session_id,
                                      credits_purchased=res["credits"])
                st.success(
                    f"""
                    🎉 Business Plan Activated Successfully

                    Plan: {plan_name}

                    Credits Added: {res['credits']}
                    """
                )
            else:
                save_business_payment(user_email, amount_paid, "business_plan", session_id, credits_purchased=0)
                st.error(f"Could not activate business plan: {res.get('reason')}")

    except Exception as e:
        st.error(f"Stripe verification failed: {e}")
    finally:
        st.session_state.processed_sessions.add(session_id)
        try:
            st.query_params.clear()
        except:
            pass


import stripe
stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")

import traceback

def _sanitize_db_text(s: str) -> str:
    """Make text safe for PostgreSQL INSERT (strip NULs, cap size)."""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    return s.replace("\x00", "")[:500_000]  # keep first ~500k chars


# Initialize database
try:
    init_db()
    from database import seed_discount_codes
    seed_discount_codes()
    backfill_wallets()
except Exception as e:
    from database import get_db_config_summary
    cfg_summary = get_db_config_summary()
    st.error("🚨 **Database Connection Error**")
    st.warning(
        "Could not connect to the PostgreSQL database.\n\n"
        "**If you deployed on Streamlit Cloud:**\n"
        "1. Go to your app dashboard on Streamlit Cloud.\n"
        "2. Click **Manage app** (lower right) -> **Settings** -> **Secrets**.\n"
        "3. Paste your database credentials into **Secrets**:\n\n"
        "```toml\n"
        'DATABASE_URL = "postgresql://user:password@your-db-host.com:5432/cvolvepro?sslmode=require"\n'
        "```\n"
        "**OR**\n\n"
        "```toml\n"
        'DB_HOST = "your-db-host.com"\n'
        'DB_PORT = "5432"\n'
        'DB_NAME = "cvolvepro"\n'
        'DB_USER = "postgres"\n'
        'DB_PASSWORD = "your_db_password"\n'
        'DB_SSLMODE = "require"\n'
        "```"
    )
    with st.expander("🔍 Detected Connection Settings & Error Log"):
        st.write("**Resolved Connection Config:**", cfg_summary)
        st.code(str(e))
    st.stop()

def get_allowed_ai_models_for_user():
    """
    Returns the AI model options the current user is allowed to see
    based on their active subscription.
    - Premium                 -> ["Premium"]
    - Premium + Premium Classic -> ["Premium", "Premium Classic"]
    - No active subscription  -> both (uses credits)
    """
    try:
        user = st.session_state.get("user_data")
        if not user:
            return ["Premium", "Premium Classic"]

        sub = check_subscription(user["email"])
        if not sub:
            # No sub – allow both; they’ll spend credits
            return ["Premium", "Premium Classic"]

        plan = (sub.get("plan") or "").strip()
        if "Premium + Premium Classic" in plan:
            return ["Premium", "Premium Classic"]
        # Default to Premium-only
        return ["Premium"]
    except Exception:
        # On any error, fall back to both (don’t block usage)
        return ["Premium", "Premium Classic"]


# Page config
st.set_page_config(
    page_title="CVOLVE PRO - AI-Powered Resume Optimization",
    page_icon="logo.jpeg",
    layout="wide",
    initial_sidebar_state="expanded"
)

add_home_button()

# Load custom CSS
with open("styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Initialize session state
if 'user_data' not in st.session_state:
    st.session_state.user_data = None
if 'cv_preview' not in st.session_state:
    st.session_state.cv_preview = None
if 'cv_optimization_metadata' not in st.session_state:
    st.session_state.cv_optimization_metadata = None
if 'auto_save' not in st.session_state:
    st.session_state.auto_save = {}
if 'selected_template' not in st.session_state:
    st.session_state.selected_template = "professional"
# CV↔JD alignment state machine (Phase 2)
if 'alignment_stage' not in st.session_state:
    st.session_state.alignment_stage = "idle"   # idle → questions → generating
if 'alignment_gaps' not in st.session_state:
    st.session_state.alignment_gaps = None
if 'alignment_jd_hash' not in st.session_state:
    st.session_state.alignment_jd_hash = None
if 'target_job_title' not in st.session_state:
    st.session_state.target_job_title = None
if 'target_job_company' not in st.session_state:
    st.session_state.target_job_company = None
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""
if 'optimizer_mode' not in st.session_state:
    st.session_state.optimizer_mode = "smart"  # "smart" or "manual"
if 'manual_jd' not in st.session_state:
    st.session_state.manual_jd = ""
if 'manual_title' not in st.session_state:
    st.session_state.manual_title = ""
if 'manual_company' not in st.session_state:
    st.session_state.manual_company = ""
if 'manual_ats_analysis' not in st.session_state:
    st.session_state.manual_ats_analysis = None
if 'cover_letter_content' not in st.session_state:
    st.session_state.cover_letter_content = None
if 'interview_qa_content' not in st.session_state:
    st.session_state.interview_qa_content = None
if 'cv_pdf_bytes' not in st.session_state:
    st.session_state.cv_pdf_bytes = None
if 'cv_docx_bytes' not in st.session_state:
    st.session_state.cv_docx_bytes = None


def auto_save_progress():
    """Auto-save user progress"""
    if st.session_state.user_data and st.session_state.auto_save:
        try:
            save_user_session(st.session_state.user_data['email'], st.session_state.auto_save)
        except Exception as e:
            # Silently handle auto-save errors to not interrupt user flow
            pass


_cvolve_cookie_mgr = None


def _cookie_manager():
    """Lazily-instantiated singleton CookieManager component."""
    global _cvolve_cookie_mgr
    if _cvolve_cookie_mgr is None:
        _cvolve_cookie_mgr = stx.CookieManager(key="cvolve_cookie_mgr")
    return _cvolve_cookie_mgr


def _persist_login_cookie(raw_token):
    """Write the session token to a 30-day browser cookie (refresh-safe)."""
    if not raw_token:
        return
    try:
        _cookie_manager().set(
            SESSION_COOKIE_NAME, raw_token, max_age=60 * 60 * 24 * 30
        )
    except Exception:
        pass


def restore_session_from_cookie():
    """Re-hydrate the logged-in session from the persisted auth cookie.

    Streamlit wipes st.session_state on every refresh, which is why users were
    logged out. This reads the cookie set at login, validates the token against
    the auth_sessions table, and repopulates the session before the login gate.
    """
    if st.session_state.get("user_data"):
        return
    try:
        cm = _cookie_manager()
        cookies = cm.get_all()
        if cookies is None:
            # The cookie component hasn't mounted yet this session; one extra
            # run lets the browser deliver the cookie for reading.
            if "cvolve_cookie_retried" not in st.session_state:
                st.session_state.cvolve_cookie_retried = True
                st.rerun()
            return
        raw = cookies.get(SESSION_COOKIE_NAME)
        if not raw:
            return
        data = validate_session(raw)
        if not data:
            try:
                cm.delete(SESSION_COOKIE_NAME)
            except Exception:
                pass
            return
        user = data["user"]
        st.session_state.user_data = user
        st.session_state.account_type = data["account_type"]
        if data["account_type"] == "business":
            st.session_state.business_user = user
            st.session_state.business_logged_in = True
    except Exception:
        pass


def main():

    handle_stripe_return_globally()

    # Lazy-cleanup expired persisted sessions (once per process)
    if not st.session_state.get("_sessions_purged"):
        purge_expired_sessions()
        st.session_state._sessions_purged = True

    # ✅ Restore login from cookie (refresh-safe persistence)
    restore_session_from_cookie()

    # ✅ Persist page navigation (login/register/main)
    if "page" not in st.session_state:
        st.session_state.page = "login"  # Default page

    if "portal" not in st.session_state:
        st.session_state.portal = "individual"

    # Interview module session state
    _init_interview_session()

    # ✅ Show Register Page if user clicked register
    if st.session_state.page == "register":
        show_register_page()
        return  # Stop here after rendering register page


    # ✅ If user is not logged in
    if not st.session_state.get("user_data"):

        if st.session_state.portal == "business":
            show_business_login_page()
        else:
            show_login_page()

        return
    
    # Auto-save progress only when user is logged in and has data to save
    if st.session_state.user_data and st.session_state.auto_save:
        auto_save_progress()
    
    # Header
    st.markdown(f"""
    <div class="header">
        <h1 style="display:inline-block; vertical-align:middle; margin:0;">CVOLVE PRO</h1>
        <p>Transform your resume into an ATS-optimized masterpiece</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Authentication
    current_user = get_current_user()
    if not current_user:
        show_login_page()
        return
    
    st.session_state.user_data = current_user
    # NEW: Just-in-time monthly reset so UI shows the correct balance
    try:
        reset_credits_if_expired(current_user['email'])
    except Exception:
        pass


    # If user navigated to Billing from sidebar, render only Billing and return
    if st.session_state.get("page") == "billing":
        st.markdown("## 💳 Billing")
        show_billing_page()
        if st.button("⬅ Back"):
            st.session_state.page = "home"   # any non-"billing" value works
            st.rerun()
        return

    # Admin-only Voucher management page
    if st.session_state.get("page") == "admin_vouchers":
        if not is_voucher_admin(current_user.get("email", "")):
            st.error("You don't have access to this page.")
            return
        st.markdown("## 🎫 Admin · Voucher Management")
        show_admin_vouchers_page(current_user.get("email", ""))
        if st.button("⬅ Back"):
            st.session_state.page = "home"
            st.rerun()
        return

    # Sidebar
    with st.sidebar:
        st.markdown(
        f"👋 Welcome, {current_user.get('name', current_user.get('company_name', 'User'))}"
    )
        
        # User credits/subscription status
        email = current_user['email'].strip().lower()

        if st.session_state.get("account_type") == "business":
            credits = get_business_credits(email)
            try:
                credits = wallet_balance("business", email)["total"]
            except Exception:
                pass
            subscription = None
        else:
            credits = get_user_credits(email)
            try:
                credits = wallet_balance("individual", email)["total"]
            except Exception:
                pass
            subscription = check_subscription(email)

        # 🔹 Set AI model options based on subscription
        plan = (subscription['plan'] if subscription else "") or ""
        allowed_models = ["Premium", "Premium Classic"] if "Premium + Premium Classic" in plan or plan == "" else ["Premium"]

        # Save in session; rerun if changed so dropdown updates instantly
        if st.session_state.get("ai_model_options") != allowed_models:
            st.session_state["ai_model_options"] = allowed_models
            if st.session_state.get("ai_model") == "openai" and "Premium Classic" not in allowed_models:
                st.session_state["ai_model"] = "gemini"
            st.rerun()

        
        if st.session_state.get("account_type") == "business":

            plan_info = get_business_plan_info(email)

            if plan_info and plan_info["current_plan"]:

                st.success(
                    f"🏢 {plan_info['current_plan']} Plan Active"
                )

                if plan_info["plan_expiry"]:

                    st.caption(
                        f"Valid Until: {plan_info['plan_expiry'].strftime('%d %b %Y')}"
                    )

            else:

                st.warning("⏳ No active business plan")

        else:

            if subscription:

                if subscription['plan'] == "Free":

                    st.info("🆓 Free Trial (1 Month)")

                else:

                    st.success(
                        f"✅ {subscription['plan']} Plan Active"
                    )

            else:

                st.warning("⏳ No active plan")

        # 💎 Credits — ALWAYS visible
        st.info(f"💎 Credits Available: {credits}")
            
        if st.button("🔄 Buy More Credits"):
            st.session_state.page = "billing"
            st.rerun()

        # Admin-only: Voucher management
        if is_voucher_admin(current_user.get("email", "")):
            if st.sidebar.button("🎫 Admin · Vouchers"):
                st.session_state.page = "admin_vouchers"
                st.rerun()


        if st.sidebar.button("Logout"):

            # Revoke persisted session + clear cookie so refresh stays logged out
            try:
                raw = _cookie_manager().get(SESSION_COOKIE_NAME)
                if raw:
                    revoke_session(raw)
                _cookie_manager().delete(SESSION_COOKIE_NAME)
            except Exception:
                pass

            # Individual session
            st.session_state.logged_in = False
            st.session_state.user_data = None

            # Business session
            st.session_state.business_logged_in = False
            st.session_state.business_user = None
            st.session_state.business_email = None
            st.session_state.business_company = None

            # Global
            st.session_state.account_type = None

            st.rerun()
            
        
        # ✅ Set default template to Professional Classic
        st.session_state.selected_template = "professional"
        
        # ✅ Always include default sections (all enabled)
        sections = {
            "Professional Summary": True,
            "Key Skills": True,
            "Work Experience": True,
            "Education": True,
            "Certifications": True,
            "Projects": True,
            "Awards": False,
            "Languages": False,
            "Hobbies": False
        }
        
        st.session_state.auto_save['sections'] = sections
        
        # Sidebar Resume Upload (global)
        st.markdown("---")
        st.markdown("### 📄 Resume / CV")
        sidebar_cv = st.file_uploader(
            "Upload CV (PDF/DOCX)",
            type=["pdf", "docx"],
            help="Upload your master resume to use for matching and optimization",
            key="sidebar_cv_uploader"
        )
        if sidebar_cv:
            stored = st.session_state.get("uploaded_resume")
            if stored is None or stored.name != sidebar_cv.name or stored.size != sidebar_cv.size:
                st.session_state.uploaded_resume = sidebar_cv
                st.rerun()

        current_sidebar_cv = st.session_state.get("uploaded_resume")
        if current_sidebar_cv:
            st.caption(f"✅ Active CV: **{current_sidebar_cv.name}**")
        else:
          st.caption("ℹ️ No CV loaded yet")

        # Quick links
        st.markdown("---")
        with st.sidebar.expander("📚 How It Works"):
            st.markdown("""
            1. Upload your resume (PDF/DOCX)  
            2. Paste the job description  
            3. Choose your sections & template  
            4. Click ‘Generate Optimized CV’  
            5. Download your resume or cover letter  
            """)

            with st.sidebar.expander("🔒 Privacy Policy"):
                st.markdown("""
                - Your data is processed securely  
                - Resumes and job descriptions are not stored  
                - No personal info is shared with third parties  
                """)

    # Full-page Interview Practice mode (entered from the CV optimization flow)
    if st.session_state.get("page") == "interview":
        _render_interview_full_page()
        return

    # Main content
    tab1, tab3, tab4, tab5 = st.tabs(["🚀 Smart Job Match & Optimizer", "📊 Analytics", "💳 Billing", "🎤 Interview Practice"])

    with tab1:
        show_smart_job_match_page()

    with tab3:
        show_analytics_page()

    with tab4:
        show_billing_page()

    with tab5:
        show_interview_practice_page(
            check_access_fn=check_user_access,
            deduct_credits_fn=deduct_user_credits,
            extract_resume_fn=extract_resume_text,
            export_qa_fn=export_interview_qa,
        )

def _apply_domain_label(url: str) -> str:
    """Return a short destination-domain label for the apply button."""
    if not url:
        return ""
    try:
        from urllib.parse import urlparse
        host = urlparse(url).netloc.lower()
        if host.startswith("www."):
            host = host[4:]
        return host[:40]
    except Exception:
        return ""


def _render_job_results_unified(result, resume_text):
    """Render the search result stored in session with 1-click optimization button."""
    if not result:
        return

    status = result["status"]
    counts = result["counts"]
    reason = result["empty_reason"]

    badge = {"ok": "✓", "cached": "✓ cached", "stale": "~ stale", "error": "✕", "auth": "key?"}
    src_line = " · ".join(f"{s.title()} {badge.get(v, v)}" for s, v in status.items())
    if src_line:
        st.caption(f"Sources: {src_line}")

    jobs = result["jobs"]
    if not jobs:
        if reason == "unreachable":
            st.error("Couldn't reach the job sources right now — **no credit was charged**. Please try again.")
        elif reason == "filtered_out":
            st.info("No jobs matched your work-type filters. Try loosening them.")
        else:
            st.info("No jobs found for that search. Try a different title or location.")
        return

    note = "" if resume_text else " — upload a résumé above for match scores & 1-click optimization"
    st.success(f"Showing {counts['shown']} job(s){note}.")

    email = st.session_state.user_data["email"]

    for i, j in enumerate(jobs):
        with st.container(border=True):
            j_title = getattr(j, "title", "") if not isinstance(j, dict) else j.get("title", "")
            j_company = getattr(j, "company", "") if not isinstance(j, dict) else j.get("company", "")
            j_location = getattr(j, "location", "") if not isinstance(j, dict) else j.get("location", "")
            j_match_score = getattr(j, "match_score", None) if not isinstance(j, dict) else j.get("match_score")
            j_remote_type = getattr(j, "remote_type", "") if not isinstance(j, dict) else j.get("remote_type", "")
            j_job_type = getattr(j, "job_type", "") if not isinstance(j, dict) else j.get("job_type", "")
            j_salary = getattr(j, "salary", None) if not isinstance(j, dict) else j.get("salary")
            j_seniority = getattr(j, "seniority", None) if not isinstance(j, dict) else j.get("seniority")
            j_posted_date = getattr(j, "posted_date", "") if not isinstance(j, dict) else j.get("posted_date", "")
            j_url = getattr(j, "url", "") if not isinstance(j, dict) else j.get("url", "")
            j_desc = getattr(j, "description", "") if not isinstance(j, dict) else j.get("description", "")
            j_why_matched = getattr(j, "why_matched", []) if not isinstance(j, dict) else j.get("why_matched", [])

            st.subheader(j_title or "Job Position")
            st.markdown(f"**{j_company or '—'}** · {j_location or '—'}")
            chips = []
            if j_match_score is not None:
                chips.append(f"🎯 {j_match_score}% match")
            if j_remote_type:
                chips.append(f"🌍 {j_remote_type}")
            if j_job_type and j_job_type != "—":
                chips.append(f"💼 {j_job_type}")
            if j_salary:
                chips.append(f"💰 {j_salary}")
            if j_seniority:
                chips.append(f"📈 {j_seniority}")
            if j_posted_date:
                chips.append(f"🗓 {j_posted_date}")
            st.caption(" · ".join(chips))

            if j_why_matched:
                with st.expander("💡 Why this job?"):
                    for bullet in j_why_matched:
                        st.markdown(f"- {bullet}")

            col_opt, col_src = st.columns([2, 3])
            with col_opt:
                btn_key = f"opt_job_btn_{i}"
                if st.button("⚡ Optimize CV for this Job", key=btn_key, type="primary"):
                    if not resume_text:
                        st.warning("⚠️ Please upload your resume above before clicking Optimize.")
                    else:
                        if not check_user_access(required_credits=3, feature="CV"):
                            st.error("⚠️ Insufficient credits. You need 3 credits to run Gap Analysis & CV Optimization.")
                        else:
                            st.session_state.active_job_url = j_url

                            # Extract full job description if possible
                            from job_aggregator import fetch_full_job_description
                            full_jd = None
                            if j_url:
                                with st.spinner("📄 Extracting full job description from source…"):
                                    full_jd = fetch_full_job_description(j_url)

                            st.session_state.job_description = full_jd or j_desc or f"Role: {j_title} at {j_company}"
                            st.session_state.target_job_title = j_title
                            st.session_state.target_job_company = j_company

                            jd_h = hash_jd(st.session_state.job_description)
                            saved = get_alignment_answers(email, jd_h)
                            if saved and saved.get("answers"):
                                st.session_state.alignment_stage = "generating"
                                st.session_state.alignment_gaps = None
                                st.session_state.alignment_jd_hash = jd_h
                            else:
                                with st.spinner("🔍 Analyzing your CV against this job description…"):
                                    try:
                                        gap_result = analyze_cv_jd_gaps(
                                            resume_text, st.session_state.job_description,
                                            language=st.session_state.get("selected_language", "English")
                                        )
                                    except Exception:
                                        gap_result = {"sufficient": True, "overall_match": None, "gaps": []}

                                if gap_result.get("sufficient") or not gap_result.get("gaps"):
                                    st.session_state.alignment_stage = "generating"
                                    st.session_state.alignment_gaps = None
                                    st.session_state.alignment_jd_hash = jd_h
                                else:
                                    st.session_state.alignment_stage = "questions"
                                    st.session_state.alignment_gaps = gap_result
                                    st.session_state.alignment_jd_hash = jd_h
                            st.rerun()

                if j.url:
                    domain = _apply_domain_label(j.url)
                    btn_label = f"View / Apply on {domain} ↗" if domain else "View / Apply on Source ↗"
                    st.link_button(btn_label, j.url)

            with col_src:
                also = f" (also on {', '.join(j.also_on)})" if j.also_on else ""
                st.caption(f"via {j.source_name}{also}")


def show_login_page():
    st.markdown("## 🔐 Login to CVOLVE PRO")
    colp1, colp2 = st.columns(2)

    with colp1:
        if st.button("👤 Individual Portal"):
            st.session_state.portal = "individual"
            st.rerun()

    with colp2:
        if st.button("🏢 Business Portal"):
            st.session_state.portal = "business"
            st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### Email Login")
        email = st.text_input("Email Address", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")

        # ── Buttons side-by-side ───────────────────────────────────────────────
        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            login_clicked = st.button("🔑 Login", key="login_button", use_container_width=True)
        with btn_col2:
            register_clicked = st.button("🆕 Register For Free", key="register_button", use_container_width=True)
        # ───────────────────────────────────────────────────────────────────────

        if login_clicked:
            if email.strip() and password.strip():
                email_norm = email.strip().lower()
                existing = get_user_data(email_norm)

                if not existing:
                    st.error("No account found for this email. Please register first.")
                    return

                if not existing.get("is_verified", False):
                    st.error("Your email is not verified. Please complete registration with OTP.")
                    return

                user = authenticate_user(email_norm, password, "email")
                if user:
                    st.session_state.user_data = user
                    st.session_state.account_type = "individual"
                    _persist_login_cookie(issue_session(email_norm, "individual"))
                    st.success("✅ Login successful!")
                    st.rerun()
                else:
                    st.error("Incorrect password.")
            else:
                st.error("Please enter both email and password")

        if register_clicked:
            st.session_state.page = "register"  # ✅ Persist state
            st.rerun()

    # (Optional) keep the divider if you want extra spacing
    # st.markdown("---")


def show_business_login_page():

    st.markdown("# 🏢 Business Portal")

    tab1, tab2 = st.tabs([
        "Business Login",
        "Business Register"
    ])

    # =====================================================
    # LOGIN
    # =====================================================

    with tab1:

        email = st.text_input(
            "Business Email",
            key="business_login_email"
        )

        password = st.text_input(
            "Password",
            type="password",
            key="business_login_password"
        )

        if st.button("Login to Business Portal"):

            user = authenticate_business_user(
                email,
                password
            )

            if user:

                # Main session
                st.session_state.user_data = user

                # Business session
                st.session_state.business_user = user
                st.session_state.business_logged_in = True

                # Optional helper fields
                st.session_state.business_email = user["email"]
                st.session_state.business_company = user["company_name"]
                st.session_state.account_type = "business"

                _persist_login_cookie(issue_session(user["email"], "business"))

                st.success("Business Login Successful")

                st.rerun()

            else:
                st.error("Invalid Credentials")

    # =====================================================
    # REGISTER
    # =====================================================

    with tab2:

        company_name = st.text_input("Company Name")

        owner_name = st.text_input("Owner Name")

        email = st.text_input("Business Email")

        password = st.text_input(
            "Create Password",
            type="password"
        )

        selected_plan = st.selectbox(
            "Corporate Plan",
            list(BUSINESS_PLANS.keys())
        )

        plan = BUSINESS_PLANS[selected_plan]

        st.success(
            f"""
            Credits: {plan['credits']}

            Duration: {plan['duration']}

            Price: ${plan['price']}
            """
        )

        st.markdown("## 📦 Corporate Packages")

        st.table([
            ["Starter", 500, "$149.99", "3 Months"],
            ["Growth", 1000, "$299", "3 Months"],
            ["Pro", 2500, "$449", "6 Months"],
            ["Plus", 5000, "$699", "6 Months"],
            ["Advanced", 7500, "$899", "1 Year"],
            ["Enterprise", 10000, "$999", "1 Year"],
        ])

        if st.button("Create Business Account"):

            existing = get_business_user(email)

            if existing:
                st.error("Business email already exists")
                return

            password_hash = generate_password_hash(password)

            create_business_user(
                company_name=company_name,
                owner_name=owner_name,
                email=email,
                password_hash=password_hash,
                plan_name=selected_plan
            )

            st.success("Business Account Created")



def _render_steps_1_and_2(email: str, resume_text: str, active_file) -> None:
    # Resume Upload
    st.markdown("### 📄 Step 1: Upload Your Resume")
    uploaded_file = st.file_uploader(
        "Choose your resume file",
        type=["pdf", "docx"],
        help="Upload your existing resume in PDF or DOCX format",
        key="resume_uploader_job_match"
    )

    if uploaded_file:
        stored = st.session_state.get("uploaded_resume")
        if stored is None or stored.name != uploaded_file.name or stored.size != uploaded_file.size:
            st.session_state.uploaded_resume = uploaded_file
            st.session_state.step1_ats_analysis = None
            # Trigger rerun to update the active_file detection in the caller
            st.rerun()

    if active_file:
        col_ats, col_ai = st.columns([1, 1])

        with col_ats:
            check_ats_btn = st.button(
                "📊 Check ATS Score (1 credit)",
                use_container_width=True
            )

        with col_ai:
            if st.session_state.get("account_type") == "business":
                subscription = True
            else:
                subscription = check_subscription(email)

            ai_job_btn = st.button(
                "💡 AI Job Recommendations (1 credit)",
                disabled=not bool(subscription),
                help="Available for paid users / business plans only",
                use_container_width=True
            )

        if check_ats_btn:
            if not resume_text.strip():
                st.warning("⚠️ Could not extract text from your resume.")
            else:
                try:
                    ats_result = ats_charge_or_free(
                        _credit_account_type(), email, resume_text, "", feature="ATS"
                    )
                    if not ats_result.get("ok"):
                        st.error("⚠️ You need at least 1 credit to run ATS Check.")
                    else:
                        with st.spinner("Analyzing baseline ATS score..."):
                            analysis = analyze_cv_ats_score(resume_text, "")
                        st.session_state["step1_ats_analysis"] = analysis
                        st.session_state["last_ats_charged"] = ats_result.get("charged", 0)
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ Error analyzing ATS score: {str(e)}")

        if st.session_state.get("step1_ats_analysis"):
            analysis = st.session_state["step1_ats_analysis"]
            with st.container(border=True):
                col1, col2 = st.columns(2)
                with col1:
                    score_val = analysis['score'] if analysis['score'] is not None else 0
                    st.metric("General ATS Compatibility", f"{score_val}%")
                    st.progress(min(1.0, max(0.0, float(score_val) / 100.0)))
                    if score_val < 40:
                        st.warning("⚠️ Your baseline resume ATS formatting is critically low.")
                with col2:
                    kw_val = analysis['keyword_match'] if analysis['keyword_match'] is not None else 0
                    st.metric("Keyword Structure Match", f"{kw_val}%")
                    st.progress(min(1.0, max(0.0, float(kw_val) / 100.0)))

                if analysis.get('suggestions'):
                    st.markdown("### 💡 Improvement Suggestions")
                    for suggestion in analysis['suggestions']:
                        st.markdown(f"• {suggestion}")

        if ai_job_btn:
            if not check_user_access(required_credits=1, feature="Job Match"):
                st.error("⚠️ You need at least 1 credit to generate AI Recommendations.")
            elif not resume_text.strip():
                st.warning("⚠️ Could not extract text from your resume.")
            else:
                try:
                    with st.spinner("🔍 Analyzing your resume for top matching job roles..."):
                        jobs = recommend_jobs_from_resume_ai(
                            resume_text,
                            language=st.session_state.get("selected_language", "English")
                        )
                    if jobs:
                        deduct_user_credits(email, 1, feature="Job Match")
                        st.session_state["recom_jobs_list"] = jobs
                    else:
                        st.warning("⚠️ No job recommendations could be generated. Please check resume content.")
                except Exception as e:
                    st.error(f"❌ Error generating job recommendations: {str(e)}")

        recom_jobs = st.session_state.get("recom_jobs_list")
        if recom_jobs:
            st.markdown("### 🎯 Recommended Roles (Click to Search Live Jobs Below)")
            st.caption("Clicking any recommendation will automatically fill the search bar and find matching live job openings.")
            cols = st.columns(min(len(recom_jobs), 3))
            for idx, job in enumerate(recom_jobs):
                with cols[idx % len(cols)]:
                    if st.button(f"🔍 {job}", key=f"recom_btn_{idx}", use_container_width=True):
                        st.session_state["ja_title_prefill"] = job
                        st.session_state["auto_trigger_search"] = True
                        st.rerun()

        st.success("✅ Resume loaded! Use the search below or AI recommendations to select a job and run 1-Click Optimization.")
    else:
        st.info("💡 Upload your resume above to unlock instant match scores and 1-Click optimization for any job opening.")

    st.markdown("---")
    st.markdown("### 🔎 Step 2: Search Live Roles & 1-Click Optimize")

    if "ja_title_prefill" in st.session_state and st.session_state["ja_title_prefill"]:
        if st.session_state.get("auto_trigger_search"):
            st.session_state["ja_title"] = st.session_state["ja_title_prefill"]

    with st.form("job_search_form_unified"):
        c1, c2 = st.columns(2)
        with c1:
            title = st.text_input("Job Title *", key="ja_title", placeholder="e.g. Python Developer")
            location = st.text_input("Location", key="ja_location", placeholder="e.g. London / Remote")
        with c2:
            yoe = st.number_input("Years of Experience", min_value=0, max_value=50, value=0, step=1, key="ja_yoe")
            supported_countries = getattr(ja, "SUPPORTED_COUNTRIES", {
                "all": "🌍 All Countries (Global)",
                "us": "🇺🇸 United States",
                "gb": "🇬🇧 United Kingdom",
                "in": "🇮🇳 India",
                "ca": "🇨🇦 Canada",
                "au": "🇦🇺 Australia",
                "de": "🇩🇪 Germany",
                "fr": "🇫🇷 France",
                "nl": "🇳🇱 Netherlands",
                "es": "🇪🇸 Spain",
                "it": "🇮🇹 Italy",
                "br": "🇧🇷 Brazil",
                "mx": "🇲🇽 Mexico",
                "pl": "🇵🇱 Poland",
                "za": "🇿🇦 South Africa",
                "sg": "🇸🇬 Singapore",
                "nz": "🇳🇿 New Zealand",
            })
            country = st.selectbox(
                "Country / Region",
                options=list(supported_countries.keys()),
                format_func=lambda x: str(supported_countries.get(x, (x or "").upper())),
                index=0,
                key="ja_country"
            )
        work_types = st.multiselect(
            "Work Type (leave empty for all)",
            options=[ja.REMOTE_WORLDWIDE, ja.REMOTE_IN_COUNTRY, ja.ONSITE_HYBRID, ja.CONTRACT],
            key="ja_worktypes",
        )
        submitted = st.form_submit_button("🔎 Search Jobs (1 credit)")

    auto_trig = st.session_state.pop("auto_trigger_search", False)
    if submitted or auto_trig:
        if not title.strip():
            st.warning("⚠️ Please enter a job title to search.")
        else:
            if not check_user_access(required_credits=1, feature="Job Search"):
                st.warning("⚠️ You need at least 1 credit to search. Top up in the 💳 Billing tab.")
            else:
                query = ja.SearchQuery(
                    title=title.strip(),
                    years_experience=int(yoe) if yoe else None,
                    location=location.strip(),
                    work_types=work_types,
                    country=country,
                )
                with st.spinner("Searching live job sources…"):
                    try:
                        result = ja.search_jobs(query, resume_text=resume_text or None)
                    except Exception as e:
                        result = None
                        st.error(f"❌ Job search failed: {e}")
                if result is not None:
                    reachable = any(s in ("ok", "cached", "stale") for s in result["status"].values())
                    if reachable:
                        if not deduct_user_credits(email, 1, feature="Job Search"):
                            result = None
                    st.session_state["ja_result"] = result

    _render_job_results_unified(st.session_state.get("ja_result"), resume_text)


def _show_manual_jd_mode(email: str):
    """Render the Manual JD Optimizer mode."""
    active_file = st.session_state.get("uploaded_resume")
    resume_text = ""
    if active_file is not None:
        cache_key = f"resume_text_cache_{active_file.name}_{active_file.size}"
        if cache_key in st.session_state:
            resume_text = st.session_state[cache_key]
        else:
            try:
                resume_text = extract_resume_text(active_file) or ""
                st.session_state[cache_key] = resume_text
            except Exception:
                resume_text = ""

    in_qa_mode = st.session_state.alignment_stage == "questions" and st.session_state.alignment_gaps
    in_suite_mode = bool(st.session_state.get("cv_preview"))
    in_generating_mode = st.session_state.alignment_stage == "generating"
    hide_inputs = in_qa_mode or in_suite_mode or in_generating_mode

    if not hide_inputs:
        st.markdown("### 📋 Step 1: Input Job Details & Job Description")
        
        col_t, col_c = st.columns(2)
        with col_t:
            st.text_input(
                "Target Job Title *",
                value=st.session_state.get("manual_title", ""),
                key="manual_title",
                placeholder="e.g. Python Developer"
            )
        with col_c:
            st.text_input(
                "Company Name (Optional)",
                value=st.session_state.get("manual_company", ""),
                key="manual_company",
                placeholder="e.g. Acme Corp"
            )

        st.text_area(
            "Paste Job Description here *",
            value=st.session_state.get("manual_jd", ""),
            key="manual_jd",
            height=200,
            placeholder="Paste the full job description text..."
        )

        col_clear, _ = st.columns([1, 4])
        with col_clear:
            if st.button("🧹 Clear Fields", key="clear_manual_fields_btn", use_container_width=True):
                st.session_state.manual_title = ""
                st.session_state.manual_company = ""
                st.session_state.manual_jd = ""
                st.session_state.manual_ats_analysis = None
                st.rerun()

        st.markdown("---")
        st.markdown("### 📁 Step 2: Upload CV & Check ATS Score")
        
        uploaded_manual_file = st.file_uploader(
            "Choose your CV / Resume file (PDF or DOCX)",
            type=["pdf", "docx"],
            help="Upload your existing resume in PDF or DOCX format",
            key="resume_uploader_manual_jd"
        )
        if uploaded_manual_file:
            stored = st.session_state.get("uploaded_resume")
            if stored is None or stored.name != uploaded_manual_file.name or stored.size != uploaded_manual_file.size:
                st.session_state.uploaded_resume = uploaded_manual_file
                st.rerun()

        # Re-evaluate active_file and cached resume_text
        active_file = st.session_state.get("uploaded_resume")
        resume_text = ""
        if active_file is not None:
            cache_key = f"resume_text_cache_{active_file.name}_{active_file.size}"
            if cache_key in st.session_state:
                resume_text = st.session_state[cache_key]
            else:
                try:
                    resume_text = extract_resume_text(active_file) or ""
                    st.session_state[cache_key] = resume_text
                except Exception:
                    resume_text = ""

        if active_file is None:
            st.info("💡 Upload your CV/Resume above to unlock ATS check and 1-Click optimization.")
        else:
            st.success(f"✅ Loaded CV: **{active_file.name}**")
            
            col_ats, col_opt = st.columns(2)
            with col_ats:
                if st.button("📊 Check ATS Score (1 credit)", key="manual_ats_score_btn", use_container_width=True):
                    if not st.session_state.manual_jd.strip() or not st.session_state.manual_title.strip():
                        st.error("❌ Please fill in the Target Job Title and Paste Job Description first.")
                    else:
                        try:
                            ats_result = ats_charge_or_free(
                                _credit_account_type(),
                                email,
                                resume_text,
                                st.session_state.manual_jd,
                                feature="ATS",
                            )
                            if not ats_result.get("ok"):
                                st.error("⚠️ You need at least 1 credit to run ATS Check.")
                            else:
                                with st.spinner("Analyzing ATS score..."):
                                    analysis = analyze_cv_ats_score(resume_text, st.session_state.manual_jd)
                                st.session_state.manual_ats_analysis = analysis
                                st.session_state["last_ats_charged"] = ats_result.get("charged", 0)
                                st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error analyzing ATS score: {e}")
            
            with col_opt:
                if st.button("⚡ Optimize CV (3 credits)", key="manual_optimize_btn", type="primary", use_container_width=True):
                    if not st.session_state.manual_jd.strip() or not st.session_state.manual_title.strip():
                        st.error("❌ Please fill in the Target Job Title and Paste Job Description first.")
                    elif not check_user_access(required_credits=3, feature="CV"):
                        st.error("⚠️ You need at least 3 credits to optimize your CV.")
                    else:
                        with st.spinner("Analyzing CV & Job Description gaps..."):
                            try:
                                jd_h = hash_jd(st.session_state.manual_jd)
                                st.session_state.alignment_jd_hash = jd_h
                                gaps_result = analyze_cv_jd_gaps(resume_text, st.session_state.manual_jd)
                                st.session_state.alignment_gaps = gaps_result
                                st.session_state.alignment_stage = "questions"
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Gap analysis failed: {e}")

            if st.session_state.get("manual_ats_analysis"):
                analysis = st.session_state.manual_ats_analysis
                with st.container(border=True):
                    col1, col2 = st.columns(2)
                    with col1:
                        score_val = analysis['score'] if analysis['score'] is not None else 0
                        st.metric("General ATS Compatibility", f"{score_val}%")
                        st.progress(min(1.0, max(0.0, float(score_val) / 100.0)))
                    with col2:
                        kw_val = analysis['keyword_match'] if analysis['keyword_match'] is not None else 0
                        st.metric("Keyword Structure Match", f"{kw_val}%")
                        st.progress(min(1.0, max(0.0, float(kw_val) / 100.0)))
                    if analysis.get('suggestions'):
                        st.markdown("#### 💡 Improvement Suggestions")
                        for sug in analysis['suggestions']:
                            st.markdown(f"• {sug}")

    if st.session_state.alignment_stage == "questions" and st.session_state.alignment_gaps:
        gap_result = st.session_state.alignment_gaps
        gaps = gap_result.get("gaps", [])
        overall = gap_result.get("overall_match")

        st.markdown("---")
        if st.button("↩ Back to JD Optimizer", key="back_to_manual_jd"):
            st.session_state.alignment_stage = "idle"
            st.session_state.alignment_gaps = None
            st.rerun()

        head_l, head_r = st.columns([3, 1])
        with head_l:
            target_disp = f" **{st.session_state.manual_title}**"
            st.markdown(f"### ⚡ Step 3: Boost Your Match Before Generating{target_disp}")
            st.caption(
                "We identified a few requirements where your resume shows no or weak evidence. "
                "Answer any that apply using your real experience — we'll weave your verified answers straight into the generated resume!"
            )
        with head_r:
            if overall is not None:
                st.metric("Current JD Match", f"{overall}%")

        answers = {}
        for gap in gaps:
            with st.container(border=True):
                st.markdown(f"**{gap.get('area', '')}**")
                if gap.get("why"):
                    st.caption(f"Why this matters: {gap['why']}")
                val = st.text_area(
                    gap.get("question", "Tell us more:"),
                    key=f"gap_manual_{gap.get('id', gap.get('area', ''))}"
                )
                if val.strip():
                    answers[gap.get("area", "")] = val.strip()

        st.markdown("---")
        st.markdown("#### ⚖️ Accuracy Acknowledgment & Disclaimer")
        st.warning(
            "Please provide accurate information based on your real experience. "
            "CVOLVE PRO can help structure and improve your CV, but you are responsible for the accuracy "
            "of the information you provide. If you enter false or misleading details, you accept full "
            "responsibility for the final content."
        )

        ack = st.checkbox(
            "I confirm that all information provided is accurate and reflects my real experience.",
            key="manual_alignment_ack"
        )

        btn_l, btn_r = st.columns(2)
        with btn_l:
            if st.button("🚀 Save & Generate Resume (3 credits)", key="manual_save_gen_btn", type="primary", disabled=not ack):
                try:
                    jd_h = st.session_state.alignment_jd_hash
                    save_alignment_answers(email, jd_h, gaps, answers)
                    st.session_state.alignment_stage = "generating"
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Saving answers failed: {e}")
        with btn_r:
            if st.button("⏭ Skip & Generate Resume (3 credits)", key="manual_skip_gen_btn", disabled=not ack):
                st.session_state.alignment_stage = "generating"
                st.rerun()

    if st.session_state.alignment_stage == "generating":
        loading_placeholder = st.empty()
        with loading_placeholder.container():
            with st.spinner("⚡ Tailoring your resume to maximum ATS compatibility (target match 100%)..."):
                try:
                    st.session_state.job_description = st.session_state.manual_jd
                    jd_h = st.session_state.alignment_jd_hash
                    extra_context = get_alignment_answers(email, jd_h).get("answers", {})
                    cv_result = generate_cv(
                        resume_text=resume_text,
                        job_description=st.session_state.manual_jd,
                        target_match=100,
                        language=st.session_state.get("selected_language", "English"),
                        extra_context=extra_context,
                        optimization_depth="max_ats",
                        return_metadata=True,
                    )
                    cv_content = cv_result.get("optimized_content", "") if isinstance(cv_result, dict) else str(cv_result)
                    clean_preview = cv_content.replace("**", "")
                    pdf_buf = apply_template(clean_preview, st.session_state.get("selected_template", "professional"))
                    docx_buf = create_word_document(cv_content)
                    
                    st.session_state.cv_preview = cv_content
                    st.session_state.cv_optimization_metadata = cv_result if isinstance(cv_result, dict) else None
                    st.session_state.cv_pdf_bytes = pdf_buf.getvalue()
                    st.session_state.cv_docx_bytes = docx_buf.getvalue()
                    
                    deduct_user_credits(email, 3, feature="CV")
                    st.session_state.alignment_stage = "idle"
                    st.session_state.alignment_gaps = None
                    st.rerun()
                except Exception as e:
                    st.session_state.alignment_stage = "idle"
                    st.session_state.alignment_gaps = None
                    st.error(f"❌ Resume optimization failed: {e}")

    if st.session_state.get("cv_preview"):
        title_disp = st.session_state.manual_title
        comp_disp = f" at {st.session_state.manual_company}" if st.session_state.manual_company else ""
        _render_application_suite(
            email=email,
            resume_text=resume_text,
            jd_to_use=st.session_state.manual_jd,
            title_disp=title_disp,
            comp_disp=comp_disp
        )


def _render_interview_full_page():
    """Full-page Interview Practice mode, entered from the CV optimization flow."""
    st.markdown("---")
    col_back, col_title = st.columns([1, 4])
    with col_back:
        if st.button("↩ Back to My Application", key="interview_back_to_suite", use_container_width=True):
            st.session_state.page = "home"
            st.rerun()
    with col_title:
        st.caption("Continuing from your optimized CV — practice interviews tailored to this job.")

    show_interview_practice_page(
        check_access_fn=check_user_access,
        deduct_credits_fn=deduct_user_credits,
        extract_resume_fn=extract_resume_text,
        export_qa_fn=export_interview_qa,
    )


def _render_application_suite(email: str, resume_text: str, jd_to_use: str, title_disp: str, comp_disp: str):
    """Render Step 4: Your Tailored Application Suite - shared by both modes."""
    st.markdown("---")
    st.markdown(f"### 🎯 Step 4: Your Tailored Application Suite for **{title_disp}**{comp_disp}")

    # Add navigation and apply actions at the top of the suite
    col_back, col_apply = st.columns([1, 1])
    with col_back:
        back_label = "↩ Back to Job Search" if st.session_state.optimizer_mode == "smart" else "↩ Back to JD Optimizer"
        if st.button(back_label, key="back_to_search_suite", use_container_width=True):
            st.session_state.cv_preview = None
            st.session_state.pop("cv_pdf_bytes", None)
            st.session_state.pop("cv_docx_bytes", None)
            st.session_state.alignment_stage = "idle"
            st.session_state.alignment_gaps = None
            st.session_state.cover_letter_content = None
            st.session_state.interview_qa_content = None
            st.session_state.manual_ats_analysis = None
            st.rerun()
    with col_apply:
        active_url = st.session_state.get("active_job_url")
        if active_url:
            domain = _apply_domain_label(active_url)
            btn_label = f"🚀 View / Apply on {domain} ↗" if domain else "🚀 View / Apply on Source ↗"
            st.link_button(btn_label, active_url, use_container_width=True)
        else:
            st.button("🚀 View / Apply on Source", disabled=True, use_container_width=True)

    st.markdown("---")

    # Continue into Interview Practice with this job's context
    st.markdown("#### 🎤 Next Step: Practice Interview for This Job")
    st.caption("Continue from your optimized CV to practice AI-led interview questions tailored to this job description.")
    if st.button("🎤 Practice Interview for This Job", type="primary",
                 key="continue_to_interview_btn", use_container_width=True):
        incoming_jd = (jd_to_use or st.session_state.get("job_description", "") or "").strip().lower()
        same_job = (st.session_state.get("interview_jd", "") or "").strip().lower() == incoming_jd \
            and st.session_state.get("interview_qa_bank") is not None
        if not same_job:
            st.session_state.interview_qa_bank = None
            st.session_state.interview_questions_flat = None
            st.session_state.interview_session_results = []
            st.session_state.interview_report = None
            st.session_state.interview_current_idx = 0
        st.session_state.interview_jd = jd_to_use or st.session_state.get("job_description", "")
        st.session_state.interview_resume_text = st.session_state.get("cv_preview") or resume_text or ""
        st.session_state.interview_phase = "setup"
        st.session_state.interview_qa_content = None
        st.session_state.page = "interview"
        st.rerun()

    st.markdown("---")

    # Template selection
    templates = {
        "classic_serif": "📜 Executive Serif (Classic Standard)",
        "modern_sans": "⚡ Modern Tech (Product Standard)",
    }
    col_t1, col_t2 = st.columns([1, 2])
    with col_t1:
        st.selectbox(
            "Select Layout Template",
            options=list(templates.keys()),
            format_func=lambda x: templates[x],
            key="selected_template"
        )

    pdf_bytes = st.session_state.get("cv_pdf_bytes")
    docx_bytes = st.session_state.get("cv_docx_bytes")
    if not pdf_bytes or not docx_bytes:
        try:
            clean_preview = st.session_state.cv_preview.replace("**", "")
            _pdf_buf = apply_template(clean_preview, st.session_state.selected_template)
            _docx_buf = create_word_document(st.session_state.cv_preview)
            pdf_bytes = _pdf_buf.getvalue()
            docx_bytes = _docx_buf.getvalue()
            st.session_state.cv_pdf_bytes = pdf_bytes
            st.session_state.cv_docx_bytes = docx_bytes
        except Exception as e:
            st.error(f"❌ Failed to prepare downloads: {e}")
            pdf_bytes, docx_bytes = None, None

    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            label="📥 Download Resume (PDF)",
            data=pdf_bytes or b"",
            file_name=f"resume_{title_disp.replace(' ', '_').lower()}.pdf",
            mime="application/pdf",
            key="dl_cv_pdf_persist",
            disabled=(pdf_bytes is None)
        )
    with c2:
        st.download_button(
            label="📄 Download Resume (DOCX)",
            data=docx_bytes or b"",
            file_name=f"resume_{title_disp.replace(' ', '_').lower()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_cv_docx_persist",
            disabled=(docx_bytes is None)
        )
    with c3:
        if st.button("🔄 Regenerate Resume", key="regen_cv_persist"):
            st.session_state.cv_preview = None
            st.session_state.pop("cv_pdf_bytes", None)
            st.session_state.pop("cv_docx_bytes", None)
            st.rerun()

    # Show preview content and ATS analysis
    st.markdown("### 📋 Preview Content")
    st.markdown(st.session_state.cv_preview)

    st.markdown("### 📊 ATS Analysis")
    analyze_ats_compatibility()

    st.markdown("---")

    # ─── Cover Letter ───
    st.markdown("### 📝 Cover Letter")
    if not st.session_state.get("cover_letter_content"):
        if st.button("📝 Generate Cover Letter (2 credits)", key="gen_cover_letter_btn", type="primary"):
            if not check_user_access(required_credits=2, feature="Cover Letter"):
                st.error("⚠️ Insufficient credits. You need 2 credits to generate a Cover Letter.")
            else:
                with st.spinner("Generating your tailored cover letter..."):
                    try:
                        jd_h = st.session_state.alignment_jd_hash or hash_jd(jd_to_use)
                        extra = get_alignment_answers(email, jd_h).get("answers", {})
                        cl_text = generate_cover_letter(
                            resume_text=resume_text,
                            job_description=jd_to_use,
                            language=st.session_state.get("selected_language", "English"),
                            extra_context=extra
                        )
                        st.session_state.cover_letter_content = cl_text
                        deduct_user_credits(email, 2, feature="Cover Letter")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Cover letter generation failed: {e}")
    else:
        st.markdown(st.session_state.cover_letter_content)
        cl_pdf_bytes, cl_docx_bytes = None, None
        try:
            cl_pdf_buf, cl_docx_buf = export_cover_letter(st.session_state.cover_letter_content)
            cl_pdf_bytes = cl_pdf_buf.getvalue()
            cl_docx_bytes = cl_docx_buf.getvalue()
        except Exception as e:
            st.error(f"❌ Failed to prepare cover letter downloads: {e}")

        cl_col1, cl_col2, cl_col3 = st.columns(3)
        with cl_col1:
            st.download_button(
                label="📥 Download Cover Letter (PDF)",
                data=cl_pdf_bytes or b"",
                file_name=f"cover_letter_{title_disp.replace(' ', '_').lower()}.pdf",
                mime="application/pdf",
                key="dl_cl_pdf",
                disabled=(cl_pdf_bytes is None)
            )
        with cl_col2:
            st.download_button(
                label="📄 Download Cover Letter (DOCX)",
                data=cl_docx_bytes or b"",
                file_name=f"cover_letter_{title_disp.replace(' ', '_').lower()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cl_docx",
                disabled=(cl_docx_bytes is None)
            )
        with cl_col3:
            if st.button("🧹 Clear Cover Letter", key="clear_cl_btn"):
                st.session_state.cover_letter_content = None
                st.rerun()

    st.info("🔍 Use the buttons above to download your tailored suite items.")


def show_smart_job_match_page():
    """Main CV generation interface"""
    col_header, col_dropdown = st.columns([3, 1])
    with col_header:
        st.markdown("## 🎯 Match Me to the Job")
    with col_dropdown:
        # Limit choices by subscription
        allowed_options = get_allowed_ai_models_for_user()

        options = st.session_state.get("ai_model_options", ["Premium", "Premium Classic"])
        default_label = "Premium Classic" if st.session_state.get("ai_model") == "openai" else "Premium"
        if default_label not in options:
            default_label = options[0]

        model_choice = st.selectbox(
            "AI Model",
            options=options,
            index=options.index(default_label)
        )
        st.session_state["ai_model"] = "openai" if model_choice == "Premium Classic" else "gemini"

        # Safety: if plan is Premium-only but session had Classic, force-correct it
        if "Premium Classic" not in allowed_options and st.session_state.get("ai_model") == "openai":
            st.session_state["ai_model"] = "gemini"

        # --- NEW: Language selector shown beside model selector ---
        # Initialize default language in session state if not present
        if "selected_language" not in st.session_state:
            st.session_state["selected_language"] = "English"

        # Language options — add more languages as you like
        language_options = ["English", "Français", "Español", "Deutsch"]

        # Show a compact dropdown for language choice
        st.selectbox(
            "Language",
            options=language_options,
            index=language_options.index(st.session_state["selected_language"]) if st.session_state["selected_language"] in language_options else 0,
            key="selected_language"
        )

    mode_col1, mode_col2 = st.columns(2)
    with mode_col1:
        if st.button("🔎 Smart Job Match", 
                     type="primary" if st.session_state.optimizer_mode == "smart" else "secondary",
                     use_container_width=True):
            st.session_state.optimizer_mode = "smart"
            st.rerun()
    with mode_col2:
        if st.button("📋 Manual JD Optimizer",
                     type="primary" if st.session_state.optimizer_mode == "manual" else "secondary",
                     use_container_width=True):
            st.session_state.optimizer_mode = "manual"
            st.rerun()

    email = st.session_state.user_data['email']

    if st.session_state.optimizer_mode == "manual":
        _show_manual_jd_mode(email)
        return
    active_file = st.session_state.get("uploaded_resume")
    resume_text = ""
    if active_file is not None:
        cache_key = f"resume_text_cache_{active_file.name}_{active_file.size}"
        if cache_key in st.session_state:
            resume_text = st.session_state[cache_key]
        else:
            try:
                resume_text = extract_resume_text(active_file) or ""
                st.session_state[cache_key] = resume_text
            except Exception:
                resume_text = ""

    # Conditionally render Steps 1 & 2
    in_qa_mode = st.session_state.alignment_stage == "questions" and st.session_state.alignment_gaps
    in_suite_mode = bool(st.session_state.get("cv_preview"))
    in_generating_mode = st.session_state.alignment_stage == "generating"
    hide_steps_1_2 = in_qa_mode or in_suite_mode or in_generating_mode

    if not hide_steps_1_2:
        _render_steps_1_and_2(email, resume_text, active_file)

    # =========================================================
    # Step 3: Gap Analysis & Q&A
    # =========================================================
    if st.session_state.alignment_stage == "questions" and st.session_state.alignment_gaps:
        gap_result = st.session_state.alignment_gaps
        gaps = gap_result.get("gaps", [])
        overall = gap_result.get("overall_match")

        st.markdown("---")
        if st.button("↩ Back to Job Search", key="back_to_search_qa"):
            st.session_state.alignment_stage = "idle"
            st.session_state.alignment_gaps = None
            st.rerun()

        head_l, head_r = st.columns([3, 1])
        with head_l:
            target_disp = f" **{st.session_state.target_job_title}**" if st.session_state.get("target_job_title") else ""
            st.markdown(f"### ⚡ Step 3: Boost Your Match Before Generating{target_disp}")
            st.caption(
                "We identified a few requirements where your resume shows no or weak evidence. "
                "Answer any that apply using your real experience — we'll weave your verified answers straight into the generated resume!"
            )
        with head_r:
            if overall is not None:
                st.metric("Current JD Match", f"{overall}%")

        for gap in gaps:
            with st.container(border=True):
                st.markdown(f"**{gap.get('area', '')}**")
                if gap.get("why"):
                    st.caption(f"Why this matters: {gap['why']}")
                st.text_area(
                    gap.get("question", "Tell us more:"),
                    key=f"gap_answer_{gap['id']}",
                    placeholder="Type your verified experience here…",
                    height=80,
                )
                example = gap.get("example", "")
                if example:
                    st.caption(f"💡 Example answer: {example}")

        st.info(
            "Please provide accurate information based on your real experience. CVOLVE PRO can "
            "help structure and improve your CV, but you are responsible for the accuracy of the "
            "information you provide. If you enter false or misleading details, you accept full "
            "responsibility for the final content."
        )
        st.checkbox(
            "I confirm the information I provide is accurate and based on my real experience.",
            key="alignment_ack",
        )

        col_use, col_skip = st.columns(2)
        with col_use:
            use_answers_btn = st.button("✅ Save Answers & Generate Resume", type="primary")
        with col_skip:
            skip_btn = st.button("⏭ Skip & Generate Resume")

        if use_answers_btn or skip_btn:
            if not st.session_state.get("alignment_ack"):
                st.warning("⚠️ Please confirm that the information provided is accurate by checking the confirmation box.")
            else:
                jd_h = st.session_state.alignment_jd_hash
                answers = {}
                for gap in gaps:
                    ans = st.session_state.get(f"gap_answer_{gap['id']}", "").strip()
                    if ans:
                        answers[gap.get("area", gap["id"])] = ans

                try:
                    save_alignment_answers(email, jd_h, gaps, answers)
                except Exception:
                    pass

                st.session_state.alignment_stage = "generating"
                st.rerun()

    # =========================================================
    # Step 3: Automated CV Generation
    # =========================================================
    jd_to_use = st.session_state.get("job_description", "")
    if st.session_state.alignment_stage == "generating" and active_file and jd_to_use.strip():
        st.session_state.alignment_stage = "idle"

        loading_placeholder = st.empty()
        loading_placeholder.markdown("""
            <div class="cvolve-loading">
                <div class="custom-loader"></div>
                <p class="cvolve-loading-text">Optimizing your resume with maximum ATS score (100 target) — usually takes a few seconds…</p>
            </div>
        """, unsafe_allow_html=True)

        start_time = time.time()

        try:
            resume_text = extract_resume_text(active_file)
            jd_h = st.session_state.alignment_jd_hash or hash_jd(jd_to_use)
            extra_context = get_alignment_answers(email, jd_h).get("answers", {})

            sections_to_use = st.session_state.auto_save.get('sections', {
                "Professional Summary": True,
                "Key Skills": True,
                "Work Experience": True,
                "Education": True,
                "Certifications": True,
                "Projects": True,
                "Awards": False,
                "Languages": False,
                "Hobbies": False
            })

            target_match_val = 100
            st.session_state["target_match"] = target_match_val

            cv_result = generate_cv(
                resume_text=resume_text,
                job_description=jd_to_use,
                target_match=target_match_val,
                template=st.session_state.selected_template,
                sections=sections_to_use,
                quantitative_focus=60,
                action_verb_intensity="High",
                keyword_matching="Balanced",
                language=st.session_state.get("selected_language", "English"),
                extra_context=extra_context,
                optimization_depth="max_ats",
                return_metadata=True,
            )
            cv_content = cv_result.get("optimized_content", "") if isinstance(cv_result, dict) else str(cv_result)
            st.session_state.cv_optimization_metadata = cv_result if isinstance(cv_result, dict) else None

            # Pre-display quality gate: validate CV quality
            try:
                from cv_generator import validate_cv_quality
                quality_check = validate_cv_quality(cv_content, jd_to_use)
                
                if not quality_check["is_valid"] and quality_check["should_regenerate"]:
                    st.warning("⚠️ CV quality checks still found issues after Max ATS repair. Showing measured gaps below.")
                elif quality_check["issues"]:
                    # Show warnings but don't block
                    for issue in quality_check["issues"]:
                        st.info(f"ℹ️ {issue}")
            except Exception as e:
                # Don't fail the whole process if quality gate has issues
                pass

            cv_content = enforce_page_limit(cv_content)
            st.session_state.cv_preview = cv_content

            clean_preview = st.session_state.cv_preview.replace("**", "")
            pdf_buffer = apply_template(clean_preview, st.session_state.selected_template)
            docx_buffer = create_word_document(st.session_state.cv_preview)
            st.session_state.cv_pdf_bytes = pdf_buffer.getvalue()
            st.session_state.cv_docx_bytes = docx_buffer.getvalue()

            loading_placeholder.empty()
            processing_time = time.time() - start_time
            st.success(f"✅ Resume optimized successfully in {processing_time:.1f} seconds!")

            try:
                jd_clean = _sanitize_db_text(jd_to_use)
                resume_clean = _sanitize_db_text(resume_text)
                cv_clean = _sanitize_db_text(st.session_state.cv_preview)

                metadata = st.session_state.get("cv_optimization_metadata") or {}
                ats_score_val = metadata.get("ats_score")
                if ats_score_val is None:
                    from utils import optimize_keywords
                    quick_analysis = optimize_keywords(cv_clean, jd_clean)
                    ats_score_val = int(float(quick_analysis.get("score") or target_match_val))
                else:
                    ats_score_val = int(ats_score_val)

                if st.session_state.get("account_type") != "business":
                    save_cv_generation(
                        user_email=email,
                        job_description=jd_clean,
                        original_resume=resume_clean,
                        generated_cv=cv_clean,
                        template_used=st.session_state.selected_template,
                        ats_score=ats_score_val,
                        target_match=target_match_val,
                        processing_time=float(f"{processing_time:.2f}")
                    )
            except Exception as e:
                st.error("❌ Failed to record this CV in Analytics.")
                st.exception(e)

            deduct_user_credits(email, 3, feature="CV")

        except Exception as e:
            loading_placeholder.empty()
            st.error(f"❌ Error generating resume: {str(e)}")

    # =========================================================
    # Step 4: Job-Specific Suite & Downloads
    # =========================================================
    if st.session_state.get("cv_preview"):
        title_disp = st.session_state.get("target_job_title") or "Selected Role"
        comp_disp = f" at {st.session_state.get('target_job_company')}" if st.session_state.get("target_job_company") else ""
        _render_application_suite(
            email=email,
            resume_text=resume_text,
            jd_to_use=jd_to_use,
            title_disp=title_disp,
            comp_disp=comp_disp
        )



def show_preview_page():
    """CV preview and download page"""
    st.markdown("## 📄 CV Preview")
    
    if st.session_state.cv_preview:
        st.markdown("### 👀 Your Optimized CV")
        
        # Preview options
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📥 Download as PDF"):
                pdf_buffer = apply_template(
                    st.session_state.cv_preview,
                    st.session_state.selected_template
                )
                
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_buffer,
                    file_name="optimized_cv.pdf",
                    mime="application/pdf"
                )
        
        with col2:
            if st.button("📄 Download as Word"):
                docx_buffer = create_word_document(st.session_state.cv_preview)
                
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_buffer,
                    file_name="optimized_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col3:
            if st.button("🔄 Regenerate"):
                st.session_state.cv_preview = None
                st.rerun()
        
        # Show preview
        st.markdown("---")
        st.markdown("### 📋 Preview Content")
        st.markdown(st.session_state.cv_preview)
        
        # ATS Analysis - Show automatically
        st.markdown("### 📊 ATS Analysis")
        analyze_ats_compatibility()
    
    else:
        st.info("🔍 No CV preview available. Please generate a CV first.")

def show_analytics_page():
    """Analytics dashboard (live from DB)"""
    from database import get_db_connection, get_user_credits, release_db_connection
    import plotly.graph_objects as go

    st.markdown("## 📊 Your Analytics")

    user = st.session_state.get("user_data")
    if not user:
        st.info("Please log in to see analytics.")
        return
    user_email = user["email"]

    # ---------- Aggregates from DB ----------
    total_cvs = 0
    avg_ats = 0.0
    success_rate = 0.0
    if st.session_state.get("account_type") == "business":
        credits_now = get_business_credits(user_email) or 0
        try:
            credits_now = wallet_balance("business", user_email)["total"]
        except Exception:
            pass
    else:
        credits_now = get_user_credits(user_email) or 0
        try:
            credits_now = wallet_balance("individual", user_email)["total"]
        except Exception:
            pass

    trend_dates, trend_scores = [], []
    conn, cur = None, None

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Totals
        cur.execute("SELECT COUNT(*), COALESCE(AVG(ats_score),0) FROM cv_generations WHERE user_email=%s", (user_email,))
        row = cur.fetchone() or (0, 0.0)
        total_cvs = int(row[0] or 0)
        avg_ats = float(row[1] or 0.0)

        # Success rate
        cur.execute("""
            SELECT 
                COALESCE(SUM(CASE WHEN ats_score IS NOT NULL 
                                   AND target_match IS NOT NULL 
                                   AND ats_score >= target_match THEN 1 ELSE 0 END), 0) AS successes,
                COUNT(*) AS total_rows
            FROM cv_generations
            WHERE user_email=%s
        """, (user_email,))
        srow = cur.fetchone() or (0, 0)
        successes, total_rows = int(srow[0] or 0), int(srow[1] or 0)
        success_rate = (successes / total_rows * 100.0) if total_rows > 0 else 0.0

        # Trend
        cur.execute("""
            SELECT DATE(created_at) AS d, COALESCE(AVG(ats_score), 0)
            FROM cv_generations
            WHERE user_email=%s
            GROUP BY 1
            ORDER BY 1
        """, (user_email,))
        trend = cur.fetchall() or []
        trend_dates = [str(r[0]) for r in trend]
        trend_scores = [float(r[1] or 0.0) for r in trend]

        # Credits used TOTAL via credit_usage
        cur.execute("""
            SELECT COALESCE(SUM(credits), 0) 
              FROM credit_usage 
             WHERE user_email=%s
        """, (user_email,))
        credits_used_total = int((cur.fetchone() or [0])[0] or 0)

        # Per-feature usage (includes ATS)
        cur.execute("""
            SELECT feature, COALESCE(SUM(credits), 0)
              FROM credit_usage
             WHERE user_email=%s
             GROUP BY feature
        """, (user_email,))
        rows = cur.fetchall() or []
        usage_map = {r[0]: int(r[1] or 0) for r in rows}
        used_cv = usage_map.get("CV", 0)
        used_cl = usage_map.get("CL", 0)
        used_qa = usage_map.get("Interview QA", 0)
        used_ats = usage_map.get("ATS", 0)

        # Recent activity (same connection)
        cur.execute("""
            SELECT created_at, ats_score, target_match, template_used
              FROM cv_generations
             WHERE user_email=%s
             ORDER BY created_at DESC
             LIMIT 10
        """, (user_email,))
        recent_rows = cur.fetchall() or []

        if cur: cur.close()
        if conn: release_db_connection(conn)
    except Exception as e:
        try:
            if cur: cur.close()
            if conn: release_db_connection(conn)
        except:
            pass
        st.error(f"Failed to load analytics: {e}")
        return

    # ---------- KPI Tiles ----------
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("CVs Generated", f"{total_cvs}")
    with col2:
        st.metric("Avg ATS Score", f"{avg_ats:.0f}%")
    with col3:
        st.metric("Total Credits Used", f"{credits_used_total}")   # <-- from credit_usage
    with col4:
        st.metric("Success Rate", f"{success_rate:.0f}%")

    st.markdown("---")

    # ---------- Credits by Feature ----------
    st.markdown("### Credits Used by Feature")
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("CV", used_cv)
    with c2: st.metric("Cover Letter", used_cl)
    with c3: st.metric("Interview Q&A", used_qa)
    with c4: st.metric("ATS", used_ats)



    # ---------- Trends ----------
    st.markdown("### 📈 ATS Score Trend")
    if trend_dates:
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=trend_dates, y=trend_scores, mode='lines+markers', name='Avg ATS'))
        fig.update_layout(
            margin=dict(l=10, r=10, t=40, b=10),
            title="Average ATS Score by Day",
            xaxis_title="Date", yaxis_title="Avg ATS %", yaxis=dict(range=[0, 100])
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No ATS score history yet. Generate a CV to see your trend!")

    # ---------- Recent activity ----------
    st.markdown("### 🧾 Recent CV Generations")
    if recent_rows:
        st.dataframe(
            [{
                "Date": r[0].strftime("%Y-%m-%d %H:%M"),
                "ATS": int(r[1]) if r[1] is not None else None,
                "Target": int(r[2]) if r[2] is not None else None,
                "Template": r[3] or "-"
            } for r in recent_rows],
            use_container_width=True, hide_index=True
        )
    else:
        st.write("No recent generations found.")


def show_admin_vouchers_page(admin_email):
    """Admin-only voucher management UI: generate, list/revoke, view redemptions."""
    from datetime import datetime as _dt

    # ── Generate a new voucher ───────────────────────────────────────────────
    st.markdown("#### ➕ Generate voucher")
    with st.form("generate_voucher_form"):
        c1, c2, c3 = st.columns(3)
        with c1:
            duration_days = st.number_input(
                "Duration (days)", min_value=1, max_value=365,
                value=voucher_engine.DEFAULT_DURATION_DAYS, step=1,
            )
            max_redemptions = st.number_input(
                "Max redemptions", min_value=1, max_value=1000,
                value=voucher_engine.DEFAULT_MAX_REDEMPTIONS, step=1,
            )
        with c2:
            plan_choice = st.selectbox("Plan", options=[voucher_engine.DEFAULT_PLAN], index=0)
            expires_at_input = st.date_input("Expires on (optional)", value=None)
        with c3:
            note = st.text_input("Note (optional)", placeholder="e.g. Marketing — July cohort")
        submitted = st.form_submit_button("Generate", use_container_width=True)

    if submitted:
        try:
            expires_at = (
                _dt.combine(expires_at_input, _dt.min.time())
                if expires_at_input else None
            )
            v = generate_voucher(
                plan=plan_choice,
                duration_days=int(duration_days),
                max_redemptions=int(max_redemptions),
                expires_at=expires_at,
                created_by=admin_email,
                note=note or None,
            )
            st.success(
                f"✅ Created **`{v['code']}`** — {v['plan']}, "
                f"{v['duration_days']} days, {v['max_redemptions']} redemption(s)."
            )
            st.info("📋 Copy the code above and share it with the selected user.")
        except Exception as e:
            st.error(f"❌ Could not create voucher: {e}")

    # ── List + revoke vouchers ───────────────────────────────────────────────
    st.markdown("#### 📋 Vouchers")
    try:
        rows = list_all_vouchers()
    except Exception as e:
        st.error(f"Could not load vouchers: {e}")
        return

    if not rows:
        st.info("No vouchers yet.")
    else:
        for v in rows:
            with st.container(border=True):
                col_info, col_action = st.columns([5, 1])
                with col_info:
                    expires = v.get("expires_at")
                    expires_str = expires.strftime("%Y-%m-%d") if expires else "—"
                    created = v.get("created_at")
                    created_str = created.strftime("%Y-%m-%d") if created else "—"
                    st.markdown(
                        f"**`{v['code']}`** · {v['plan']} · "
                        f"**{v['redeemed_count']}/{v['max_redemptions']}** redemptions · "
                        f"expires {expires_str} · status: `{v['status']}`"
                    )
                    meta = []
                    if v.get("note"):
                        meta.append(f"📝 {v['note']}")
                    if v.get("created_by"):
                        meta.append(f"Created by {v['created_by']} on {created_str}")
                    if meta:
                        st.caption(" · ".join(meta))
                with col_action:
                    if v["status"] == "active":
                        if st.button("Revoke", key=f"revoke_{v['code']}", use_container_width=True):
                            try:
                                admin_revoke_voucher(v["code"])
                                st.success(f"Revoked {v['code']}")
                                st.rerun()
                            except Exception as e:
                                st.error(str(e))

    # ── Redemption history ───────────────────────────────────────────────────
    st.markdown("#### 👥 Redemption history")
    try:
        reds = list_voucher_redemptions(limit=200)
    except Exception as e:
        st.error(f"Could not load redemptions: {e}")
        return

    if not reds:
        st.info("No redemptions yet.")
        return

    codes = sorted({r["voucher_code"] for r in reds})
    code_filter = st.selectbox("Filter by code", options=["(all)"] + codes)
    view = reds if code_filter == "(all)" else [r for r in reds if r["voucher_code"] == code_filter]
    st.dataframe(
        [{
            "Code": r["voucher_code"],
            "Email": r["user_email"],
            "When": r["redeemed_at"].strftime("%Y-%m-%d %H:%M") if r.get("redeemed_at") else "—",
        } for r in view],
        use_container_width=True,
        hide_index=True,
    )


def show_billing_page():
    """Billing and subscription management with Stripe + simple currency selection"""

    # ---- Currency: safe, local-friendly (no JS, no early return) ----
    cur_param = st.query_params.get("cur", None)
    if not cur_param:
        # Best-effort first guess from phone prefix while local/testing
        phone = (st.session_state.get("user_data", {}).get("phone") or "").strip()
        prefix_map = {
            "+91": "INR",   # India
            "+971": "AED",  # UAE
            "+973": "BHD",  # Bahrain
            "+61": "AUD",   # Australia
            "+44": "GBP",   # UK
        }
        guess = "USD"
        for pref, code in prefix_map.items():
            if phone.startswith(pref):
                guess = code
                break
        st.query_params["cur"] = guess  # triggers a rerun

    cur_param = st.query_params.get("cur", "USD")
    if isinstance(cur_param, list):
        cur_param = cur_param[0]

    SUPPORTED = ("INR", "EUR", "USD", "AED", "BHD", "AUD", "GBP")
    CURRENT_CURRENCY = cur_param if cur_param in SUPPORTED else "USD"

    SYMBOL = {
        "USD": "$", "INR": "₹", "EUR": "€",
        "AED": "د.إ", "BHD": "BD", "AUD": "A$", "GBP": "£",
    }[CURRENT_CURRENCY]

    # Simple static USD→local multipliers (adjust anytime)
    RATE = {
        "USD": 1.00,
        "INR": 84.00,
        "EUR": 0.92,
        "AED": 3.67,
        "BHD": 0.38,
        "AUD": 1.50,
        "GBP": 0.78,
    }[CURRENT_CURRENCY]

    THREE_DECIMAL = {"BHD"}  # BHD uses 3 decimals

    def price_local(usd_amount: float) -> float:
        v = usd_amount * RATE
        if CURRENT_CURRENCY == "INR":
            return round(v)           # whole rupees
        if CURRENT_CURRENCY in THREE_DECIMAL:
            return round(v, 3)        # e.g., BHD
        return round(v, 2)            # default 2-decimal

    def fmt(amount_local: float) -> str:
        if CURRENT_CURRENCY == "INR":
            return f"{SYMBOL}{amount_local:,.0f}"
        if CURRENT_CURRENCY in THREE_DECIMAL:
            return f"{SYMBOL}{amount_local:,.3f}"
        return f"{SYMBOL}{amount_local:,.2f}"

    # ---- Make sure we have a user while testing locally ----
    if "user_data" not in st.session_state:
        st.session_state.user_data = {"email": "local@test.com", "phone": "+91-0000000000"}
    user_email = st.session_state.user_data["email"]

    # ---- Imports / setup ----
    import os, stripe, urllib.parse
    from database import get_user_credits, save_payment, get_db_connection
    from payment import create_checkout_session, check_subscription

    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")
    LEGACY_PLAN_MAP = {
        "Premium": "Career Pro",
        "Premium + Premium Classic": "Interview Pro",
    }
    PACK_CREDITS_TO_NAME = {p["credits"]: p["name"] for p in pricing.PACKS}

    def payment_exists(stripe_payment_id: str) -> bool:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM payments WHERE stripe_payment_id=%s LIMIT 1", (stripe_payment_id,))
        exists = cur.fetchone() is not None
        cur.close(); conn.close()
        return exists

    def qp_get(key: str, default=""):
        v = st.query_params.get(key, default)
        return v[0] if isinstance(v, list) else v

    # ---- Read query params ----
    success    = qp_get("success", "").lower() == "true"
    typ        = qp_get("type", "")
    session_id = qp_get("session_id", "")
    credits_qp = int(qp_get("credits", "0") or 0)
    plan_qp    = qp_get("plan", "")

    if "processed_sessions" not in st.session_state:
        st.session_state.processed_sessions = set()

    # ---- Handle CREDITS (pack) success ----
    if success and typ == "credits" and session_id and session_id not in st.session_state.processed_sessions:
        try:
            sess = stripe.checkout.Session.retrieve(session_id)
            if sess.get("payment_status") == "paid":
                md = sess.get("metadata") or {}
                amount_paid = (sess.get("amount_total") or 0) / 100.0
                pack_name = (md.get("pack") or "").strip()
                if not pack_name:
                    pack_credits = int(md.get("credits") or credits_qp or 0)
                    pack_name = {p["credits"]: p["name"] for p in pricing.PACKS}.get(pack_credits)

                if not payment_exists(session_id) and pack_name:
                    res = purchase_pack("individual", user_email, pack_name, stripe_session_id=session_id)
                    if res.get("ok"):
                        save_payment(user_email, amount_paid, "credits", session_id,
                                     credits_purchased=res["credits"])
                        st.success(f"🎉 {res['pack']} ({res['credits']} credits) added to your account.")
                    else:
                        st.error(f"Could not add pack: {res.get('reason')}")
                else:
                    st.info("Payment already processed or no pack found.")
            else:
                st.warning("Payment not completed yet.")
        except Exception as e:
            st.error(f"Could not verify credit payment: {e}")
        finally:
            st.session_state.processed_sessions.add(session_id)
            try: st.query_params.clear()
            except: pass

    # ---- Handle SUBSCRIPTION success ----
    # After processing, redirect to fixed success link (root with Pixel → login)

    if (
        success
        and typ == "business"
        and session_id
        and session_id not in st.session_state.processed_sessions
    ):

        try:

            sess = stripe.checkout.Session.retrieve(session_id)

            if sess.get("payment_status") == "paid":

                md = sess.get("metadata") or {}

                plan_name = md.get("plan_name", "Starter")

                amount_paid = (
                    sess.get("amount_total") or 0
                ) / 100.0

                # ==========================================
                # ACTIVATE BUSINESS PLAN VIA CREDIT ENGINE
                # ==========================================

                if not payment_exists(session_id):

                    res = purchase_plan(
                        "business",
                        user_email,
                        plan_name,
                        stripe_session_id=session_id
                    )

                    if res.get("ok"):

                        save_payment(
                            user_email,
                            amount_paid,
                            "business_plan",
                            session_id,
                            credits_purchased=res["credits"]
                        )

                        st.success(
                            f"""
                            🎉 Business Plan Activated

                            Plan: {plan_name}

                            Credits Added: {res['credits']}
                            """
                        )

                    else:

                        st.error(
                            f"Could not activate business plan: "
                            f"{res.get('reason')}"
                        )

                else:

                    st.info(
                        "Business payment already processed."
                    )

            else:

                st.warning(
                    "Payment not completed yet."
                )

        except Exception as e:

            st.error(
                f"Could not verify business payment: {e}"
            )

        finally:

            st.session_state.processed_sessions.add(
                session_id
            )

            try:
                st.query_params.clear()
            except:
                pass
    redirect_after_success_url = None
    if success and typ == "subscription" and session_id and session_id not in st.session_state.processed_sessions:
        try:
            sess = stripe.checkout.Session.retrieve(session_id)
            if sess.get("payment_status") == "paid":
                md   = sess.get("metadata") or {}
                plan = md.get("plan") or plan_qp or "Career Pro"
                plan = LEGACY_PLAN_MAP.get(plan, plan)
                amount_paid = (sess.get("amount_total") or 0) / 100.0
                if not payment_exists(session_id):
                    res = purchase_plan("individual", user_email, plan, stripe_session_id=session_id)
                    if res.get("ok"):
                        credits_to_add = res["credits"]
                        save_payment(user_email, amount_paid, "subscription", session_id, credits_purchased=credits_to_add)

                        # ✅ Log out and send to Pixel success hop (index.html handles redirect to login)
                        st.session_state.pop("user_data", None)
                        redirect_after_success_url = "https://cvolvepro.com/?trk=payment_success"

                        st.success(f"🎉 {plan} active! {credits_to_add} credits added. Redirecting to login…")
                    else:
                        st.error(f"Could not activate plan: {res.get('reason')}")
                else:
                    st.info("Subscription payment already processed.")
            else:
                st.warning("Payment not completed yet.")
        except Exception as e:
            st.error(f"Could not verify subscription payment: {e}")
        finally:
            st.session_state.processed_sessions.add(session_id)
            try: st.query_params.clear()
            except: pass

        if redirect_after_success_url:
            st.markdown(f'<meta http-equiv="refresh" content="0; url={redirect_after_success_url}">', unsafe_allow_html=True)
            st.stop()

    # ---- Current status ----
    subscription = check_subscription(user_email)
    credits_now  = get_user_credits(user_email)
    try:
        bal = wallet_balance("individual", user_email)
        credits_now = bal["total"]
        pack_now = bal["pack_credits"]
    except Exception:
        pack_now = 0
    if subscription:
        st.success(f"✅ Current Plan: {subscription['plan']}")
        st.info(f"📅 Next billing: {subscription['next_billing']}")
    st.info(f"💎 Current Credits: {credits_now}")
    if pack_now:
        st.caption(f"🧩 Of which {pack_now} are pack credits (valid 90 days).")

    # Show cycle window (start → +30 days)
    try:
        from datetime import timedelta
        conn = get_db_connection(); cur = conn.cursor()
        cur.execute("SELECT credit_cycle_start FROM users WHERE email=%s", (user_email,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if row and row[0]:
            cycle_start = row[0]
            ends_on = (cycle_start + timedelta(days=30)).strftime('%Y-%m-%d')
            st.caption(f"🗓️ Credit cycle started: {cycle_start.strftime('%Y-%m-%d')} • Ends: {ends_on}")
    except Exception:
        pass

    st.markdown("### 💰 Purchase Options")
    # Define once so both columns can use it
    base_url = st.secrets.get("BASE_URL", "http://localhost:8501")
    col1, col2 = st.columns(2)

    is_business_user = st.session_state.get(
        "business_logged_in",
        False
    )

    # LEFT: Credit Packages
    with col1:
        # ==================================================
        # BUSINESS USER BILLING
        # ==================================================

        if is_business_user:

            st.markdown("## 🏢 Business Plans")

            st.caption(
                "Corporate plans designed for recruiters, agencies, HR teams, and enterprises."
            )

            plan_info = get_business_plan_info(
                st.session_state.business_email
            )

            if plan_info and plan_info["current_plan"]:

                expiry_text = ""

                if plan_info["plan_expiry"]:
                    expiry_text = plan_info["plan_expiry"].strftime(
                        "%d %b %Y"
                    )

                st.info(
                    f"""
            🏢 Current Plan: {plan_info['current_plan']}

            💎 Available Credits: {plan_info['credits']}

            📅 Expires On: {expiry_text}
                    """
                )

            business_prices = {
                "Starter": 149, "Growth": 299, "Pro": 449, "Plus": 699, "Enterprise": 999,
            }
            _duration_label = lambda days: ("3 Months" if days <= 90
                                            else ("6 Months" if days <= 180 else "1 Year"))
            business_plans = [
                {
                    "name": name,
                    "credits": cfg["credits"],
                    "price": business_prices.get(name, 149),
                    "duration": _duration_label(cfg["duration_days"]),
                }
                for name, cfg in pricing.CORPORATE_PLANS.items()
            ]

            for plan in business_plans:

                local_price = price_local(plan["price"])

                with st.container(border=True):

                    left_col, right_col = st.columns(
                        [6, 4],
                        vertical_alignment="center"
                    )

                    with left_col:

                        st.markdown(
                            f"### 🏢 {plan['name']}"
                        )

                        st.markdown(
                            f"## {fmt(local_price)}"
                        )

                        st.write(f"✅ {plan['credits']} AI Credits")
                        st.write(f"✅ {plan['duration']} Access")
                        st.write("✅ ATS Resume Optimization")
                        st.write("✅ Bulk Resume Generation")
                        st.write("✅ Team Hiring Support")
                        st.write("✅ Priority Business Support")

                    with right_col:

                        st.markdown("###")

                        if st.button(
                            f"Purchase",
                            key=f"business_purchase_{plan['name']}",
                            use_container_width=True
                        ):

                            checkout_url = create_checkout_session(
                                user_email=st.session_state.business_email,
                                amount=local_price,
                                payment_type="business_plan",
                                success_url=f"{base_url}?success=true&type=business",
                                cancel_url=f"{base_url}?canceled=true",
                                credits=plan["credits"],
                                currency=CURRENT_CURRENCY,
                                plan_name=plan["name"],
                                duration=plan["duration"]
                            )

                            if checkout_url:

                                st.link_button(
                                    "Continue Checkout",
                                    checkout_url,
                                    use_container_width=True
                                )

                st.markdown("")

            st.stop()
        st.markdown("#### 💎 Credit Packs")
        for pack in pricing.PACKS:
            local_amount = price_local(pack["price_usd"])
            with st.container(border=True):
                st.markdown(f"### 🧩 {pack['name']}")
                st.markdown(f"## {fmt(local_amount)}")
                st.write(f"✅ {pack['credits']} AI Credits")
                st.write(f"✅ Valid {pack['valid_days']} days")
                if st.button(f"Buy {pack['name']}", key=f"buy_{pack['name'].replace(' ', '_')}"):
                    url = create_checkout_session(
                        user_email=user_email,
                        amount=local_amount,                     # local currency amount
                        payment_type="credits",
                        success_url=f"{base_url}?success=true&type=credits",
                        cancel_url=f"{base_url}?canceled=true",
                        credits=pack["credits"],
                        pack=pack["name"],
                        currency=CURRENT_CURRENCY                # INR/EUR/USD/AED/BHD/AUD/GBP
                    )
                    if url:
                        st.markdown(f"💳 [Pay securely via Stripe]({url})", unsafe_allow_html=True)

    # RIGHT: Subscriptions
    with col2:
        if is_business_user:

            st.markdown("#### 🏢 Business Subscription")

            st.info(
                "Business accounts use corporate plans instead of individual subscriptions."
            )

            st.stop()
        st.markdown("#### 🔄 Subscription Plans")
        from database import validate_discount_code, use_discount_code, record_user_coupon_usage

        st.markdown("### 🎟️ Apply Coupon Code")
        c1, c2 = st.columns([3,1])
        with c1:
            discount_code = st.text_input("Enter coupon code")
        with c2:
            apply_now = st.button("Apply")

        coupon_msg = st.empty()
        discount_pct = 0

        if apply_now and discount_code:
            dc = discount_code.strip().upper()
            row = validate_discount_code(dc)
            if row:
                discount_pct = int(row["discount_percent"] or 0)

                if dc == "PREMIUM599":
                    coupon_msg.success("✅ Coupon applied: Career Pro for $5.99")
                    st.session_state["active_coupon"] = {
                        "code": dc,
                        "discount_pct": 0,
                        "special_offer": "premium599"
                    }
                else:
                    coupon_msg.success(f"✅ Coupon applied: {discount_pct}% off!")
                    st.session_state["active_coupon"] = {
                        "code": dc,
                        "discount_pct": discount_pct
                    }
            else:
                st.session_state.pop("active_coupon", None)
                coupon_msg.warning("❌ Invalid or expired coupon code")
        elif "active_coupon" in st.session_state:
            discount_pct = int(st.session_state["active_coupon"].get("discount_pct", 0))

        active_coupon = st.session_state.get("active_coupon", {})
        active_code = (active_coupon.get("code") or "").strip().upper()

        if active_code == "PREMIUM599":
            st.markdown("### 🎉 Special Offer")
            st.success("Career Pro subscription for $5.99")

            if st.button("Buy Career Pro Promo – $5.99", key="buy_premium599"):
                success_url = f"{base_url}?success=true&type=subscription&plan={urllib.parse.quote_plus('Career Pro')}"
                cancel_url  = f"{base_url}?canceled=true"

                session_url = create_checkout_session(
                    user_email=user_email,
                    amount=5.99,
                    payment_type="subscription",
                    success_url=success_url,
                    cancel_url=cancel_url,
                    plan="Career Pro",
                    currency="USD"
                )

                fixed_root_hop = "https://cvolvepro.com/?trk=subscribe_click"
                click_hop = f"{fixed_root_hop}&next={urllib.parse.quote_plus(session_url or '')}"
                st.markdown(
                    f'<meta http-equiv="refresh" content="0; url={click_hop}">',
                    unsafe_allow_html=True
                )
                st.stop()

        # ── Voucher redemption (Voucher Pro — 1 month, no F2F) ───────────────
        st.markdown("### 🎫 Have a voucher code?")
        st.caption(
            "Selected users get 1 month of full platform access via a voucher code. "
            "Live voice interview (ElevenLabs) is excluded; everything else is included."
        )
        v1, v2 = st.columns([3, 1])
        with v1:
            voucher_input = st.text_input(
                "Voucher code",
                key="voucher_code_input",
                placeholder="CV-XXXX-XXXX",
            )
        with v2:
            redeem_clicked = st.button("Redeem", key="redeem_voucher_btn")
        voucher_msg = st.empty()

        if redeem_clicked and voucher_input:
            result = redeem_voucher(voucher_input.strip(), user_email)
            if result.get("ok"):
                voucher_msg.success(
                    f"✅ Voucher redeemed! {result['plan']} activated for "
                    f"{result['duration_days']} days ({result['credits']} credits)."
                )
                try:
                    reset_credits_if_expired(user_email)
                except Exception:
                    pass
                st.rerun()
            else:
                reason = result.get("reason", "unknown")
                friendly = {
                    "not_found": "This code doesn't exist. Please double-check.",
                    "inactive": "This voucher has been revoked.",
                    "expired": "This voucher has expired.",
                    "max_redemptions": "This voucher has reached its redemption limit.",
                    "already_redeemed": "You've already redeemed this code.",
                    "missing_code_or_email": "Please enter a code.",
                    "plan_activation_failed": "Could not activate the plan. Please contact support.",
                }.get(reason, f"Could not redeem (reason: {reason}).")
                voucher_msg.error(f"❌ {friendly}")

        phone = (st.session_state.get("user_data", {}).get("phone") or "").strip()
        is_india_user = phone.startswith("+91")

        for plan_name, cfg in pricing.PLANS.items():
            # Voucher-only plans aren't user-buyable — activated via redemption above.
            if cfg.get("voucher_only"):
                continue
            price_usd = cfg["price_usd"]

            special_discount = get_user_special_discount(
                user_email,
                plan_name
            )

            effective_discount = max(
                discount_pct,
                special_discount
            )

            if is_india_user:

                base_local = cfg["price_inr"]

                final_local = round(
                    base_local * (1 - effective_discount / 100.0)
                )

                display_currency = "INR"
                display_symbol = "₹"

            else:

                base_local = price_local(price_usd)

                effective_usd = price_usd * (
                    1 - effective_discount / 100.0
                )

                final_local = price_local(
                    effective_usd
                )

                display_currency = CURRENT_CURRENCY
                display_symbol = SYMBOL

            with st.expander(f"{plan_name} – {display_symbol}{base_local:,.0f}" if display_currency == "INR"
                            else f"{plan_name} – {fmt(base_local)}"):

                st.markdown("✅ Premium AI Model")
                st.markdown(f"✅ {cfg['monthly_credits']} Credits / month")
                st.markdown("✅ ATS Score Checker")
                st.markdown("✅ CV Generator")
                st.markdown("✅ CL Generator")
                st.markdown("✅ Interview Q&A")
                if cfg.get("f2f"):
                    st.markdown(f"✅ Live F2F Mock Interview (up to {cfg['f2f_max_minutes']} min)")
                else:
                    st.markdown("❌ Live F2F Mock Interview (Interview Pro only)")

                if effective_discount:

                    st.success(
                        f"🎉 {effective_discount}% OFF Applied"
                    )

                    st.markdown(
                        f"~~{display_symbol}{base_local:,.0f}~~ → "
                        f"**{display_symbol}{final_local:,.0f}**"
                    )

                if st.button(
                    f"Subscribe to {plan_name} – "
                    f"{display_symbol}{final_local:,.0f}" if display_currency == "INR"
                    else f"Subscribe to {plan_name} – {fmt(final_local)}",
                    key=f"sub_{plan_name.replace(' ','_')}"
                ):
                    success_url = f"{base_url}?success=true&type=subscription&plan={urllib.parse.quote_plus(plan_name)}"
                    cancel_url  = f"{base_url}?canceled=true"

                    
                    session_url = create_checkout_session(
                        user_email=user_email,
                        amount=final_local,              # ✅ ₹699 / ₹1499 or converted value
                        payment_type="subscription",
                        success_url=success_url,
                        cancel_url=cancel_url,
                        plan=plan_name,
                        currency=display_currency        # ✅ INR for India, else converted currency
                    )

                    fixed_root_hop = "https://cvolvepro.com/?trk=subscribe_click"
                    click_hop = f"{fixed_root_hop}&next={urllib.parse.quote_plus(session_url or '')}"
                    st.markdown(
                        f'<meta http-equiv="refresh" content="0; url={click_hop}">',
                        unsafe_allow_html=True
                    )
                    st.stop()



def create_word_document(content):
    current_section = ""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # Set base font and spacing
    style = doc.styles['Normal']
    font = style.font  # type: ignore[reportAttributeAccessIssue]
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in content.split('\n'):
        if not line.strip():
            continue

        text = line.strip()
        clean_text = text.replace("**", "")  # ✅ Remove markdown asterisks only

        # Detect if it's a section header (fully uppercase and ends with ":")
        is_section_header = clean_text.endswith(':') and clean_text == clean_text.upper()

        if is_section_header:
            current_section = clean_text[:-1].lower()
            doc.add_paragraph()

        if current_section == "work experience" and "|" in clean_text and not clean_text.startswith("•"):
            spacer_para = doc.add_paragraph()
            spacer_para.paragraph_format.space_after = Pt(1)

        para = doc.add_paragraph()
        run = para.add_run(clean_text)

        # ✅ Keep formatting rules
        if is_section_header:
            run.bold = True
            add_bottom_border(para)

        elif current_section == "work experience" and "|" in clean_text and not clean_text.startswith("•"):
            run.bold = True

        elif current_section == "projects" and not clean_text.startswith("•"):
            run.bold = True

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.0

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def analyze_ats_compatibility():
    """Analyze ATS compatibility of generated CV using AI-based scorer with local fallback.
    """
    if st.session_state.cv_preview:
        jd = st.session_state.get('job_description', '') or st.session_state.get('manual_jd', '')
        metadata = st.session_state.get("cv_optimization_metadata") or {}

        def _to_numeric(val):
            if val is None:
                return 0
            if isinstance(val, (int, float)):
                return int(val)
            s = str(val).strip().rstrip('%')
            try:
                return int(float(s))
            except (ValueError, TypeError):
                return 0

        if metadata:
            analysis = {
                "score": metadata.get("ats_score"),
                "keyword_match": metadata.get("keyword_match"),
                "missing_keywords": metadata.get("missing_keywords") or [],
                "suggestions": metadata.get("fixes_applied") or [],
            }
        else:
            analysis = None
            if jd:
                try:
                    from cv_generator import analyze_cv_ats_score
                    analysis = analyze_cv_ats_score(st.session_state.cv_preview, jd)
                except Exception:
                    analysis = None

            if not analysis or analysis.get("score") is None:
                analysis = optimize_keywords(st.session_state.cv_preview, jd)

        score_val = _to_numeric(analysis.get('score', 0))
        kw_val = _to_numeric(analysis.get('keyword_match', 0))

        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Target ATS", f"{metadata.get('target_ats_score', 100)}%")

        with col2:
            st.metric("Measured ATS Score", f"{score_val}%")
            st.progress(min(1.0, max(0.0, float(score_val) / 100.0)))

        with col3:
            st.metric("Keyword Match", f"{kw_val}%")
            st.progress(min(1.0, max(0.0, float(kw_val) / 100.0)))

        if metadata.get("repair_passes_used") is not None:
            st.caption(f"Optimization repair passes used: {metadata.get('repair_passes_used', 0)}")

        if metadata.get('fixes_applied'):
            st.markdown("### ✅ Optimization Fixes Applied")
            for fix in metadata['fixes_applied']:
                st.markdown(f"• {fix}")
        elif analysis.get('suggestions'):
            st.markdown("### 💡 Improvement Suggestions")
            for suggestion in analysis['suggestions']:
                st.markdown(f"• {suggestion}")

        if analysis.get('missing_keywords'):
            st.markdown("### 🔍 Missing Keywords")
            for keyword in analysis['missing_keywords'][:5]:  # Show only first 5
                st.markdown(f"• {keyword}")

        if metadata.get("unsupported_gaps"):
            st.markdown("### ⚠️ JD Terms Needing Candidate Evidence")
            for keyword in metadata["unsupported_gaps"][:5]:
                st.markdown(f"• {keyword}")

def _credit_account_type():
    return "business" if st.session_state.get("account_type") == "business" else "individual"


def check_user_access(required_credits=2, feature=None):
    email = st.session_state.user_data['email']
    if email and ("tester@cvolvepro.com" in email.lower() or "test" in email.lower()):
        return True
    account_type = "business" if st.session_state.get("account_type") == "business" else "individual"
    try:
        return has_enough(account_type, email, amount=required_credits, feature=feature)
    except Exception:
        return False


def deduct_user_credits(email, amount, feature=None):
    """Deduct credits for individual or business users via the credit engine."""
    try:
        # Test accounts receive free/unlimited credits without deduction
        if email and ("tester@cvolvepro.com" in email.lower() or "test" in email.lower()):
            if feature:
                try:
                    record_credit_usage(email, feature, 0)
                except Exception:
                    pass
            return True

        account_type = "business" if st.session_state.get("account_type") == "business" else "individual"
        result = spend_credits(
            account_type,
            email,
            feature or "general",
            amount=amount,
        )
        if result.get("ok"):
            if feature:
                try:
                    record_credit_usage(email, feature, amount)
                except Exception:
                    pass
            st.session_state["last_credit_activity"] = {
                "feature": feature or "general",
                "amount": amount,
                "timestamp": time.time(),
            }
            return True

        st.warning("You don’t have enough credits to complete this action.")
        return False

    except Exception as e:
        st.error(
            f"Error deducting credits: {str(e)}"
        )
        return False



def show_payment_page():
    """Show payment processing page"""
    st.markdown("## 💳 Purchase Credits")
    # Implementation would show Stripe payment form
    pass

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')     # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    borders.append(bottom)
    pPr.append(borders)

def send_otp_email(email: str, otp: str) -> bool:
    """Send a 6-digit OTP via Resend."""
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,  # e.g., "CVOLVE PRO <verify@yourdomain.com>"
            "to": [email],
            "subject": "Your CVOLVE PRO verification code",
            "html": f"""
            <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
              <p>Hello,</p>
              <p>Your <b>CVOLVE PRO</b> verification code is:</p>
              <p style="font-size:22px;font-weight:700;letter-spacing:3px">{otp}</p>
              <p>This code will expire in <b>10 minutes</b>.</p>
              <p>Thanks,<br/>CVOLVE PRO</p>
            </div>
            """,
        })
        return True
    except Exception as e:
        import streamlit as st
        st.error(f"Email send failed: {e}")
        return False



def show_register_page():
    st.markdown("## 🆕 Create Your Account")

    # --- Helper: build full country list once ---
    @st.cache_data
    def get_all_country_dial_codes(default_region="IN"):
        items = []
        try:
            regions = sorted(phonenumbers.SUPPORTED_REGIONS)
        except Exception:
            regions = ["IN", "US", "GB"]  # tiny fallback
        for region in regions:
            try:
                code = phonenumbers.country_code_for_region(region)
                country = pycountry.countries.get(alpha_2=region)
                name = getattr(country, "name", region)
                label = f"{name} (+{code})"
                items.append((label, region, f"+{code}"))
            except Exception:
                continue
        items.sort(key=lambda x: (x[1] != default_region, x[0]))
        return items

    # --- Session state initialization ---
    if "register_name" not in st.session_state:
        st.session_state["register_name"] = ""
    if "register_email_address" not in st.session_state:
        st.session_state["register_email_address"] = ""
    if "register_password" not in st.session_state:
        st.session_state["register_password"] = ""
    if "register_phone" not in st.session_state:
        st.session_state["register_phone"] = ""
    if "register_region" not in st.session_state:
        st.session_state["register_region"] = "IN"
    if "register_country_code" not in st.session_state:
        st.session_state["register_country_code"] = "+91"
    # Pending registration (stored until OTP is verified)
    if "pending_registration" not in st.session_state:
        st.session_state["pending_registration"] = None
    if "awaiting_otp_email" not in st.session_state:
        st.session_state["awaiting_otp_email"] = None

    # --- Widget on_change handlers ---
    def update_name():
        st.session_state["register_name"] = st.session_state.name_input
    def update_email():
        st.session_state["register_email_address"] = st.session_state.email_input
    def update_phone():
        st.session_state["register_phone"] = st.session_state.phone_input
    def update_password():
        st.session_state["register_password"] = st.session_state.password_input

    # --- Inputs ---
    st.text_input("Full Name", key="name_input",
                  value=st.session_state["register_name"], on_change=update_name)

    st.text_input("Email Address", key="email_input",
                  value=st.session_state["register_email_address"], on_change=update_email)
    # Gmail-only hint
    st.markdown("<small>We’ll email you a 6‑digit code to verify.</small>", unsafe_allow_html=True)

    countries = get_all_country_dial_codes(default_region="IN")
    col_code, col_number = st.columns([2, 3])
    with col_code:
        selected_label = st.selectbox(
            "Country / Code",
            options=[c[0] for c in countries],
            index=0,
            key="register_country_label"
        )
        sel = next(c for c in countries if c[0] == selected_label)
        st.session_state["register_region"] = sel[1]
        st.session_state["register_country_code"] = sel[2]
    with col_number:
        st.text_input("Phone (without country code) – optional",
                      key="phone_input",
                      value=st.session_state["register_phone"],
                      on_change=update_phone,
                      placeholder="e.g., 9876543210")

    st.text_input("Password", type="password",
                  key="password_input",
                  value=st.session_state["register_password"],
                  on_change=update_password)

    # --- Register: send OTP first; DO NOT create user yet ---
    if st.button("Register", key="register_button"):
        # Read the widget keys directly. Streamlit keeps these current on every rerun
        # (including this button click), whereas the register_* mirrors are only synced
        # by the on_change callbacks — which don't fire if the user clicks Register while
        # a field still has focus, leaving the mirrors empty and triggering a false
        # "invalid email" / "fill in all fields" error.
        name = st.session_state.get("name_input", "").strip()
        email = st.session_state.get("email_input", "").strip().lower()
        raw_phone = st.session_state.get("phone_input", "").strip()
        region = st.session_state.get("register_region", "IN")
        e164_phone = ""

        if not re.match(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$", email):
            st.error("Please enter a valid email address.")
            return

        # Prevent duplicate registration
        if get_user_data(email):
            st.error("This email is already registered. Please log in or use a different Email address.")
            return

        # Validate & format phone if provided
        if raw_phone:
            try:
                pn = pn_parse(raw_phone, region)
                if is_valid_number(pn):
                    e164_phone = format_number(pn, PhoneNumberFormat.E164)
                else:
                    st.warning("The phone number looks invalid for the selected country.")
                    e164_phone = f"{st.session_state['register_country_code']} {raw_phone}"
            except Exception:
                st.warning("Could not parse the phone number. Please check it.")
                e164_phone = f"{st.session_state['register_country_code']} {raw_phone}"

        password = st.session_state.get("password_input", "").strip()
        if not (name and email and password):
            st.error("Please fill in all required fields.")
            return

        # Build pending registration payload
        password_hash = hash_password(password)
        otp = f"{secrets.randbelow(1000000):06d}"
        expires_at = time.time() + 10*60  # 10 minutes

        # Store pending registration in session (not in DB)
        st.session_state["pending_registration"] = {
            "name": name,
            "email": email,
            "phone": e164_phone,
            "password_hash": password_hash,
            "otp": otp,
            "expires_at": expires_at
        }

        # Send OTP
        if send_otp_email(email, otp):
            st.session_state["awaiting_otp_email"] = email
            st.success("✅ We sent a 6‑digit OTP to your email. Enter it below to verify and complete registration.")
        else:
            st.error("Could not send OTP.")

    # --- OTP verification UI ---
    if st.session_state.get("awaiting_otp_email"):
        st.markdown("### 🔒 Verify your email")
        v_email = st.session_state["awaiting_otp_email"]
        otp_input = st.text_input("Enter the 6-digit OTP", max_chars=6, key="otp_input")

        col_v1, col_v2 = st.columns(2)
        with col_v1:
            if st.button("Verify OTP"):
                code = (otp_input or "").strip()
                pend = st.session_state.get("pending_registration")
                # Validate session payload
                if not pend or pend.get("email") != v_email:
                    st.error("Registration session expired. Please register again.")
                    st.session_state["pending_registration"] = None
                    st.session_state["awaiting_otp_email"] = None
                    return
                if time.time() > pend["expires_at"]:
                    st.error("OTP expired. Please click 'Resend OTP'.")
                    return
                if not (len(code) == 6 and code.isdigit()):
                    st.error("Please enter a valid 6-digit code.")
                    return
                if code != pend["otp"]:
                    st.error("Invalid OTP. Please try again.")
                    return

                # ✅ OTP OK: now create the user in DB and mark verified
                try:
                    register_user(pend["name"], pend["email"], pend["phone"], pend["password_hash"], "")
                    # Mark verified immediately
                    conn = get_db_connection()
                    cur = conn.cursor()
                    cur.execute("UPDATE users SET is_verified=TRUE, verification_token=NULL WHERE email=%s", (pend["email"],))
                    conn.commit(); cur.close(); conn.close()

                    st.success("🎉 Email verified! Your account is created. You can now log in.")
                    # Clear pending state
                    st.session_state["pending_registration"] = None
                    st.session_state["awaiting_otp_email"] = None

                    # Optionally auto-redirect to login
                    st.session_state.page = "login"
                    st.rerun()
                except Exception as e:
                    st.error(f"Error completing registration: {str(e)}")

        with col_v2:
            if st.button("Resend OTP"):
                pend = st.session_state.get("pending_registration")
                if not pend:
                    st.warning("Registration session not found. Please start again.")
                else:
                    new_otp = f"{secrets.randbelow(1000000):06d}"
                    pend["otp"] = new_otp
                    pend["expires_at"] = time.time() + 10*60
                    st.session_state["pending_registration"] = pend
                    if send_otp_email(pend["email"], new_otp):
                        st.info("📩 A new OTP has been sent.")
                    else:
                        st.error("Could not resend OTP.")

    if st.button("⬅ Back to Login"):
        st.session_state.page = "login"
        st.rerun()



if __name__ == "__main__":
    main()
