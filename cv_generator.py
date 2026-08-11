import os
import re
import json
import hashlib
import time
import logging
from datetime import datetime
from typing import Any, Optional
import PyPDF2 as pdf
from docx import Document
import google.generativeai as genai
from google.generativeai import types
from pydantic import BaseModel
from utils import optimize_keywords, enforce_page_limit
from dotenv import load_dotenv
from streamlit import session_state as st_session
import openai

logger = logging.getLogger(__name__)

# Cache for ATS scores (max 100 entries, TTL 1 hour)
_ats_cache = {}
_ats_cache_ttl = 3600

def _get_cache_key(cv_content, job_description=""):
    """Generate cache key from CV + JD content"""
    content = f"{cv_content}|{job_description}"
    return hashlib.md5(content.encode()).hexdigest()

def _get_session_ai_model():
    """Safely get ai_model from Streamlit session state without warnings when running outside Streamlit."""
    try:
        import threading
        t = threading.current_thread()
        if hasattr(t, "streamlit_script_run_ctx") and getattr(t, "streamlit_script_run_ctx") is not None:
            return st_session.get("ai_model")
    except Exception:
        pass
    return None

from utils import get_gemini_response
    
openai.api_key = os.getenv("OPENAI_API_KEY")


# Initialize Gemini client
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))  # type: ignore
model = genai.GenerativeModel("gemini-2.5-flash")  # type: ignore


# =============================================================================
# Phase 2 — CV ↔ JD alignment: truthfulness guardrail + gap analysis
# See docs/CV_JD_ALIGNMENT_PLAN.md. The guardrail replaces the old prompts that
# explicitly instructed the model to fabricate experience.
# =============================================================================
TRUTHFULNESS_GUARDRAIL = (
    "HARD RULE — TRUTHFULNESS (highest priority, overrides any other instruction below):\n"
    "Use facts from the candidate's résumé and VERIFIED EXPERIENCE block (if present) as your foundation.\n\n"
    "You MAY:\n"
    "- Reframe and rephrase existing experience to align with the job description's terminology\n"
    "- Surface genuinely-held skills that are reasonably inferred from the candidate's demonstrated experience\n"
    "- Emphasize relevant aspects of existing roles that match the target position\n"
    "- Connect related skills and experiences to show applicability to the JD\n"
    "- Highlight transferable skills that are logically supported by the candidate's background\n\n"
    "You MUST NOT:\n"
    "- Invent new employers, job titles, dates, degrees, or certifications\n"
    "- Fabricate metrics, achievements, or project outcomes that aren't supported by the candidate's history\n"
    "- Present placeholder or example numbers as real achievements\n"
    "- Add skills or tools the candidate has never used or demonstrated\n\n"
    "The goal is to present the candidate's authentic experience in the most compelling way for the target role — not to invent experience they don't have."
)


def build_verified_context_block(extra_context) -> str:
    """Format the candidate's verified follow-up answers for prompt injection.

    Accepts a preformatted string or a {area: answer} dict. Returns "" when empty
    so prompts are byte-identical to the no-alignment path when there are no answers.
    """
    if not extra_context:
        return ""
    if isinstance(extra_context, dict):
        lines = [
            f"- {area}: {str(ans).strip()}"
            for area, ans in extra_context.items()
            if ans and str(ans).strip()
        ]
        body = "\n".join(lines)
    else:
        body = str(extra_context).strip()
    if not body:
        return ""
    return (
        "VERIFIED EXPERIENCE (provided by the candidate in follow-up answers; treat as true "
        "and use it, but still do not invent anything beyond it):\n" + body
    )


def hash_jd(job_description) -> str:
    """Stable short hash of a JD (whitespace/case-insensitive) — keys stored answers."""
    import hashlib
    norm = re.sub(r"\s+", " ", (job_description or "").strip().lower())
    return hashlib.sha256(norm.encode()).hexdigest()[:16]


def parse_gap_analysis(raw_text, max_gaps: int = 5) -> dict:
    """Parse gap-analysis model output into a safe dict. NEVER raises.

    Returns {"sufficient": bool, "overall_match": int|None,
             "gaps": [{"id","area","why","question","example"}]}.
    Fails OPEN (sufficient=True, no gaps) on any parse problem so a bad model
    response can never block generation.
    """
    fallback = {"sufficient": True, "overall_match": None, "gaps": []}
    if not raw_text or not str(raw_text).strip():
        return fallback
    text = str(raw_text).strip()
    if text.startswith("```"):                       # strip ```json ... ``` fences
        text = re.sub(r"^```[a-zA-Z]*\s*", "", text)
        text = re.sub(r"\s*```$", "", text).strip()
    if not text.startswith("{"):                     # pull first {...} out of prose
        m = re.search(r"\{.*\}", text, re.DOTALL)
        if m:
            text = m.group(0)
    try:
        data = json.loads(text)
    except Exception:
        return fallback
    if not isinstance(data, dict):
        return fallback

    gaps = []
    for i, g in enumerate(data.get("gaps") or []):
        if not isinstance(g, dict):
            continue
        question = str(g.get("question") or "").strip()
        if not question:                             # a gap with no question is useless
            continue
        area = str(g.get("area") or "").strip() or question
        example_text = str(g.get("example") or "").strip()
        if len(example_text) < 15:
            example_text = ""  # Too short to be useful; leave text area empty
        gaps.append({
            "id": str(g.get("id") or f"gap_{i + 1}").strip(),
            "area": area,
            "why": str(g.get("why") or "").strip(),
            "question": question,
            "example": example_text,
        })
    gaps = gaps[:max_gaps]

    try:
        val = data.get("overall_match")
        overall = int(str(val)) if val is not None else None
    except (TypeError, ValueError):
        overall = None

    # No actionable gaps → treat as sufficient regardless of the model's flag.
    sufficient = True if not gaps else bool(data.get("sufficient", False))
    return {"sufficient": sufficient, "overall_match": overall, "gaps": gaps}


def analyze_cv_jd_gaps(resume_text, job_description, language: str = "English", max_gaps: int = 5) -> dict:
    """One LLM call → structured gap analysis (see parse_gap_analysis for shape).

    Fail-open: any error returns sufficient=True so generation is never blocked.
    Cost: exactly +1 model call; callers should memoize per (resume, JD) per session.
    """
    if not resume_text or not job_description:
        return {"sufficient": True, "overall_match": None, "gaps": []}

    prompt = f"""
    You are a career coach comparing a candidate's résumé against a target job description.
    Identify AT MOST {max_gaps} areas the JD clearly requires but the résumé shows no or weak
    evidence for. For each gap, write:
    1. ONE clear question the candidate can answer to supply real evidence.
    2. A strong, ELABORATIVE, and PERSONALISED example answer written AS IF the candidate is replying,
       using specific details extracted from THEIR résumé (real project names, company names,
       tech stack, metrics, dates). The example answer MUST be:
       - Written in first person ("I led…", "At [Company], I…")
       - 3-5 elaborative sentences providing comprehensive context (scale, responsibilities, tools, and results)
       - Directly reference real experience from the résumé that is closest to the gap
       - Detailed, complete, and professional so the candidate sees an ideal response
       - If the résumé has NO related experience at all for a gap, write
         "No direct experience found in your CV — please describe any related experience you have."
         instead of fabricating details.

    Return STRICT JSON ONLY (no prose, no code fences) in exactly this shape:
    {{
      "sufficient": <true if the résumé already has enough evidence and needs no questions>,
      "overall_match": <integer 0-100 estimate>,
      "gaps": [
        {{"id": "<short_slug>", "area": "<missing area>", "why": "<why it is a gap>",
          "question": "<one clear question>",
          "example": "<an elaborative, personalised, first-person draft answer grounded in the candidate's CV>"}}
      ]
    }}
    If the résumé already covers the JD well, return "sufficient": true and an empty "gaps" list.
    Write the "question" and "example" text in {language}.

    RÉSUMÉ:
    {resume_text}

    JOB DESCRIPTION:
    {job_description}
    """
    try:
        if _get_session_ai_model() == "openai":
            response = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You output only strict JSON."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
            )
            raw = response.choices[0].message.content
        else:
            response = model.generate_content(
                prompt,
                generation_config={"temperature": 0.1, "response_mime_type": "application/json"},
            )
            raw = response.text if (response and getattr(response, "text", None)) else ""
    except Exception:
        return {"sufficient": True, "overall_match": None, "gaps": []}
    return parse_gap_analysis(raw, max_gaps=max_gaps)


class CVOptimization(BaseModel):
    """CV optimization response model"""
    ats_score: int
    missing_keywords: list
    optimized_content: str
    suggestions: list


class CVOptimizationResult(BaseModel):
    """Structured result for the closed-loop CV optimizer."""
    optimized_content: str
    ats_score: Optional[int] = None
    keyword_match: Optional[int] = None
    missing_keywords: list[str] = []
    repair_passes_used: int = 0
    fixes_applied: list[str] = []
    unsupported_gaps: list[str] = []
    target_ats_score: int = 100


def _unique_preserve_order(items: list[str], limit: Optional[int] = None) -> list[str]:
    seen = set()
    out = []
    for item in items:
        val = str(item or "").strip()
        key = val.lower()
        if not val or key in seen:
            continue
        seen.add(key)
        out.append(val)
        if limit and len(out) >= limit:
            break
    return out


def _term_present(term: str, text: str) -> bool:
    """ATS-ish phrase check: direct phrase or all core words present."""
    if not term or not text:
        return False
    term_low = term.lower().strip()
    text_low = text.lower()
    if term_low in text_low:
        return True
    words = [w for w in re.findall(r"[a-zA-Z0-9\+#\.\-/]+", term_low) if len(w) > 2]
    generic = {"developer", "engineer", "specialist", "manager", "role", "team"}
    core = [w for w in words if w not in generic]
    return bool(core and all(w in text_low for w in core))


def _context_text(extra_context: Any) -> str:
    if not extra_context:
        return ""
    if isinstance(extra_context, dict):
        return "\n".join(str(v) for v in extra_context.values() if v)
    if isinstance(extra_context, list):
        return "\n".join(str(v) for v in extra_context if v)
    return str(extra_context)


def build_jd_keyword_plan(resume_text: str, job_description: str, extra_context: Any = "") -> dict[str, list[str]]:
    """Build a deterministic keyword plan and evidence classification for the JD."""
    if not job_description:
        return {
            "supported": [],
            "candidate_verified": [],
            "missing_evidence": [],
            "required_keywords": [],
            "role_title_terms": [],
        }

    try:
        from utils import extract_ats_phrases, extract_domain_keywords, filter_keywords
    except Exception:
        extract_ats_phrases = lambda text: []  # type: ignore
        extract_domain_keywords = lambda text: []  # type: ignore
        filter_keywords = lambda kws: kws  # type: ignore

    phrases = list(extract_ats_phrases(job_description) or [])
    phrases.extend(extract_domain_keywords(job_description) or [])

    # Pull title-like terms from the first JD lines. This is intentionally narrow:
    # role terms help header/summary alignment without polluting the skills list.
    first_lines = " ".join(str(job_description).splitlines()[:4])
    title_terms = filter_keywords(re.findall(r"\b[A-Za-z][A-Za-z0-9\+#\.\-]{2,}\b", first_lines.lower()))
    phrases.extend(title_terms[:6])

    terms = _unique_preserve_order([p.lower() for p in phrases], limit=35)
    resume = resume_text or ""
    verified = _context_text(extra_context)

    supported = [t for t in terms if _term_present(t, resume)]
    candidate_verified = [t for t in terms if t not in supported and _term_present(t, verified)]
    missing = [t for t in terms if t not in supported and t not in candidate_verified]

    required = _unique_preserve_order(supported + candidate_verified, limit=25)
    return {
        "supported": supported,
        "candidate_verified": candidate_verified,
        "missing_evidence": missing[:15],
        "required_keywords": required,
        "role_title_terms": _unique_preserve_order(title_terms[:6]),
    }


def _format_keyword_plan_for_prompt(keyword_plan: dict[str, list[str]]) -> str:
    if not keyword_plan:
        return ""
    required = keyword_plan.get("required_keywords") or []
    missing = keyword_plan.get("missing_evidence") or []
    role_terms = keyword_plan.get("role_title_terms") or []
    lines = [
        "DETERMINISTIC ATS KEYWORD PLAN:",
        "- Supported/verified JD terms to include naturally wherever truthful: "
        + (", ".join(required) if required else "None found; rely only on resume evidence."),
        "- Target role/title terms to surface in the header or professional summary when truthful: "
        + (", ".join(role_terms) if role_terms else "None detected."),
        "- JD terms with missing evidence. Do NOT claim these as experience unless the resume or verified answers support them: "
        + (", ".join(missing) if missing else "None."),
    ]
    return "\n".join(lines)

def standardize_cv_formatting(cv_text: str) -> str:
    """Standardize bullet symbols to '•' and clean up section headers per docs/resume_fix_prompt.md."""
    if not cv_text:
        return ""
    lines = cv_text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^[•\-\*]\s*(SUMMARY|EDUCATION|EXPERIENCE|WORK EXPERIENCE|PROJECTS|TECHNICAL SKILLS|ACHIEVEMENTS)\b', stripped, re.IGNORECASE):
            line = re.sub(r'^[•\-\*]\s*', '', stripped)
        elif re.match(r'^[-\*]\s+', stripped) and not stripped.startswith("•"):
            line = re.sub(r'^[-\*]\s+', '• ', stripped)
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

def _finalize_optimized_cv(raw_cv: str, job_description: str = "") -> str:
    """Clean, format, and standardize the generated CV content."""
    if not raw_cv:
        return ""
    cleaned = clean_cv_content(raw_cv)
    standardized = standardize_cv_formatting(cleaned)
    enhanced = enhance_action_verbs(standardized, intensity="High")
    return enhanced

def extract_resume_text(uploaded_file):
    """Extract text from uploaded resume file"""
    if not uploaded_file:
        return ""
    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
    except Exception:
        pass

    if uploaded_file.name.lower().endswith(".pdf"):
        try:
            reader = pdf.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text.strip()
        except Exception:
            try:
                uploaded_file.seek(0)
                reader = pdf.PdfReader(uploaded_file)
                text = ""
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
                return text.strip()
            except Exception:
                return ""
    elif uploaded_file.name.lower().endswith(".docx"):
        try:
            doc = Document(uploaded_file)
            return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception:
            return ""
    else:
        return ""

def _generate_cv_once(
    resume_text,
    job_description,
    target_match=100,
    template="professional",
    sections=None,
    quantitative_focus=60,
    action_verb_intensity="High",
    keyword_matching="Balanced",
    language="English",
    model_choice="premium",
    extra_context="",
    keyword_plan=None,
    repair_instructions="",
    **kwargs
):
    """Generate one optimized CV draft using Gemini AI.

    NOTE:
      - `model_choice` is a string (e.g., "premium" or "premium_classic") coming
        from the extension/UI. It does NOT replace the module-level `model`
        object used to call Gemini; it is only used for prompt/config choices.
      - We normalize `sections` so the function safely accepts dict / list / str.
    """
    if target_match is None:
        target_match = 100

    # --- Normalize `sections` (accept list/tuple/str/dict safely) ---
    if sections is None:
        sections = {}
    elif isinstance(sections, (list, tuple)):
        # convert ['Summary','Skills'] -> {'Summary': True, 'Skills': True}
        try:
            sections = {str(s): True for s in sections}
        except Exception:
            sections = {}
    elif isinstance(sections, str):
        sections = {sections: True}
    elif isinstance(sections, dict):
        # ensure boolean values for includes
        sections = {str(k): bool(v) for k, v in sections.items()}
    else:
        # unexpected type — fallback to empty dict
        sections = {}

    # --- Normalize language ---
    if not language:
        language = "English"
    else:
        language = str(language)

    # --- Normalize model_choice (string from UI) ---
    try:
        model_choice = str(model_choice).lower() if model_choice is not None else "premium"
    except Exception:
        model_choice = "premium"

    # Build sections string (safe now)
    sections_list = [section for section, include in sections.items() if include]
    sections_string = ", ".join(sections_list)

    
    # Adjust prompt based on settings
    intensity_mapping = {
        "Moderate": "moderate use of action verbs",
        "High": "strong emphasis on action verbs",
        "Very High": "maximum use of powerful action verbs"
    }
    
    matching_mapping = {
        "Conservative": "maintain authenticity while incorporating key terms",
        "Balanced": "strategically integrate job description keywords",
        "Aggressive": "maximize keyword density and exact phrase matching"
    }

    language_instruction = f"""
    Respond ONLY in {language}.
    Generate the entire resume and all section headers in {language}.
    Use native {language} formatting for dates, section names, and style.
    """
    
    verified_block = build_verified_context_block(extra_context)
    keyword_plan_block = _format_keyword_plan_for_prompt(keyword_plan or {})
    repair_block = str(repair_instructions or "").strip()
    prompt_5 = f"""
    {language_instruction}

    {TRUTHFULNESS_GUARDRAIL}

    {keyword_plan_block}

    {repair_block}

    You are an expert resume writer and ATS optimization specialist.
    Your objective is to achieve a BENCHMARK ATS SCORE OF AT LEAST 85-100% against the target Job Description by reframing the candidate's real experience and incorporating all candidate-verified answers, WITHOUT inventing any unverified companies, job titles, or fake metrics.

    REWRITING & ATS OPTIMIZATION DIRECTIVES:
    1. VERIFIED ANSWERS & FOLLOW-UP Q&A INTEGRATION (CRITICAL):
       - If a VERIFIED EXPERIENCE block is provided below, treat all candidate answers, skills, and tools in that block as candidate-verified truths.
       - You MUST seamlessly embed EVERY skill, tool, methodology, and achievement mentioned in the VERIFIED EXPERIENCE block into the KEY SKILLS section and the relevant bullet points under WORK EXPERIENCE or PROJECTS.
       - These verified answers were provided specifically to resolve ATS gaps identified for this Job Description. Failing to include them will cause the resume to fail ATS cutoff screening.

    2. EXACT TERMINOLOGY & KEYWORD ALIGNMENT:
       - Identify every hard technical skill, platform, framework, library, tool, database, architecture pattern, and methodology mentioned in the Job Description.
       - Use exact terminology and spelling from the Job Description across the resume only when supported by the résumé or VERIFIED EXPERIENCE block.
       - Maintain both full terms and standard acronyms (e.g. "Continuous Integration/Continuous Deployment (CI/CD)").

    3. PROFESSIONAL SUMMARY:
       - Open directly with a strong target profile headline referencing the exact role title from the Job Description (e.g. "[Exact Target Role Title] with X+ years of experience in...").
       - Weave the top 4-6 critical technical keywords and domain skills from the JD into this 3-4 sentence paragraph.

    4. KEY SKILLS SECTION:
       - Create bulleted categories: Technical Skills, Tools & Platforms, Methodologies & Core Competencies.
       - Ensure every hard skill and tool from the JD that the candidate possesses or verified in follow-up answers is explicitly listed here.

    5. WORK EXPERIENCE & PROJECTS:
       - {intensity_mapping.get(action_verb_intensity, intensity_mapping["High"])}
       - {matching_mapping.get(keyword_matching, matching_mapping["Balanced"])}
       - Keep original company names, job titles, and employment dates intact.
       - Rewrite bullet points to start with strong, high-impact action verbs (e.g., Spearheaded, Architected, Engineered, Optimized, Implemented).
       - In every bullet point, explicitly reference the tools, technologies, and exact JD keywords used, and highlight quantifiable outcomes (%, numbers, performance gains).
       - Each bullet must tell a mini-story: challenge → action → result (e.g., "Identified bottleneck in data pipeline, engineered Spark-based solution, reduced processing time by 60%").
       - Integrate keywords naturally into achievement sentences — don't just list them. Write as if describing real accomplishments to a hiring manager.
       - Vary sentence openings and structure to avoid repetitive patterns that ATS parsers may flag as keyword stuffing.
       - Include context: don't just say "Used Python for automation" — say "Developed Python-based ETL pipeline that reduced data processing time by 40% and eliminated manual errors".

    6. CLEAN ATS FORMATTING:
       - Use standard, uppercase section headers EXACTLY as: PROFESSIONAL SUMMARY, KEY SKILLS, WORK EXPERIENCE, EDUCATION, PROJECTS, CERTIFICATIONS.
       - Output only standard bullet points ('• '). Do not add markdown code fences, commentary, or tags.

    OUTPUT FORMAT TEMPLATE:

    [CANDIDATE FULL NAME]
    [Phone Number] | [Email Address] | [Location / Address]
    [LinkedIn / Portfolio Link]

    PROFESSIONAL SUMMARY:
    [Professional summary paragraph matching target role keywords]

    KEY SKILLS:
    • Technical Skills: [Skill 1, Skill 2, Skill 3...]
    • Tools & Platforms: [Tool 1, Tool 2, Tool 3...]
    • Core Competencies: [Competency 1, Competency 2...]

    WORK EXPERIENCE:
    [Company Name] | [Job Title] | [MM/YYYY - MM/YYYY]
    • [Action verb] [Responsibility/Achievement rephrased with JD keywords]...
    • [Action verb] [Responsibility/Achievement rephrased with JD keywords]...

    EDUCATION:
    • [Degree] | [Institution] | [Graduation Year]

    PROJECTS:
    [Project Title]
    • [Bullet point detailing scope, technologies used, and outcomes]

    CERTIFICATIONS:
    • [Certification Name] - [Issuing Organization] ([Year])

    Resume Content:
    {resume_text}

    {verified_block}

    Job Description:
    {job_description}

    IMPORTANT: Output ONLY the final ATS-optimized resume. Do NOT include any analysis, markdown code blocks, comments, or notes.
    """

    
    try:
        start_time = time.time()
        model_choice = _get_session_ai_model() or "gemini"
        logger.info(f"Starting CV generation: model={model_choice}, prompt_length={len(prompt_5)}")
        
        # ✅ OpenAI Flow
        if model_choice == "openai":
            response = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": prompt_5}
                ],
                temperature=0.2
            )
            raw_cv = response.choices[0].message.content or ""
            logger.info(f"OpenAI response: tokens={response.usage.total_tokens if response.usage else 'unknown'}")
            elapsed = time.time() - start_time
            logger.info(f"CV generation completed in {elapsed:.2f}s")
            return _finalize_optimized_cv(raw_cv, job_description)

        # ✅ Gemini Flow
        else:
            response = model.generate_content(
                prompt_5,
                generation_config={
                    "temperature": 0.2,
                    "response_mime_type": "text/plain"
                }
            )

        if not response:
            raise Exception("No response received from AI")

        raw_cv = ""
        if hasattr(response, "text") and response.text:
            raw_cv = response.text
        elif response.candidates and len(response.candidates) > 0:
            candidate = response.candidates[0]
            if candidate.content and candidate.content.parts:
                raw_cv = "".join([p.text for p in candidate.content.parts if hasattr(p, 'text') and p.text])

        if not raw_cv:
            raise Exception("AI response was empty")

        elapsed = time.time() - start_time
        logger.info(f"CV generation completed in {elapsed:.2f}s")
        return _finalize_optimized_cv(raw_cv, job_description)
        
    except Exception as e:
        logger.exception(f"CV generation failed: {e}")
        raise Exception(f"Failed to generate CV: {str(e)}")


def _to_int_or_none(value: Any) -> Optional[int]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(value)
    try:
        return int(float(str(value).strip().rstrip("%")))
    except Exception:
        return None


def _measure_cv_against_jd(cv_content: str, job_description: str) -> dict[str, Any]:
    """Return measured ATS data, preferring AI analysis when it succeeds."""
    local = optimize_keywords(cv_content, job_description)
    ai = None
    try:
        ai = analyze_cv_ats_score(cv_content, job_description)
    except Exception:
        ai = None

    ai_score = _to_int_or_none((ai or {}).get("score"))
    ai_kw = _to_int_or_none((ai or {}).get("keyword_match"))
    local_score = _to_int_or_none(local.get("score")) or 0
    local_kw = _to_int_or_none(local.get("keyword_match")) or 0

    score = ai_score if ai_score is not None else local_score
    keyword_match = ai_kw if ai_kw is not None else local_kw

    missing = _unique_preserve_order(
        [str(k) for k in ((local.get("missing_keywords") or []) + ((ai or {}).get("missing_keywords") or []))],
        limit=12,
    )
    suggestions = _unique_preserve_order(
        [str(s) for s in ((local.get("suggestions") or []) + ((ai or {}).get("suggestions") or []))],
        limit=10,
    )
    return {
        "score": score,
        "keyword_match": keyword_match,
        "missing_keywords": missing,
        "suggestions": suggestions,
        "local": local,
        "ai": ai or {},
    }


def _build_repair_instructions(
    cv_content: str,
    job_description: str,
    keyword_plan: dict[str, list[str]],
    measurement: dict[str, Any],
    pass_number: int,
) -> tuple[str, list[str]]:
    quality = validate_cv_quality(cv_content, job_description)
    missing = measurement.get("missing_keywords") or []
    unsupported = keyword_plan.get("missing_evidence") or []
    actionable_missing = [
        kw for kw in missing
        if kw not in unsupported and (
            _term_present(kw, " ".join(keyword_plan.get("required_keywords") or []))
            or _term_present(kw, cv_content)
        )
    ]

    fixes = []
    lines = [
        f"TARGETED ATS REPAIR PASS {pass_number}:",
        "Revise the resume below, preserving truthful facts and all original employers, job titles, dates, degrees, and contact details.",
    ]
    if actionable_missing:
        fixes.append("Added supported JD keywords missing from generated CV")
        lines.append(
            "Naturally add these supported or candidate-verified JD terms where they fit the summary, KEY SKILLS, WORK EXPERIENCE, or PROJECTS: "
            + ", ".join(_unique_preserve_order(actionable_missing, limit=10))
        )
    if measurement.get("keyword_match") is not None and int(measurement.get("keyword_match") or 0) < 90:
        fixes.append("Improved measured keyword match")
        lines.append("Increase exact JD terminology coverage without keyword stuffing.")
    if measurement.get("score") is not None and int(measurement.get("score") or 0) < 85:
        fixes.append("Improved measured ATS score")
        lines.append("Strengthen ATS score by improving relevance, quantification, structure, and supported keyword placement.")
    if quality.get("issues"):
        fixes.append("Fixed quality gate issues")
        lines.append("Fix these quality issues: " + "; ".join(str(i) for i in quality.get("issues", [])[:6]))
    lines.append(
        "If a JD term has missing evidence, do not claim it as experience. Leave it out or phrase it only as an honest transferable-adjacent strength supported by the resume."
    )
    lines.append("\nCURRENT GENERATED RESUME TO REPAIR:\n" + cv_content)
    return "\n".join(lines), _unique_preserve_order(fixes)


def _needs_repair(cv_content: str, job_description: str, measurement: dict[str, Any]) -> bool:
    quality = validate_cv_quality(cv_content, job_description)
    score = _to_int_or_none(measurement.get("score")) or 0
    keyword_match = _to_int_or_none(measurement.get("keyword_match")) or 0
    return (
        quality.get("should_regenerate", False)
        or score < 85
        or keyword_match < 90
        or bool(measurement.get("missing_keywords"))
    )


def generate_cv(
    resume_text,
    job_description,
    target_match=100,
    template="professional",
    sections=None,
    quantitative_focus=60,
    action_verb_intensity="High",
    keyword_matching="Balanced",
    language="English",
    model_choice="premium",
    extra_context="",
    optimization_depth: str = "max_ats",
    return_metadata: bool = False,
    **kwargs
):
    """Generate a CV through a measured Max ATS optimization loop.

    Backward compatible: returns plain text unless `return_metadata=True`.
    """
    target_score = int(target_match or 100)
    keyword_plan = build_jd_keyword_plan(resume_text, job_description, extra_context)
    max_repairs = 2 if str(optimization_depth or "max_ats").lower() in {"max_ats", "max", "maximum"} else 1

    cv_content = _generate_cv_once(
        resume_text=resume_text,
        job_description=job_description,
        target_match=target_score,
        template=template,
        sections=sections,
        quantitative_focus=quantitative_focus,
        action_verb_intensity=action_verb_intensity,
        keyword_matching=keyword_matching,
        language=language,
        model_choice=model_choice,
        extra_context=extra_context,
        keyword_plan=keyword_plan,
        **kwargs,
    )

    fixes_applied: list[str] = []
    repair_passes = 0
    measurement = _measure_cv_against_jd(cv_content, job_description)

    for pass_number in range(1, max_repairs + 1):
        if not _needs_repair(cv_content, job_description, measurement):
            break
        repair_instructions, fixes = _build_repair_instructions(
            cv_content=cv_content,
            job_description=job_description,
            keyword_plan=keyword_plan,
            measurement=measurement,
            pass_number=pass_number,
        )
        if not fixes and pass_number > 1:
            break
        cv_content = _generate_cv_once(
            resume_text=resume_text,
            job_description=job_description,
            target_match=target_score,
            template=template,
            sections=sections,
            quantitative_focus=quantitative_focus,
            action_verb_intensity=action_verb_intensity,
            keyword_matching="Aggressive",
            language=language,
            model_choice=model_choice,
            extra_context=extra_context,
            keyword_plan=keyword_plan,
            repair_instructions=repair_instructions,
            **kwargs,
        )
        repair_passes = pass_number
        fixes_applied.extend(fixes)
        measurement = _measure_cv_against_jd(cv_content, job_description)

    result = CVOptimizationResult(
        optimized_content=cv_content,
        ats_score=_to_int_or_none(measurement.get("score")),
        keyword_match=_to_int_or_none(measurement.get("keyword_match")),
        missing_keywords=measurement.get("missing_keywords") or [],
        repair_passes_used=repair_passes,
        fixes_applied=_unique_preserve_order(fixes_applied),
        unsupported_gaps=keyword_plan.get("missing_evidence") or [],
        target_ats_score=target_score,
    )

    if return_metadata:
        return result.model_dump() if hasattr(result, "model_dump") else result.dict()
    return result.optimized_content


def generate_cover_letter(resume_text, job_description, language="English", extra_context=""):
    """Generate cover letter using Gemini AI"""
    verified_block = build_verified_context_block(extra_context)


    language_instruction = f"""
    Respond ONLY in {language}.
    Generate the entire cover letter and all section headers in {language}.
    Use native {language} formatting, salutations, and polite forms appropriate for professional letters
    (e.g., "Madame, Monsieur" or "Monsieur/Madame" for French depending on recipient gender conventions).
    """
    
    prompt = f"""
    You are an expert ATS-optimized cover letter writer.
    
    Objective:
    Generate a personalized, professional cover letter that achieves **90%+ ATS compatibility** and aligns precisely with the provided Job Description.
    
    Rules:
    - Dynamically adjust keyword placement to ensure **high ATS score**.
    - Start with: “Hello Hiring Manager,” and include the line: “I am applying for the [exact job title] position.”
    - Use a tone that reflects professionalism and enthusiasm.
    
    Structure:
    1. **Paragraph 1**: Express genuine enthusiasm using the company's mission and JD language. Include company-specific values, vision, and relevant projects to show personalization.
    2. **Paragraph 2**: Align with the top 5 responsibilities in the JD. Provide **metrics-rich accomplishments** from the resume that demonstrate capability. Integrate at least **10 relevant keywords** from the JD (e.g., Python, SQL, machine learning, A/B testing, scikit-learn, AWS, Snowflake, dashboards, Spark, data-driven decision-making).
    3. **Paragraph 3**: Highlight **2-3 JD outcome-based goals** (e.g., predictive models, actionable insights, collaboration with cross-functional teams) using similar phrasing and past success examples. Include any statistical analysis, testing, or ML exposure.
    4. **Paragraph 4**: Reaffirm **2 key JD priorities**. Close by offering measurable value in JD terms. Request an interview and include a polite sign-off.
    
    Additional Requirements:
    - Use **identical terminology from the JD** wherever possible (e.g., "predictive models," "statistical analyses," "machine learning frameworks").
    - Mention **preferred skills** if applicable (e.g., AWS, Snowflake, Power BI).
    - Keep tone formal yet engaging, max 4 paragraphs.
    - After sign-off, include candidate's email and phone number (extract from resume).
    
    Inputs:
    Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Output:
    Generate the final cover letter in **plain text** format without extra commentary.
    """

    prompt_2 = f"""
    You are an expert ATS-optimized cover letter writer.
    {language_instruction}

    {TRUTHFULNESS_GUARDRAIL}

    {verified_block}

    1. Alignment & Conciseness Prompt
    “Generate a cover letter tightly aligned with the provided job description. Mention the exact job title, reference 2–3 JD responsibilities, and avoid repetition by summarizing key points once. Keep length to ≤ one A4 page.”

    2. ATS Keyword Enrichment Prompt
    “Extract 20–25 top ATS keywords from the JD (skills, tools, responsibilities) and integrate them naturally throughout the cover letter to maximize ATS relevance without keyword stuffing.”

    3. Measurable Achievements Prompt
    “Highlight 7–8 distinct, quantifiable achievements relevant to the JD (e.g., learner engagement %, performance improvements, cost/time savings). Spread them across the middle paragraphs for impact.”

    4. Portfolio & Link Integration Prompt
    “If the candidate has a portfolio, authored book, GitHub, Amazon, or research work, summarize its relevance in one concise sentence and provide a short, clickable hyperlink (≤100 characters) in the second or closing paragraph.”

    5. Non-Repetitive Value Demonstration Prompt
    “Do not duplicate CV bullets. Instead, showcase how the candidate’s achievements map directly to major JD requirements with unique, outcome-driven examples.”

    6. Storytelling with Impact Prompt
    “Include one brief, compelling success story (4–5 lines para) that demonstrates problem-solving or impact directly tied to the role’s core functions.”

    7. Employer Mission & Culture Link Prompt
    “Reference at least one element from the employer’s JD, mission, or values (e.g., innovation, inclusivity, customer focus) and connect it to the candidate’s philosophy or past work.”

    8. Tone & Culture Match Prompt
    “Adopt a professional yet personable tone that reflects the company’s culture as conveyed in the JD, ensuring recruiter resonance.”

    9. Call-to-Action & Closing Prompt
    “Conclude with a confident call-to-action: express enthusiasm, willingness to discuss contributions further, restate contact details, and thank the recruiter.”

    10. Formatting & Flow Prompt
    “Maintain ATS-friendly formatting: single font, 1-inch margins, left alignment. Use 3–4 paragraphs (Intro, Achievements/Portfolio, JD Alignment, Closing). Keep sentences ≤20 words. Ensure portfolio/research links are clean and clickable.”
    
    Objective:
    Generate a personalized, professional cover letter that achieves **90%+ ATS compatibility** and aligns precisely with the provided Job Description.
    
    Rules:
    - Dynamically adjust keyword placement to ensure **high ATS score**.
    - Start with: “Hello Hiring Manager,” and include the line: “I am applying for the [exact job title] position.”
    - Use a tone that reflects professionalism and enthusiasm.
    
    Structure:
    1. **Paragraph 1**: Express genuine enthusiasm using the company's mission and JD language. Include company-specific values, vision, and relevant projects to show personalization.
    2. **Paragraph 2**: Align with the top 5 responsibilities in the JD. Provide **metrics-rich accomplishments** from the resume that demonstrate capability. Integrate at least **10 relevant keywords** from the JD (e.g., Python, SQL, machine learning, A/B testing, scikit-learn, AWS, Snowflake, dashboards, Spark, data-driven decision-making).
    3. **Paragraph 3**: Highlight **2-3 JD outcome-based goals** (e.g., predictive models, actionable insights, collaboration with cross-functional teams) using similar phrasing and past success examples. Include any statistical analysis, testing, or ML exposure.
    4. **Paragraph 4**: Reaffirm **2 key JD priorities**. Close by offering measurable value in JD terms. Request an interview and include a polite sign-off.
    
    Additional Requirements:
    - Use **identical terminology from the JD** wherever possible (e.g., "predictive models," "statistical analyses," "machine learning frameworks").
    - Mention **preferred skills** if applicable (e.g., AWS, Snowflake, Power BI).
    - Keep tone formal yet engaging, max 4 paragraphs.
    - After sign-off, include candidate's email and phone number (extract from resume).
    
    Inputs:
    Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Output:
    Generate the final cover letter in **plain text** format without extra commentary.
    """

    try:
        if _get_session_ai_model() == "openai":
            response = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are a professional cover letter writer."},
                    {"role": "user", "content": prompt_2}
                ],
                temperature=0.2
            )
            content = response.choices[0].message.content
            cover_letter = content if content is not None else ""
            return re.sub(r'\*{1,2}', '', cover_letter).strip()

        else:

            response = model.generate_content(
            contents=prompt_2,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "text/plain"
            }
        )
        if not response or not response.text:
            raise Exception("AI response was empty")

        return re.sub(r'\*{1,2}', '', response.text).strip()

    except Exception as e:
        raise Exception(f"Failed to generate cover letter: {str(e)}")

def clean_cv_content(content):
    """Clean and format CV content for standard ATS parsing"""
    if not content:
        return "Error: No content received from AI"
    
    # Strip markdown backticks if AI wrapped the entire response
    content = re.sub(r'^```[a-zA-Z]*\n', '', content)
    content = re.sub(r'\n```$', '', content)

    # Remove prompt comment leaks (e.g. # Make sure NAME...)
    content = re.sub(r'^\s*#.*$', '', content, flags=re.MULTILINE)

    # Strip parenthetical suffixes from section headers (e.g. PROJECTS:(if any) -> PROJECTS:)
    content = re.sub(r'^([A-Z\s]+):\s*\([^)]*\)', r'\1:', content, flags=re.MULTILINE)

    # Remove markdown bold/italic decorators inside text body for clean text parsing
    content = re.sub(r'\*{1,2}', '', content)
    content = re.sub(r'__', '', content)
    
    # Remove excessive blank lines
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # Remove hidden HTML comment markers
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    # Remove unwanted markdown horizontal rule separators
    content = re.sub(r'^\s*---\s*$', '', content, flags=re.MULTILINE)
    
    # Ensure proper section header spacing
    content = re.sub(r'^([A-Z][A-Z\s]+):', r'\n\1:', content, flags=re.MULTILINE)

    # Ensure clean spacing around bullets
    lines = content.splitlines()
    cleaned_lines = []
    for ln in lines:
        stripped = ln.strip()
        if stripped:
            cleaned_lines.append(ln)
        elif cleaned_lines and cleaned_lines[-1] != "":
            cleaned_lines.append("")

    return "\n".join(cleaned_lines).strip()

def analyze_cv_ats_score(cv_content, job_description=""):
    """Analyze CV ATS compatibility score using Gemini AI with caching.
    
    If job_description is provided, performs a target JD match analysis.
    If job_description is empty or missing, performs a baseline ATS structural and formatting check.
    """
    has_jd = bool(job_description and str(job_description).strip())
    
    # Check cache first
    cache_key = _get_cache_key(cv_content, job_description)
    if cache_key in _ats_cache:
        entry = _ats_cache[cache_key]
        if time.time() - entry["timestamp"] < _ats_cache_ttl:
            return entry["result"]
    
    if has_jd:
        prompt = f"""
    You are an ATS analysis expert. Score the CV against the job description using this exact weighted rubric:
    
    **SCORING RUBRIC (Total: 100 points)**
    
    1. KEYWORD PRESENCE (30 points):
       - Are the top 10-15 critical JD keywords explicitly present in the CV?
       - Are both full terms and acronyms included (e.g., "Continuous Integration/Continuous Deployment (CI/CD)")?
       - Deduct points for missing critical hard skills.
    
    2. EXPERIENCE RELEVANCE (25 points):
       - Do work experience bullet points demonstrate the skills required by the JD?
       - Are achievements relevant to the target role?
       - Are there specific, measurable outcomes (%, numbers, performance gains)?
    
    3. QUANTIFICATION (15 points):
       - Are there specific numbers, percentages, or metrics?
       - Are achievements measurable (e.g., "reduced processing time by 40%", "managed $2M budget")?
       - Deduct points for generic claims without evidence.
    
    4. FORMAT COMPLIANCE (15 points):
       - Standard section headers (PROFESSIONAL SUMMARY, KEY SKILLS, WORK EXPERIENCE, EDUCATION)?
       - ATS-parseable bullet points ('• ' format)?
       - Clean contact information structure?
       - No markdown, code fences, or non-ATS elements?
    
    5. PROFESSIONAL NARRATIVE (15 points):
       - Is the content readable and coherent?
       - Does it tell a compelling career story?
       - Are there varied sentence structures (not repetitive keyword stuffing)?
       - Does it sound authentic (like a real person describing accomplishments)?
    
    Calculate the total score and return:
    {{
        "ats_score": total_score,
        "keyword_match": percentage_ofJD_keywords_found_in_CV,
        "missing_keywords": [list of critical missing keywords],
        "suggestions": [list of specific actionable suggestions based on rubric gaps]
    }}
    
    CV Content:
    {cv_content}
    
    Job Description:
    {job_description}
    """
    else:
        prompt = f"""
    You are an ATS analysis expert. No specific job description was provided.
    Score the CV using this baseline structural rubric:
    
    **SCORING RUBRIC (Total: 100 points)**
    
    1. STRUCTURAL COMPLETENESS (30 points):
       - Has all standard sections: PROFESSIONAL SUMMARY, KEY SKILLS, WORK EXPERIENCE, EDUCATION?
       - Contact information present (name, phone, email, location)?
       - Optional sections present (PROJECTS, CERTIFICATIONS)?
    
    2. KEYWORD ORGANIZATION (25 points):
       - KEY SKILLS section lists technical skills, tools, and competencies clearly?
       - Skills are categorized logically (Technical Skills, Tools & Platforms, Core Competencies)?
       - Industry-relevant keywords are present?
    
    3. EXPERIENCE QUALITY (20 points):
       - Work experience uses strong action verbs?
       - Achievements are quantified with specific metrics?
       - Bullet points follow ATS-compatible format ('• ' prefix)?
    
    4. ATS PARSEABILITY (15 points):
       - Clean formatting without markdown, tables, or graphics?
       - Consistent section headers in uppercase?
       - Proper spacing and line breaks?
    
    5. PROFESSIONAL IMPACT (10 points):
       - Professional summary is concise and compelling?
       - Content flows logically from summary → skills → experience → education?
       - Overall impression is professional and well-organized?
    
    Calculate the total score and return:
    {{
        "ats_score": total_score,
        "keyword_match": overall_keyword_structural_match_percentage,
        "missing_keywords": [list of missing standard sections or essential skills],
        "suggestions": [list of specific actionable suggestions based on rubric gaps]
    }}
    
    CV Content:
    {cv_content}
    """
    
    try:
        if _get_session_ai_model() == "openai":
            # ✅ GPT-based analysis
            response = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are an ATS scoring and resume optimization expert."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0  # Keep it deterministic for scoring
            )
            content = response.choices[0].message.content
            raw_text = content.strip() if content is not None else ""

        else:
            # ✅ Gemini-based analysis
            response = model.generate_content(
                contents=prompt,
                generation_config={
                    "temperature": 0.0,
                    "response_mime_type": "application/json"
                }
            )
            raw_text = response.text.strip()

        # Clean markdown code blocks if present
        if raw_text.startswith("```"):
            raw_text = re.sub(r"^```[a-zA-Z]*\s*", "", raw_text)
            raw_text = re.sub(r"\s*```$", "", raw_text).strip()

        # ✅ Parse JSON
        try:
            parsed = json.loads(raw_text)
        except Exception as parse_err:
            raise Exception(f"Invalid JSON response: {raw_text}")

        # Coerce LLM-returned values defensively: the prompt asks for numbers but
        # Gemini/OpenAI sometimes returns them as strings like "67" or "67%".
        # Without coercion, downstream `float(value)` calls crash on strings.
        result = {
            "score": _to_int_or_none(parsed.get("ats_score")) or 0,
            "keyword_match": _to_int_or_none(parsed.get("keyword_match")) or 0,
            "missing_keywords": parsed.get("missing_keywords", []) or [],
            "suggestions": parsed.get("suggestions", []) or []
        }
        
        # Cache result
        _ats_cache[cache_key] = {
            "result": result,
            "timestamp": time.time()
        }
        
        return result

    except Exception as e:
        return {
            "score": None,
            "keyword_match": None,
            "missing_keywords": [],
            "suggestions": [f"Error analyzing CV: {str(e)}"]
        }

def extract_key_metrics(cv_content):
    """Extract quantifiable metrics from CV"""
    # Pattern to find numbers and percentages
    metrics_pattern = r'(\d+(?:\.\d+)?(?:%|K|M|B|k|m|b|\+|,\d+)*)'
    
    metrics = re.findall(metrics_pattern, cv_content)
    
    return {
        'total_metrics': len(metrics),
        'metrics_found': metrics,
        'quantification_score': min(100, len(metrics) * 5)  # 5 points per metric, max 100
    }

def enhance_action_verbs(content, intensity="High"):
    """Enhance action verbs in CV content with context-aware replacement"""
    
    # Weak verb patterns (multi-word included)
    weak_patterns = {
        "worked on": "developed",
        "did": "executed",
        "made": "created",
        "helped": "collaborated on",
        "was responsible for": "managed",
        "handled": "oversaw",
        "participated in": "contributed to",
        "assisted with": "supported",
        "took part in": "engaged in",
        "was in charge of": "directed",
        "used": "utilized",
        "improved": "optimized",
    }
    
    # Professional replacements by intensity
    intensity_map = {
        "Moderate": {
            "developed": "developed",
            "created": "created",
            "managed": "managed",
            "improved": "enhanced",
        },
        "High": {
            "developed": "engineered",
            "created": "architected",
            "managed": "spearheaded",
            "improved": "optimized",
        },
        "Very High": {
            "developed": "pioneered",
            "created": "revolutionized",
            "managed": "orchestrated",
            "improved": "transformed",
        }
    }
    
    enhanced = content
    for weak, replacement in weak_patterns.items():
        # Case-insensitive replacement
        enhanced = re.sub(
            r'\b' + re.escape(weak) + r'\b',
            replacement,
            enhanced,
            flags=re.IGNORECASE
        )
    
    return enhanced

def generate_interview_qa(resume_text, job_description, extra_context: Any = ""):
    """Generate interview Q&A using Gemini AI"""
    verified_block = build_verified_context_block(extra_context)
    prompt = f"""
    You are “Ultra-Strict JD Full-Coverage Practice Pack (Non-Interactive)”.

    {TRUTHFULNESS_GUARDRAIL}

    {verified_block}




    ## QUESTION BREAKDOWN:
    1. Tell us something about yourself?
    2. Why are you applying for this job?
    - **8 Behavioral Questions**
        - Focus on: leadership, teamwork, conflict resolution, problem-solving, adaptability, failure/learning, achievement, cultural fit
        - Use starters like:
        - "Tell me about a time when..."
        - "Describe a situation where..."
    - **12 Technical Questions**
        - Role-specific and scenario-based
        - Include advanced problem-solving situations
        - Cover tools, frameworks, and methodologies from the job description and resume
        

    ## ANSWER FORMAT:
    
    1. Tell us something about yourself?
    2. Why are you applying for this job?

    Behavioral Questions

    1. Question line....
    Situation: Brief context and background
    Task: Specific challenge or responsibility
    Action: Key steps you took(Atleast 5-6 Steps)
        Step1:
        Step2:
        .....
    Result: Positive outcome with measurable impact (include metrics where possible)

    Technical Questions

    1. Question line.....
    Situation: Technical context or project background
    Task: Technical challenge or requirement
    Action: Detailed technical approach, tools, and methods implemented(Atleast 5-6 actions)
        Step1:
        Step2:
        .....
    Result: Technical outcomes, improvements, or measurable impact

    ✅ **Response Length:** Each answer should be concise yet complete (approx. **200-250 words**).

    ## ROLE-LEVEL ADAPTATION:
    - If role is **Entry-Level**: Use simpler technical challenges and emphasize learning and adaptability
    - If role is **Mid-Level**: Include mix of technical depth and soft skills like collaboration and initiative
    - If role is **Senior/Lead**: Focus on leadership, architecture-level decisions, influencing stakeholders, and scaling solutions

    ## INPUTS:
    **Resume:**
    {resume_text}

    **Job Description:**
    {job_description}

    ## ADDITIONAL INSTRUCTIONS:
    - Prioritize critical skills and achievements from the resume
    - Ensure technical depth matches role level
    - Use industry-specific terminology and best practices
    - Keep answers **authentic, achievement-oriented, and concise**
    """

    try:
        if _get_session_ai_model() == "openai":
            response = openai.chat.completions.create(
                model="gpt-5",
                messages=[
                    {"role": "system", "content": "You are an expert career coach and interviewer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=1
            )
            return response.choices[0].message.content

        else:

            response = model.generate_content(
            contents=prompt,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "text/plain"
            }
        )
        if not response or not response.text:
            raise Exception("AI response was empty")

        return response.text

    except Exception as e:
        raise Exception(f"Failed to generate Q&A: {str(e)}")




def export_interview_qa(content):
    """Export Q&A with bold section headings, bold questions, and STAR keywords, robust for GPT & Gemini."""
    from io import BytesIO
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from docx import Document
    from docx.shared import Pt
    import re

    # ---------- Styles (PDF) ----------
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        'HeadingStyle', fontSize=16, leading=20, textColor=colors.darkblue,
        spaceAfter=12, fontName='Helvetica-Bold'
    )
    question_style = ParagraphStyle(
        'QuestionStyle', fontSize=13, leading=15, spaceAfter=8,
        fontName='Helvetica-Bold'
    )
    normal_style = ParagraphStyle(
        'NormalStyle', fontSize=11, leading=14, spaceAfter=6
    )

    # ---------- Helpers ----------
    def normalize_line(s: str) -> str:
        if not s: return ""
        s = s.strip()
        # skip pure separators like *** --- ___
        if re.fullmatch(r'(\*{3,}|-{3,}|_{3,})', s):
            return ""
        # strip leading markdown heading/bullet markers
        s = re.sub(r'^\s*(#{1,6}|\*{1,3}|-|\u2022|\u25CF)\s*', '', s)
        # unwrap **bold** around whole line
        s = re.sub(r'^\*\*(.+?)\*\*$','\\1', s)
        # collapse double spaces
        s = re.sub(r'\s{2,}', ' ', s)
        return s.strip()

    # Accept headings even if wrapped in ** or ##
    heading_re = re.compile(
        r'^\s*(?:\*{0,3}|#{0,3})\s*(Behavioral Questions|Technical Questions|Resume-based Questions|General Questions)\s*:?\s*(?:\*{0,3})?\s*$',
        re.IGNORECASE
    )

    # Accept STAR labels even if wrapped in ** and with variable spaces
    star_label_re = re.compile(
        r'^\s*(?:\*{0,3}|#{0,3})\s*(Situation|Task|Action|Result)\s*:\s*(.*)$',
        re.IGNORECASE
    )

    # Question lines can be "1. ...", "1) ...", and may have leading ###/** etc.
    question_re = re.compile(
        r'^\s*(?:#{1,6}|\*{1,3})?\s*\d+[\.\)]\s+.+'
    )

    # We'll avoid marking numbered items as questions if they appear immediately after "Action:"
    in_action_block = False

    # ---------- Build PDF ----------
    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer)
    story = []

    lines = [normalize_line(l) for l in content.split('\n')]
    for raw in lines:
        line = raw.strip()
        if not line:
            # leaving an empty line ends any action-list context
            in_action_block = False
            continue

        # Section headings
        m_heading = heading_re.match(line)
        if m_heading:
            title = m_heading.group(1)
            story.append(Paragraph(title, heading_style))
            story.append(Spacer(1, 6))
            in_action_block = False
            continue

        # STAR labels
        m_star = star_label_re.match(line)
        if m_star:
            label = m_star.group(1).title()
            rest  = m_star.group(2).strip()
            story.append(Paragraph(f"<b>{label}:</b> {rest}", normal_style))
            story.append(Spacer(1, 4))
            in_action_block = (label.lower() == "action")  # enter action-block
            continue

        # Numbered questions (guard against action-step lists)
        if question_re.match(line) and not in_action_block:
            story.append(Paragraph(line, question_style))
            story.append(Spacer(1, 4))
            continue

        # Any other paragraph
        story.append(Paragraph(line, normal_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    pdf_buffer.seek(0)

    # ---------- Build DOCX ----------
    docx_buffer = BytesIO()
    word_doc = Document()

    in_action_block = False  # reset state
    for raw in lines:
        line = raw.strip()
        if not line:
            in_action_block = False
            continue

        m_heading = heading_re.match(line)
        if m_heading:
            title = m_heading.group(1)
            p = word_doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(16)
            in_action_block = False
            continue

        m_star = star_label_re.match(line)
        if m_star:
            label = m_star.group(1).title()
            rest  = m_star.group(2).strip()
            p = word_doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            if rest:
                p.add_run(rest)
            in_action_block = (label.lower() == "action")
            continue

        if question_re.match(line) and not in_action_block:
            p = word_doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(13)
            continue

        word_doc.add_paragraph(line)

    word_doc.save(docx_buffer)
    docx_buffer.seek(0)

    return pdf_buffer, docx_buffer

def export_cover_letter(content):
    """Export cover letter to PDF and Word DOCX formats."""
    from io import BytesIO
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from docx import Document
    from docx.shared import Pt
    import re

    # ---------- Styles (PDF) ----------
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        'NormalStyle', fontSize=11, leading=15, spaceAfter=10
    )

    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer)
    story = []

    lines = content.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            story.append(Spacer(1, 10))
            continue
        story.append(Paragraph(line_str, normal_style))

    doc.build(story)
    pdf_buffer.seek(0)

    # ---------- DOCX ----------
    docx_buffer = BytesIO()
    word_doc = Document()
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
        word_doc.add_paragraph(line_str)
    
    word_doc.save(docx_buffer)
    docx_buffer.seek(0)

    return pdf_buffer, docx_buffer

# AFTER (fixed):
def recommend_jobs_from_resume_ai(resume_text: str, language: str = "English"):
    """Recommend job roles from resume using Gemini AI directly."""
    prompt = f"""
    Analyze the resume below and return EXACTLY 5 job titles
    the candidate is best suited for.

    Rules:
    - Return ONLY job titles
    - No explanation, no numbering text
    - No extra words
    - Each job title on a new line
    - Respond in {language}

    Resume:
    {resume_text}
    """
    try:
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.3,
                "max_output_tokens": 512
            }
        )
        # Safely extract text from response
        response_text = ""
        if response and response.candidates and len(response.candidates) > 0:
            candidate = response.candidates[0]
            if candidate.content and candidate.content.parts:
                for part in candidate.content.parts:
                    if hasattr(part, "text") and part.text:
                        response_text += part.text
        elif response and hasattr(response, "text") and response.text:
            response_text = response.text

        jobs = [line.strip() for line in response_text.split("\n") if line.strip()]
        # Remove any numbering artifacts like "1." "1)" "-" that the model might add
        cleaned_jobs = []
        for job in jobs:
            job = re.sub(r"^[\d\.\)\-\*]+\s*", "", job).strip()
            if job:
                cleaned_jobs.append(job)
        return cleaned_jobs[:5]

    except Exception as e:
        # Re-raise so app.py can catch it and show a proper error
        raise Exception(f"Job recommendation failed: {str(e)}")

def validate_cv_quality(cv_content: str, job_description: str = "") -> dict:
    """Validate CV quality before display. Returns dict with validation results."""
    from utils import extract_ats_phrases, validate_cv_format
    
    results = {
        "is_valid": True,
        "issues": [],
        "should_regenerate": False
    }
    
    if not cv_content or not cv_content.strip():
        results["is_valid"] = False
        results["issues"].append("CV content is empty")
        results["should_regenerate"] = True
        return results
    
    cv_lower = cv_content.lower()
    
    # Check 1: Required sections present
    required_sections = ["professional summary", "key skills", "work experience", "education"]
    missing_sections = []
    for section in required_sections:
        if section not in cv_lower:
            missing_sections.append(section.upper())
    
    if missing_sections:
        results["is_valid"] = False
        results["issues"].append(f"Missing required sections: {', '.join(missing_sections)}")
        results["should_regenerate"] = True
    
    # Check 2: KEY SKILLS section has sufficient keywords
    if job_description:
        jd_phrases = extract_ats_phrases(job_description)
        
        # Find KEY SKILLS section
        key_skills_match = re.search(r'key skills:?\s*\n(.*?)(?=\n[A-Z]+:|\n*$)', cv_content, re.IGNORECASE | re.DOTALL)
        if key_skills_match:
            key_skills_content = key_skills_match.group(1).lower()
            keywords_found = sum(1 for phrase in jd_phrases if phrase.lower() in key_skills_content)
            keyword_minimum = min(10, max(3, len(jd_phrases)))
            
            if keywords_found < keyword_minimum:
                results["issues"].append(f"KEY SKILLS section has only {keywords_found} JD keywords (minimum: {keyword_minimum})")
                # Don't force regeneration for this - just warn
        else:
            results["issues"].append("Could not find KEY SKILLS section content")
    
    # Check 3: Professional summary length
    summary_match = re.search(r'professional summary:?\s*\n(.*?)(?=\n[A-Z]+:|\n*$)', cv_content, re.IGNORECASE | re.DOTALL)
    if summary_match:
        summary_content = summary_match.group(1).strip()
        word_count = len(summary_content.split())
        
        if word_count > 100:
            results["issues"].append(f"Professional summary is {word_count} words (recommended: under 100)")
            # Don't force regeneration - just warn
    
    # Check 4: Basic format validation
    format_validation = validate_cv_format(cv_content)
    if not format_validation.get("valid", False):
        results["issues"].append(f"Format issues: {format_validation.get('issues', [])}")

    # Check 5: ATS-clean output only
    if "```" in cv_content or re.search(r"^\s*#{1,6}\s+", cv_content, re.MULTILINE):
        results["is_valid"] = False
        results["issues"].append("Markdown/code-fence formatting detected")
        results["should_regenerate"] = True

    # Check 6: Bullet structure and repetition
    bullet_lines = [line.strip() for line in cv_content.splitlines() if line.strip().startswith("•")]
    if len(bullet_lines) < 6:
        results["is_valid"] = False
        results["issues"].append(f"Only {len(bullet_lines)} ATS bullet points found (minimum: 6)")
        results["should_regenerate"] = True
    normalized_bullets = [re.sub(r"\s+", " ", b.lower()) for b in bullet_lines]
    duplicate_count = len(normalized_bullets) - len(set(normalized_bullets))
    if duplicate_count:
        results["issues"].append(f"{duplicate_count} repeated bullet point(s) detected")

    # Check 7: simple keyword stuffing detection
    words = re.findall(r"\b[a-zA-Z][a-zA-Z0-9\+#\.\-]{2,}\b", cv_lower)
    if words:
        total = len(words)
        counts = {}
        for word in words:
            if word in {"and", "the", "with", "for", "from", "that", "this"}:
                continue
            counts[word] = counts.get(word, 0) + 1
        overused = [w for w, c in counts.items() if c >= 10 and (c / total) > 0.035]
        if overused:
            results["issues"].append("Possible keyword stuffing: " + ", ".join(overused[:5]))
    
    # Check 8: Minimum content length
    non_empty_lines = [line for line in cv_content.split('\n') if line.strip()]
    if len(non_empty_lines) < 20:
        results["is_valid"] = False
        results["issues"].append(f"CV has only {len(non_empty_lines)} non-empty lines (minimum: 20)")
        results["should_regenerate"] = True
    
    return results
